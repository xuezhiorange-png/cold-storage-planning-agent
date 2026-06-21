"""Knowledge parser tests — text, markdown, CSV, DOCX, XLSX, PDF parsers.

All test fixtures are generated in code (small strings/bytes), NOT from external files.

Covers 14 parser properties:
 1. txt UTF-8 basic parsing
 2. txt BOM handling
 3. txt encoding error
 4. markdown heading path
 5. markdown code block as separate block
 6. CSV row ranges
 7. CSV headers captured
 8. DOCX paragraphs extracted
 9. DOCX tables extracted
10. XLSX multi-sheet
11. XLSX formula not executed
12. PDF multi-page text
13. PDF empty scan (requires_ocr detection)
14. Unsupported format rejected
"""

from __future__ import annotations

import io

import pytest

from cold_storage.modules.knowledge.infrastructure.parsers.base import get_parser
from cold_storage.modules.knowledge.infrastructure.parsers.csv_parser import CsvParser
from cold_storage.modules.knowledge.infrastructure.parsers.markdown_parser import MarkdownParser
from cold_storage.modules.knowledge.infrastructure.parsers.text_parser import TextParser

# ---------------------------------------------------------------------------
# 1-3. Text parser tests
# ---------------------------------------------------------------------------


class TestTextParser:
    def test_txt_utf8(self) -> None:
        """Basic UTF-8 text file is parsed into paragraph blocks."""
        parser = TextParser()
        content = b"Hello world.\n\nThis is a second paragraph."
        blocks = parser.parse(content, "test.txt")
        assert len(blocks) >= 1
        assert blocks[0].block_type == "paragraph"
        assert "Hello world" in blocks[0].text

    def test_txt_bom(self) -> None:
        """UTF-8 BOM is handled and stripped."""
        parser = TextParser()
        content = b"\xef\xbb\xbfHello with BOM"
        blocks = parser.parse(content, "bom.txt")
        assert len(blocks) >= 1
        assert "Hello with BOM" in blocks[0].text
        # BOM should not appear in text
        assert "\ufeff" not in blocks[0].text

    def test_txt_encoding_error(self) -> None:
        """Invalid bytes that are not valid UTF-8 raises an error."""
        parser = TextParser()
        content = b"\x80\x81\x82\x83"
        with pytest.raises(UnicodeDecodeError):
            parser.parse(content, "bad.txt")


# ---------------------------------------------------------------------------
# 4-5. Markdown parser tests
# ---------------------------------------------------------------------------


class TestMarkdownParser:
    def test_markdown_heading_path(self) -> None:
        """Section path correctly tracks heading hierarchy."""
        parser = MarkdownParser()
        md = "# Title\n## Section A\nContent under A\n## Section B\nMore content"
        blocks = parser.parse(md.encode("utf-8"), "doc.md")
        # Find the paragraph blocks (not headings)
        paragraphs = [b for b in blocks if b.block_type == "paragraph"]
        assert len(paragraphs) >= 1
        # The first paragraph should be under "Title > Section A"
        assert "Title" in paragraphs[0].section_path
        assert "Section A" in paragraphs[0].section_path

    def test_markdown_code_block(self) -> None:
        """Fenced code blocks are parsed as separate 'code' blocks."""
        parser = MarkdownParser()
        md = "# Title\nSome text\n```python\nprint('hello')\n```\nAfter code"
        blocks = parser.parse(md.encode("utf-8"), "code.md")
        code_blocks = [b for b in blocks if b.block_type == "code"]
        assert len(code_blocks) >= 1
        assert "print" in code_blocks[0].text


# ---------------------------------------------------------------------------
# 6-7. CSV parser tests
# ---------------------------------------------------------------------------


class TestCsvParser:
    def test_csv_rows(self) -> None:
        """Row ranges are correctly recorded for each data row."""
        parser = CsvParser()
        csv_content = "Name,Value\nAlpha,100\nBeta,200\nGamma,300\n"
        blocks = parser.parse(csv_content.encode("utf-8"), "data.csv")
        # First block is metadata (headers), rest are data rows
        data_blocks = [b for b in blocks if b.block_type == "paragraph"]
        assert len(data_blocks) == 3
        # Row ranges should be sequential
        for i, block in enumerate(data_blocks):
            assert block.row_start == i + 2  # header is row 1
            assert block.row_end == i + 2

    def test_csv_headers(self) -> None:
        """Headers are captured in a metadata block."""
        parser = CsvParser()
        csv_content = "Name,Value,Unit\nAlpha,100,kg\n"
        blocks = parser.parse(csv_content.encode("utf-8"), "headers.csv")
        metadata_blocks = [b for b in blocks if b.block_type == "metadata"]
        assert len(metadata_blocks) == 1
        assert "Name" in metadata_blocks[0].text
        assert "Value" in metadata_blocks[0].text
        assert "Unit" in metadata_blocks[0].text


# ---------------------------------------------------------------------------
# 8-9. DOCX parser tests
# ---------------------------------------------------------------------------


class TestDocxParser:
    def test_docx_paragraphs(self) -> None:
        """Paragraphs are extracted from a minimal DOCX file."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from cold_storage.modules.knowledge.infrastructure.parsers.docx_parser import DocxParser

        parser = DocxParser()
        blocks = parser.parse(buf.read(), "test.docx")
        text_blocks = [b for b in blocks if b.block_type in ("paragraph", "heading")]
        assert len(text_blocks) >= 2
        texts = " ".join(b.text for b in text_blocks)
        assert "First paragraph" in texts
        assert "Second paragraph" in texts

    def test_docx_tables(self) -> None:
        """Tables are extracted as 'table' blocks."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Before table")
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(0, 2).text = "C"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        table.cell(1, 2).text = "3"
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        from cold_storage.modules.knowledge.infrastructure.parsers.docx_parser import DocxParser

        parser = DocxParser()
        blocks = parser.parse(buf.read(), "table.docx")
        table_blocks = [b for b in blocks if b.block_type == "table"]
        assert len(table_blocks) >= 1
        assert "A" in table_blocks[0].text
        assert "1" in table_blocks[0].text


# ---------------------------------------------------------------------------
# 10-11. XLSX parser tests
# ---------------------------------------------------------------------------


class TestXlsxParser:
    def test_xlsx_multi_sheet(self) -> None:
        """Multiple sheets are parsed with correct sheet_name metadata."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1.append(["Name", "Value"])
        ws1.append(["A", 100])
        ws2 = wb.create_sheet("Sheet2")
        ws2.append(["Item", "Count"])
        ws2.append(["X", 5])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        from cold_storage.modules.knowledge.infrastructure.parsers.xlsx_parser import XlsxParser

        parser = XlsxParser()
        blocks = parser.parse(buf.read(), "multi.xlsx")
        sheet_names = {b.sheet_name for b in blocks if b.sheet_name is not None}
        assert "Sheet1" in sheet_names
        assert "Sheet2" in sheet_names

    def test_xlsx_formula_not_executed(self) -> None:
        """Formula text is preserved as-is (openpyxl data_only=False)."""
        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        ws = wb.active
        ws.append(["A", "B", "Formula"])
        ws.append([10, 20, "=A1+B1"])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        from cold_storage.modules.knowledge.infrastructure.parsers.xlsx_parser import XlsxParser

        parser = XlsxParser()
        blocks = parser.parse(buf.read(), "formula.xlsx")
        all_text = " ".join(b.text for b in blocks)
        # Formula should be present as text, not executed value
        assert "=A1+B1" in all_text


# ---------------------------------------------------------------------------
# 12-13. PDF parser tests
# ---------------------------------------------------------------------------


class TestPdfParser:
    def test_pdf_multipage(self) -> None:
        """Multi-page PDF produces blocks with correct page numbers."""
        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        doc = pymupdf.open()
        doc.new_page()
        page0 = doc[0]
        page0.insert_text((72, 72), "Page one content here with enough text to be meaningful")

        doc.new_page()
        page1 = doc[1]
        page1.insert_text((72, 72), "Page two content here with enough text to be meaningful")

        pdf_bytes = doc.tobytes()
        doc.close()

        from cold_storage.modules.knowledge.infrastructure.parsers.pdf_parser import PdfParser

        parser = PdfParser()
        blocks = parser.parse(pdf_bytes, "multi.pdf")
        pages = {b.page_start for b in blocks if b.page_start is not None}
        assert 1 in pages
        assert 2 in pages

    def test_pdf_empty_scan(self) -> None:
        """Empty/image-only PDF flags requires_ocr detection."""
        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        # Create a PDF with an empty page (no text, no images)
        doc = pymupdf.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        from cold_storage.modules.knowledge.infrastructure.parsers.pdf_parser import PdfParser

        parser = PdfParser()
        result = parser.detect_ocr_needed(pdf_bytes)
        # Empty page has no text and no images — avg text is 0
        assert result is True


# ---------------------------------------------------------------------------
# 14. Unsupported format test
# ---------------------------------------------------------------------------


class TestUnsupportedFormat:
    def test_unsupported_format(self) -> None:
        """No parser registered for .zip returns None."""
        parser = get_parser(".zip")
        assert parser is None


# ---------------------------------------------------------------------------
# 15-16. PDF parser statelessness tests
# ---------------------------------------------------------------------------


class TestPdfParserStatelessness:
    def test_pdf_parse_with_metadata_returns_result(self) -> None:
        """parse_with_metadata returns a ParseResult with correct fields."""
        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        from cold_storage.modules.knowledge.infrastructure.parsers.base import ParseResult
        from cold_storage.modules.knowledge.infrastructure.parsers.pdf_parser import PdfParser

        doc = pymupdf.open()
        doc.new_page()
        page = doc[0]
        page.insert_text((72, 72), "Test content for parse result verification")
        pdf_bytes = doc.tobytes()
        doc.close()

        parser = PdfParser()
        result = parser.parse_with_metadata(pdf_bytes, "test.pdf")

        assert isinstance(result, ParseResult)
        assert isinstance(result.blocks, list)
        assert len(result.blocks) >= 1
        assert result.page_count == 1
        assert isinstance(result.warnings, list)
        assert isinstance(result.ocr_page_numbers, list)

    def test_pdf_no_shared_state(self) -> None:
        """parse_with_metadata leaves no shared state between calls."""
        try:
            import pymupdf
        except ImportError:
            pytest.skip("pymupdf not installed")

        from cold_storage.modules.knowledge.infrastructure.parsers.pdf_parser import PdfParser

        # First PDF
        doc1 = pymupdf.open()
        doc1.new_page()
        page1 = doc1[0]
        page1.insert_text((72, 72), "First document content alpha")
        pdf1 = doc1.tobytes()
        doc1.close()

        # Second PDF
        doc2 = pymupdf.open()
        doc2.new_page()
        page2 = doc2[0]
        page2.insert_text((72, 72), "Second document content beta")
        pdf2 = doc2.tobytes()
        doc2.close()

        parser = PdfParser()
        result1 = parser.parse_with_metadata(pdf1, "first.pdf")
        result2 = parser.parse_with_metadata(pdf2, "second.pdf")

        # No state leakage: results are independent
        assert result1.page_count == 1
        assert result2.page_count == 1
        assert result1.blocks is not result2.blocks
        assert result1.blocks[0].text != result2.blocks[0].text
        assert "alpha" in result1.blocks[0].text
        assert "beta" in result2.blocks[0].text
