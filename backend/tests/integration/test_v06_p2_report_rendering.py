"""V0.6 P2 integration tests for report rendering and localization hardening."""

from __future__ import annotations

from io import BytesIO

import fitz
import pytest
from docx import Document

from cold_storage.modules.reports.application.canonical_render_model_builder import (
    build_canonical_render_model,
)
from cold_storage.modules.reports.application.render_model_localizer import localize_render_model
from cold_storage.modules.reports.domain.enums import ReportLocale
from cold_storage.modules.reports.localization.catalog import translate
from cold_storage.modules.reports.renderers.docx_renderer import DocxRenderer
from cold_storage.modules.reports.renderers.pdf_renderer import PdfRenderer


def _scheme_metric_row(
    table: object,
    scheme_name: str,
    metric_label: str,
) -> tuple:
    """Return the long-format row matching scheme name and metric label."""
    for row in table.rows:
        if row[0].display_value == scheme_name and row[1].display_value == metric_label:
            return row
    raise AssertionError(
        f"No row for scheme={scheme_name!r}, metric={metric_label!r}; "
        f"rows={[(r[0].display_value, r[1].display_value) for r in table.rows]!r}"
    )


def _build_p2_fixture_content() -> dict:
    return {
        "report_metadata": {
            "project_id": "p2-project",
            "schema_version": "cold_storage_concept_design@1.0.0",
        },
        "cooling_load": {
            "total_design_refrigeration_load": {
                "value": 250.0,
                "unit": "kW(r)",
                "source_result_id": "calc-001",
                "source_tool": "cooling_load_calculator",
                "source_tool_version": "2.1.0",
                "source_content_hash": "abc123hashvalue",
            },
        },
        "scheme_comparison": {
            "run_id": "run-001",
            "schemes": [
                {
                    "scheme_id": "s1",
                    "name": "Scheme Alpha",
                    "rank": 1,
                    "total_score": {
                        "value": "85.5",
                        "source_tool_version": "3.0.0",
                        "source_content_hash": "scorehash99",
                    },
                },
                {
                    "scheme_id": "s2",
                    "name": "Scheme Beta",
                    "rank": 2,
                    "total_score": "72.0",
                },
            ],
            "recommended_scheme": "s1",
        },
        "risks_and_missing_information": {
            "risks": [
                {
                    "description": "Risk about missing survey",
                    "severity": "high",
                    "mitigation": "Wait for survey",
                }
            ],
            "missing_information": [
                {
                    "description": "Site geotechnical survey not completed",
                    "impact": "foundation_design",
                }
            ],
        },
        "quality_summary": {
            "findings": [
                {
                    "code": "WARN_001",
                    "severity": "warning",
                    "message": "First distinct message",
                    "section_key": "cooling_load",
                    "field_path": "cooling_load.total",
                },
                {
                    "code": "WARN_001",
                    "severity": "warning",
                    "message": "Second distinct message",
                    "section_key": "equipment_selection",
                    "field_path": "equipment_selection.count",
                },
            ],
        },
    }


def _build_canonical() -> object:
    return build_canonical_render_model(
        content=_build_p2_fixture_content(),
        report_id="p2-report",
        revision_number=1,
        content_hash="a" * 64,
        generated_by="p2-test",
        generated_at="2025-01-01T00:00:00Z",
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
    )


def _docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _pdf_text(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text


def _pdf_scheme_row_y_bands(
    pdf_bytes: bytes,
    scheme_name: str,
    metric_label: str,
    value_token: str,
) -> list[int]:
    """Return y-band keys where scheme name, metric label, and value token co-occur."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    y_bands: dict[str, set[int]] = {"scheme": set(), "metric": set(), "value": set()}
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                y_key = int(line["bbox"][1] / 5.0)
                if scheme_name in line_text:
                    y_bands["scheme"].add(y_key)
                if metric_label in line_text:
                    y_bands["metric"].add(y_key)
                if value_token in line_text:
                    y_bands["value"].add(y_key)
    doc.close()
    shared = y_bands["scheme"] & y_bands["metric"] & y_bands["value"]
    return sorted(shared)


@pytest.mark.parametrize("locale", [ReportLocale.ZH_CN, ReportLocale.EN_US])
def test_p2_scheme_table_long_format_row_coordinates(locale: ReportLocale) -> None:
    """Scheme name, metric label, and value share one logical PDF row per scheme metric."""
    canonical = _build_canonical()
    model = localize_render_model(canonical, locale=locale)
    scheme_section = next(s for s in model.sections if s.section_key == "scheme_comparison")
    assert scheme_section.table is not None
    assert len(scheme_section.table.headers) == 3

    first_row = scheme_section.table.rows[0]
    scheme_name = first_row[0].display_value
    metric_label = first_row[1].display_value
    value_token = first_row[2].display_value.split(" ", 1)[0]

    pdf_bytes = PdfRenderer().render(model)
    shared_y = _pdf_scheme_row_y_bands(pdf_bytes, scheme_name, metric_label, value_token)
    assert shared_y, (
        f"scheme name, metric, and value must share PDF row coordinates; "
        f"scheme={scheme_name!r}, metric={metric_label!r}, value_token={value_token!r}"
    )


@pytest.mark.parametrize("locale", [ReportLocale.ZH_CN, ReportLocale.EN_US])
def test_p2_quality_findings_same_code_distinguishable(locale: ReportLocale) -> None:
    """Findings with the same code but different message/section render as separate rows."""
    canonical = _build_canonical()
    model = localize_render_model(canonical, locale=locale)
    quality_section = next(s for s in model.sections if s.section_key == "quality_summary")
    assert quality_section.table is not None
    assert len(quality_section.table.rows) == 2
    messages = [row[4].display_value for row in quality_section.table.rows]
    assert "First distinct message" in messages
    assert "Second distinct message" in messages

    docx_bytes = DocxRenderer().render(model)
    docx_text = _docx_text(docx_bytes)
    assert "First distinct message" in docx_text
    assert "Second distinct message" in docx_text

    pdf_text = _pdf_text(PdfRenderer().render(model))
    assert "First distinct" in pdf_text
    assert "Second distinct" in pdf_text


@pytest.mark.parametrize("locale", [ReportLocale.ZH_CN, ReportLocale.EN_US])
def test_p2_provenance_renders_when_present_and_omits_when_absent(locale: ReportLocale) -> None:
    """Per-source version/hash provenance appears only when present on the render model."""
    canonical = _build_canonical()
    model = localize_render_model(canonical, locale=locale)

    cooling_section = next(s for s in model.sections if s.section_key == "cooling_load")
    if cooling_section.metrics:
        cooling_display = cooling_section.metrics[0].display_value
    else:
        assert cooling_section.table is not None
        assert cooling_section.table.rows, "cooling_load compact table is empty"
        cooling_display = cooling_section.table.rows[0][1].display_value
    assert "2.1.0" in cooling_display
    assert "abc123hashvalue" in cooling_display

    scheme_section = next(s for s in model.sections if s.section_key == "scheme_comparison")
    total_score_label = translate(locale, "header.total_score")
    provenance_row = _scheme_metric_row(scheme_section.table, "Scheme Alpha", total_score_label)
    provenance_cell = provenance_row[2]
    assert "3.0.0" in provenance_cell.display_value
    assert "scorehash99" in provenance_cell.display_value
    plain_row = _scheme_metric_row(scheme_section.table, "Scheme Beta", total_score_label)
    plain_cell = plain_row[2]
    assert "72.0" in plain_cell.display_value
    assert "Version:" not in plain_cell.display_value
    assert "Content Hash:" not in plain_cell.display_value
    assert "版本" not in plain_cell.display_value
    assert "内容哈希" not in plain_cell.display_value

    docx_text = _docx_text(DocxRenderer().render(model))
    pdf_text = _pdf_text(PdfRenderer().render(model))
    assert "2.1.0" in docx_text
    assert "2.1.0" in pdf_text
    assert "3.0.0" in docx_text
    assert "3.0.0" in pdf_text


def test_p2_missing_information_not_duplicated_when_risks_cover_it() -> None:
    """Do not emit missing-information fallback when risks already cover the content."""
    content = {
        "report_metadata": {
            "project_id": "p2-project",
            "schema_version": "cold_storage_concept_design@1.0.0",
        },
        "risks_and_missing_information": {
            "risks": [
                {
                    "description": "Site geotechnical survey not completed",
                    "severity": "high",
                    "mitigation": "",
                }
            ],
            "missing_information": [
                {
                    "description": "Site geotechnical survey not completed",
                    "impact": "foundation_design",
                }
            ],
        },
    }
    canonical = build_canonical_render_model(
        content=content,
        report_id="p2-dedup",
        revision_number=1,
        content_hash="b" * 64,
        generated_by="p2-test",
        generated_at="2025-01-01T00:00:00Z",
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
    )
    for locale in (ReportLocale.ZH_CN, ReportLocale.EN_US):
        model = localize_render_model(canonical, locale=locale)
        risks_section = next(
            s for s in model.sections if s.section_key == "risks_and_missing_information"
        )
        assert "Site geotechnical survey not completed" in risks_section.text
        assert "Missing Information:" not in risks_section.text
        assert "缺失信息" not in risks_section.text

        docx_text = _docx_text(DocxRenderer().render(model))
        pdf_text = _pdf_text(PdfRenderer().render(model))
        assert "Site geotechnical survey not completed" in docx_text
        assert "Site geotechnical survey not completed" in pdf_text
        assert "Missing Information:" not in docx_text
        assert "Missing Information:" not in pdf_text


@pytest.mark.parametrize("locale", [ReportLocale.ZH_CN, ReportLocale.EN_US])
def test_p2_risks_render_in_docx_and_pdf(locale: ReportLocale) -> None:
    """Localized risks render in both DOCX and PDF when not covered by missing-info dedup."""
    canonical = _build_canonical()
    model = localize_render_model(canonical, locale=locale)
    risks_section = next(
        s for s in model.sections if s.section_key == "risks_and_missing_information"
    )
    assert "Risk about missing survey" in risks_section.text

    docx_text = _docx_text(DocxRenderer().render(model))
    pdf_text = _pdf_text(PdfRenderer().render(model))
    assert "Risk about missing survey" in docx_text
    assert "Risk about missing survey" in pdf_text
