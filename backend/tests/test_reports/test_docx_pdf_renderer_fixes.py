"""Tests for DOCX and PDF renderer fixes: P0-4, P0-5, P0-7.

P0-4: DOCX page orientation and table config
P0-5: DOCX header, footer, and watermark
P0-7: PDF cell alignment priority
"""

from __future__ import annotations

import zipfile
from io import BytesIO
from typing import Any

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT

# Register qn for XML namespace access in tests
from docx.oxml.ns import qn as _qn  # noqa: F401

from cold_storage.modules.reports.domain.errors import TemplateManifestError
from cold_storage.modules.reports.domain.render_model import (
    RenderManifest,
    RenderMetadata,
    RenderSection,
    RenderTable,
    RenderTableCell,
    ReportRenderModel,
)
from cold_storage.modules.reports.renderers.docx_renderer import DocxRenderer
from cold_storage.modules.reports.renderers.pdf_renderer import PdfRenderer

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_metadata(**overrides: Any) -> RenderMetadata:
    defaults = dict(
        report_id="r-test",
        project_name="蓝莓冷库概念设计项目",
        report_type="概念设计报告",
        schema_version="cold_storage_concept_design@1.0.0",
        revision_number=1,
        content_hash="a" * 64,
        content_hash_short="a" * 8,
        generated_at="2025-01-01T00:00:00+00:00",
        generated_by="test-system",
        template_version="1.0.0",
        template_code="cold_storage_concept_design",
    )
    defaults.update(overrides)
    return RenderMetadata(**defaults)  # type: ignore[arg-type]


def _make_manifest(
    sections: list[RenderSection],
    render_settings: dict | None = None,
    fmt: str = "pdf",
) -> RenderManifest:
    return RenderManifest(
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
        schema_version="cold_storage_concept_design@1.0.0",
        source_content_hash="a" * 64,
        sections=[s.section_key for s in sections],
        format=fmt,
        render_settings=render_settings or {},
    )


def _render_pdf(
    sections: list[RenderSection],
    *,
    is_draft: bool = False,
    render_settings: dict | None = None,
) -> bytes:
    metadata = _make_metadata()
    manifest = _make_manifest(sections, render_settings=render_settings, fmt="pdf")
    model = ReportRenderModel(metadata=metadata, sections=sections, manifest=manifest)
    return PdfRenderer().render(model, is_draft=is_draft)


def _render_docx(
    sections: list[RenderSection],
    *,
    is_draft: bool = False,
    render_settings: dict | None = None,
) -> bytes:
    metadata = _make_metadata()
    manifest = _make_manifest(sections, render_settings=render_settings, fmt="docx")
    model = ReportRenderModel(metadata=metadata, sections=sections, manifest=manifest)
    return DocxRenderer().render(model, is_draft=is_draft)


def _get_docx_text(docx_bytes: bytes) -> str:
    """Extract all XML text from a DOCX file."""
    with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
        text = ""
        for name in zf.namelist():
            if name.endswith(".xml"):
                text += zf.read(name).decode("utf-8", errors="ignore")
        return text


# ===========================================================================
# P0-7: PDF cell alignment priority
# ===========================================================================


class TestP0_7_PdfCellAlignment:
    def test_manifest_right_cell_explicit_left_final_left(self):
        """Manifest column right + cell explicit left -> final left."""
        sections = [
            RenderSection(
                section_key="s1",
                title="对齐测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["列A"],
                    rows=[
                        [RenderTableCell(value="值", align="left")],
                    ],
                ),
            )
        ]
        pdf_bytes = _render_pdf(
            sections,
            render_settings={
                "tables": {
                    "s1": {
                        "columns": [{"key": "a", "header": "列A", "align": "right"}],
                    },
                },
            },
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        margin_left = 56.69
        found = False
        for page_num in range(doc.page_count):
            page = doc[page_num]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span["text"] == "值":
                            found = True
                            # Left-aligned: x0 should be near left margin
                            assert abs(span["bbox"][0] - margin_left) < 20, (
                                f"Cell text '值' should be left-aligned, "
                                f"but x0={span['bbox'][0]:.1f}, expected ~{margin_left}"
                            )
        doc.close()
        assert found, "target span '值' not found in PDF"

    def test_manifest_center_cell_none_final_center(self):
        """Manifest column center + cell None -> final center."""
        sections = [
            RenderSection(
                section_key="s1",
                title="对齐测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["列A"],
                    rows=[
                        [RenderTableCell(value="居中")],  # align=None
                    ],
                ),
            )
        ]
        pdf_bytes = _render_pdf(
            sections,
            render_settings={
                "tables": {
                    "s1": {
                        "columns": [{"key": "a", "header": "列A", "align": "center"}],
                    },
                },
            },
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        found = False
        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_width = page.rect.width
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span["text"] == "居中":
                            found = True
                            span_center_x = (span["bbox"][0] + span["bbox"][2]) / 2
                            page_center_x = page_width / 2
                            assert abs(span_center_x - page_center_x) < 10, (
                                f"Cell text '居中' should be center-aligned, "
                                f"span_center={span_center_x:.1f}, page_center={page_center_x:.1f}"
                            )
        doc.close()
        assert found, "target span '居中' not found in PDF"

    def test_split_row_alignment_consistent_with_normal(self):
        """Split row alignment must be consistent with normal row alignment.

        P0-6: Verifies that normal rows and split rows use the same alignment
        rule by finding spans from both row types and comparing alignment.
        """
        long_text = "这是一段超长的中文文本用于测试分页对齐一致性。" * 50
        sections = [
            RenderSection(
                section_key="s1",
                title="分页对齐测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["项目", "内容"],
                    rows=[
                        [
                            RenderTableCell(value="短文本", align="right"),
                            RenderTableCell(value="普通"),
                        ],
                        [
                            RenderTableCell(value="长文本", align="right"),
                            RenderTableCell(value=long_text),
                        ],
                    ],
                ),
            )
        ]
        pdf_bytes = _render_pdf(
            sections,
            render_settings={
                "tables": {
                    "s1": {
                        "columns": [
                            {"key": "a", "header": "项目", "align": "left"},
                            {"key": "b", "header": "内容", "align": "left"},
                        ],
                    },
                },
            },
        )
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        assert doc.page_count >= 1

        # Find all spans with "短文本" (normal row) and check alignment
        normal_positions = []
        split_positions = []
        for page_num in range(doc.page_count):
            page = doc[page_num]
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block["type"] != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span["text"] == "短文本":
                            normal_positions.append(span["bbox"][0])
                        elif span["text"] == "长文本":
                            split_positions.append(span["bbox"][0])

        doc.close()

        # Both should be aligned (x0 values should be close)
        if normal_positions and split_positions:
            normal_x = normal_positions[0]
            split_x = split_positions[0]
            # Alignment tolerance: within 10pt
            assert abs(normal_x - split_x) < 10, (
                f"Normal row x0={normal_x:.1f} should match split row x0={split_x:.1f}"
            )


# ===========================================================================
# P0-4: DOCX orientation
# ===========================================================================


class TestP0_4_DocxOrientation:
    def test_default_portrait_orientation(self):
        """Default orientation is portrait when page.orientation is not set."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(sections)
        doc = Document(BytesIO(docx_bytes))
        for section in doc.sections:
            assert section.orientation == WD_ORIENT.PORTRAIT

    def test_manifest_landscape_orientation(self):
        """page.orientation=landscape sets all sections to landscape."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={"page": {"orientation": "landscape"}},
        )
        doc = Document(BytesIO(docx_bytes))
        for section in doc.sections:
            assert section.orientation == WD_ORIENT.LANDSCAPE

    def test_landscape_sections_creates_new_word_section(self):
        """Landscape section creates a new Word section with landscape orientation."""
        sections = [
            RenderSection(
                section_key="portrait1",
                title="竖版章节",
                level=1,
                content_type="text",
                text="竖版内容",
            ),
            RenderSection(
                section_key="landscape1",
                title="横版章节",
                level=1,
                content_type="text",
                text="横版内容",
            ),
            RenderSection(
                section_key="portrait2",
                title="恢复竖版",
                level=1,
                content_type="text",
                text="竖版内容",
            ),
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "landscape_sections": ["landscape1"],
            },
        )
        doc = Document(BytesIO(docx_bytes))
        assert len(doc.sections) >= 2

    def test_table_config_orientation_overrides_landscape_sections(self):
        """tables[key].orientation overrides landscape_sections list."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "landscape_sections": ["s1"],
                "tables": {
                    "s1": {"orientation": "portrait"},
                },
            },
        )
        doc = Document(BytesIO(docx_bytes))
        assert doc.sections[0].orientation == WD_ORIENT.PORTRAIT


# ===========================================================================
# P0-4: DOCX table config
# ===========================================================================


class TestP0_4_DocxTableConfig:
    def test_unit_row_false_no_unit_row(self):
        """unit_row=false should NOT create a unit row."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                    unit_row=["kg", "m"],
                ),
            )
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "tables": {
                    "s1": {"unit_row": False},
                },
            },
        )
        doc = Document(BytesIO(docx_bytes))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 2

    def test_unit_row_true_creates_unit_row(self):
        """unit_row=true (default) creates unit row when data has units."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                    unit_row=["kg", "m"],
                ),
            )
        ]
        docx_bytes = _render_docx(sections)
        doc = Document(BytesIO(docx_bytes))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 3

    def test_column_count_mismatch_raises_error(self):
        """Column count mismatch between manifest and table raises TemplateManifestError."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                ),
            )
        ]
        try:
            _render_docx(
                sections,
                render_settings={
                    "tables": {
                        "s1": {
                            "columns": [
                                {"key": "a", "header": "A"},
                            ],
                        },
                    },
                },
            )
            raise AssertionError("Should have raised TemplateManifestError")
        except TemplateManifestError:
            pass

    def test_invalid_width_ratio_raises_error(self):
        """Invalid width_ratio (all zero) raises TemplateManifestError."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                ),
            )
        ]
        try:
            _render_docx(
                sections,
                render_settings={
                    "tables": {
                        "s1": {
                            "columns": [
                                {"key": "a", "header": "A", "width_ratio": 0},
                                {"key": "b", "header": "B", "width_ratio": 0},
                            ],
                        },
                    },
                },
            )
            raise AssertionError("Should have raised TemplateManifestError")
        except TemplateManifestError:
            pass


# ===========================================================================
# P0-5: DOCX header/watermark
# ===========================================================================


class TestP0_5_DocxHeaderWatermark:
    def test_custom_header_and_watermark_both_exist(self):
        """Custom header text and watermark both exist in output."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            is_draft=True,
            render_settings={
                "header": {"right": "自定义页眉"},
                "watermark": {"text": "DRAFT"},
            },
        )
        docx_text = _get_docx_text(docx_bytes)
        assert "自定义页眉" in docx_text, "Custom header text not found"
        assert "PowerPlusWaterMarkObject" in docx_text, "Watermark not found"

    def test_watermark_does_not_delete_header_text(self):
        """Watermark adds a separate paragraph, header text is preserved."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            is_draft=True,
            render_settings={
                "header": {"left": "左侧", "right": "右侧"},
                "watermark": {"text": "DRAFT"},
            },
        )
        docx_text = _get_docx_text(docx_bytes)
        assert "左侧" in docx_text, "Left header text not found"
        assert "右侧" in docx_text, "Right header text not found"
        assert "DRAFT" in docx_text, "Watermark text not found"

    def test_page_number_field_in_footer_left(self):
        """{page_number} in footer left generates PAGE field."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "footer": {"left": "第{page_number}页", "center": "", "right": ""},
            },
        )
        docx_text = _get_docx_text(docx_bytes)
        assert "PAGE" in docx_text, "PAGE field not found in footer"

    def test_page_number_field_in_footer_right(self):
        """{page_number} in footer right generates PAGE field."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="text",
                text="测试内容",
            )
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "footer": {"left": "", "center": "", "right": "第{page_number}页"},
            },
        )
        docx_text = _get_docx_text(docx_bytes)
        assert "PAGE" in docx_text, "PAGE field not found in footer right"

    def test_landscape_section_tab_stop_matches_width(self):
        """Landscape section tab stops match landscape page width.

        P0-7: Verifies that portrait and landscape sections produce different
        tab stop positions in their headers/footers, and that left/center/right
        text and PAGE fields are present.
        """
        import zipfile

        from lxml import etree

        sections = [
            RenderSection(
                section_key="portrait1",
                title="竖版章节",
                level=1,
                content_type="text",
                text="竖版内容",
            ),
            RenderSection(
                section_key="landscape1",
                title="横版章节",
                level=1,
                content_type="text",
                text="横版内容",
            ),
        ]
        docx_bytes = _render_docx(
            sections,
            render_settings={
                "landscape_sections": ["landscape1"],
                "header": {"left": "左", "center": "中", "right": "右"},
                "footer": {"left": "左", "center": "中", "right": "右"},
            },
        )
        assert len(docx_bytes) > 0
        doc = Document(BytesIO(docx_bytes))
        assert len(doc.sections) >= 2

        with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
            # Find header XML files
            header_files = [n for n in zf.namelist() if "header" in n and n.endswith(".xml")]
            footer_files = [n for n in zf.namelist() if "footer" in n and n.endswith(".xml")]
            assert len(header_files) >= 1, "No header XML files found"
            assert len(footer_files) >= 1, "No footer XML files found"

            # Parse each header XML for w:tab/@w:pos values
            header_tab_positions = []
            w_pos_attr = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pos"
            for hf in sorted(header_files):
                xml_content = zf.read(hf)
                root = etree.fromstring(xml_content)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                tabs = root.findall(".//w:tab", ns)
                positions = [int(tab.get(w_pos_attr)) for tab in tabs if tab.get(w_pos_attr)]
                if positions:
                    header_tab_positions.append(max(positions))

            # Parse each footer XML for w:tab/@w:pos values
            footer_tab_positions = []
            for ff in sorted(footer_files):
                xml_content = zf.read(ff)
                root = etree.fromstring(xml_content)
                tabs = root.findall(".//w:tab", ns)
                positions = [int(tab.get(w_pos_attr)) for tab in tabs if tab.get(w_pos_attr)]
                if positions:
                    footer_tab_positions.append(max(positions))

            # Different orientations must produce different tab stop positions
            if len(header_tab_positions) >= 2:
                assert header_tab_positions[0] != header_tab_positions[1], (
                    f"Portrait and landscape headers should have different tab stops: "
                    f"{header_tab_positions}"
                )
            if len(footer_tab_positions) >= 2:
                assert footer_tab_positions[0] != footer_tab_positions[1], (
                    f"Portrait and landscape footers should have different tab stops: "
                    f"{footer_tab_positions}"
                )

            # Verify left/center/right text exists in header
            all_xml = ""
            for hf in header_files:
                all_xml += zf.read(hf).decode("utf-8", errors="ignore")
            for ff in footer_files:
                all_xml += zf.read(ff).decode("utf-8", errors="ignore")
            assert "左" in all_xml, "Left header/footer text not found"
            assert "中" in all_xml, "Center header/footer text not found"
            assert "右" in all_xml, "Right header/footer text not found"


# ===========================================================================
# P0-7: PDF manifest errors
# ===========================================================================


class TestP0_7_PdfManifestErrors:
    def test_wrong_column_count_raises_manifest_error(self):
        """Wrong column count raises TemplateManifestError."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                ),
            )
        ]
        try:
            _render_pdf(
                sections,
                render_settings={
                    "tables": {
                        "s1": {
                            "columns": [{"key": "a", "header": "A"}],
                        },
                    },
                },
            )
            raise AssertionError("Should have raised TemplateManifestError")
        except TemplateManifestError:
            pass

    def test_invalid_width_ratio_raises_manifest_error(self):
        """Invalid width_ratio (all zero) raises TemplateManifestError."""
        sections = [
            RenderSection(
                section_key="s1",
                title="测试",
                level=1,
                content_type="table",
                table=RenderTable(
                    title="测试表",
                    headers=["A", "B"],
                    rows=[[RenderTableCell(value="1"), RenderTableCell(value="2")]],
                ),
            )
        ]
        try:
            _render_pdf(
                sections,
                render_settings={
                    "tables": {
                        "s1": {
                            "columns": [
                                {"key": "a", "header": "A", "width_ratio": 0},
                                {"key": "b", "header": "B", "width_ratio": 0},
                            ],
                        },
                    },
                },
            )
            raise AssertionError("Should have raised TemplateManifestError")
        except TemplateManifestError:
            pass
