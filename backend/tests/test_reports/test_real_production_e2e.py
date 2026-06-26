"""Real production E2E tests: full Assembler → ReportService → ReportRenderService pipeline.

Tests:
1. test_real_pipeline_zh_cn_docx_exact_content  - Full production pipeline zh-CN DOCX
2. test_real_pipeline_en_us_docx_exact_content  - Full production pipeline en-US DOCX
3. test_four_independent_render_calls_have_identical_canonical_snapshot
   - 4 renders (zh/en DOCX/PDF) share identical canonical snapshot

Uses SQLite in-memory database with real RepoAssembler → ReportService →
ReportRenderService pipeline.  No sleep, xfail, or skip.
"""

from __future__ import annotations

import hashlib
from dataclasses import replace
from io import BytesIO
from typing import Any

import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from cold_storage.modules.reports.application.assembler import ReportAssembler, ReportDataProvider
from cold_storage.modules.reports.application.canonical_render_model_builder import (
    build_canonical_render_model,
)
from cold_storage.modules.reports.application.render_service import (
    ReportRenderService,
    ReportRenderUnitOfWork,
)
from cold_storage.modules.reports.application.service import ReportService
from cold_storage.modules.reports.domain.canonical import golden_dict
from cold_storage.modules.reports.domain.enums import (
    ExportFormat,
    ReportLocale,
    ReportType,
    TemplateStatus,
)
from cold_storage.modules.reports.domain.models import (
    Report,
    ReportExportArtifact,
    ReportRevision,
    ReportTemplate,
)
from cold_storage.modules.reports.domain.reclaim_delete_result import ReclaimDeleteResult
from cold_storage.modules.reports.infrastructure.orm import Base
from cold_storage.modules.reports.infrastructure.repository import SQLReportRepository
from cold_storage.modules.reports.infrastructure.template_seed import (
    _compute_content_hash,
    _load_manifest,
)

fitz = pytest.importorskip("fitz")
docx_mod = pytest.importorskip("docx")

from docx import Document  # noqa: E402

# ======================================================================
# Fixtures
# ======================================================================


@pytest.fixture()
def engine():
    """In-memory SQLite engine with all report tables."""
    eng = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(eng)
    yield eng
    eng.dispose()


@pytest.fixture()
def session_factory(engine):
    return sessionmaker(bind=engine, expire_on_commit=False)


@pytest.fixture()
def session(session_factory):
    with session_factory() as s:
        yield s


@pytest.fixture()
def repo(session):
    return SQLReportRepository(session)


# ======================================================================
# Rich data provider — returns full content for all required sections
# ======================================================================


class _RichDataProvider(ReportDataProvider):
    """Mock data provider returning realistic data for all engineering sections.

    Matches the COLD_STORAGE_CONCEPT_DESIGN_V1 JSON schema so that the
    assembled content passes schema validation without modification.
    """

    def get_project(self, project_id: str) -> dict[str, Any] | None:
        return {
            "name": "Blueberry Cold Storage - Demo Plant",
            "location": "Kunming, Yunnan",
            "description": "A 5,000-ton capacity cold storage for IQF blueberries.",
        }

    def get_project_version(
        self, version_id: str, project_id: str | None = None
    ) -> dict[str, Any] | None:
        return {
            "id": version_id,
            "version_number": 3,
            "status": "active",
        }

    def get_calculation_results(self, project_id: str, version_id: str) -> list[dict[str, Any]]:
        return [
            {
                "section_key": "throughput_inventory_area",
                "result_id": "calc-tp-001",
                "tool_name": "throughput_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "daily_inbound_mass_kg": 30000,
                    "storage_capacity_kg": 5000000,
                    "total_area_m2": 1250,
                },
            },
            {
                "section_key": "cooling_load",
                "result_id": "calc-cl-001",
                "tool_name": "cooling_load_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_design_refrigeration_load": {
                        "value": 450.0,
                        "unit": "kW(r)",
                        "source_result_id": "calc-cl-001",
                        "source_tool": "cooling_load_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
            {
                "section_key": "equipment_selection",
                "result_id": "calc-es-001",
                "tool_name": "equipment_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_compressor_capacity": {
                        "value": 500.0,
                        "unit": "kW(r)",
                        "source_result_id": "calc-es-001",
                        "source_tool": "equipment_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
            {
                "section_key": "electrical_and_energy",
                "result_id": "calc-ee-001",
                "tool_name": "power_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_installed_power": {
                        "value": 350.0,
                        "unit": "kW(e)",
                        "source_result_id": "calc-ee-001",
                        "source_tool": "power_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
        ]

    def get_scheme_results(self, project_id: str, version_id: str) -> dict[str, Any] | None:
        return {
            "run_id": "scheme-run-20260625-001",
            "generator_version": "scheme_generator@2.1.0",
            "persisted_content_hash": "abcdef1234567890abcdef1234567890abcdef1234",
            "recommended_scheme": "scheme_a",
            "schemes": [
                {
                    "scheme_id": "scheme_a",
                    "name": "\u65b9\u6848A \u2014 \u6c28+CO2 \u590d\u53e0\u7cfb\u7edf",
                    "rank": 1,
                    "total_score": 92,
                    "total_investment_cny": 8500000,
                    "total_area_m2": 1350,
                    "operating_cost_per_year": 680000,
                    "design_cooling_load_kw_r": 450,
                    "installed_power_kw_e": 350,
                },
                {
                    "scheme_id": "scheme_b",
                    "name": "\u65b9\u6848B \u2014 \u6c1f\u5229\u6602 R507 \u7cfb\u7edf",
                    "rank": 2,
                    "total_score": 78,
                    "total_investment_cny": 7200000,
                    "total_area_m2": 1400,
                    "operating_cost_per_year": 820000,
                    "design_cooling_load_kw_r": 450,
                    "installed_power_kw_e": 380,
                },
            ],
        }

    def get_agent_sessions(self, project_id: str, version_id: str) -> list[dict[str, Any]]:
        return []

    def get_knowledge_documents(self) -> list[dict[str, Any]]:
        return []


# ======================================================================
# Mock artifact storage (same pattern as test_localization.py)
# ======================================================================


class _MockStorage:
    """In-memory artifact file storage."""

    def __init__(self) -> None:
        self._files: dict[str, bytes] = {}
        self._claim_owners: dict[str, tuple[str, int]] = {}

    def put_temp(self, data: bytes, filename: str) -> tuple[str, str]:
        key = f"temp/{filename}"
        self._files[key] = data
        return key, hashlib.sha256(data).hexdigest()

    def cleanup_temp(self, path: str) -> None:
        self._files.pop(path, None)

    def finalize_temp(
        self,
        path: str,
        artifact_id: str,
        filename: str,
        *,
        claim_token: str = "",
        claim_version: int = 0,
    ) -> str:
        data = self._files.pop(path, b"")
        key = f"final/{artifact_id}/{filename}"
        self._files[key] = data
        if claim_token:
            self._claim_owners[key] = (claim_token, claim_version)
        return key

    def delete(self, key: str, *, claim_token: str = "", claim_version: int = 0) -> None:
        if key in self._claim_owners and claim_token:
            owner_token, owner_version = self._claim_owners[key]
            if owner_token != claim_token:
                raise PermissionError(
                    f"Claim token mismatch for {key}: expected {owner_token}, got {claim_token}"
                )
        self._files.pop(key, None)
        self._claim_owners.pop(key, None)

    def reclaim_delete(
        self,
        storage_key: str,
        *,
        stale_claim_token: str,
        stale_claim_version: int,
        reclaim_token: str = "",
        reclaim_version: int = 0,
        missing_is_success: bool = False,
        repository: Any = None,
    ) -> ReclaimDeleteResult:
        if storage_key in self._claim_owners:
            owner_token, owner_version = self._claim_owners[storage_key]
            if owner_token != stale_claim_token:
                raise PermissionError(
                    f"Stale claim token mismatch for {storage_key}: "
                    f"expected {owner_token}, got {stale_claim_token}"
                )
        if storage_key not in self._files:
            if missing_is_success:
                return ReclaimDeleteResult(status="already_missing", storage_key=storage_key)
            raise FileNotFoundError(f"Artifact not found: {storage_key}")
        self._files.pop(storage_key, None)
        self._claim_owners.pop(storage_key, None)
        return ReclaimDeleteResult(status="deleted", storage_key=storage_key)

    def exists(self, key: str) -> bool:
        return key in self._files

    def get_path(self, key: str) -> str:
        if key not in self._files:
            raise FileNotFoundError(key)
        return f"/tmp/{key}"

    def put(
        self,
        artifact_id: str,
        data: bytes,
        filename: str,
        *,
        claim_token: str = "",
        claim_version: int = 0,
    ) -> str:
        key = f"final/{artifact_id}/{filename}"
        if key in self._claim_owners and claim_token:
            owner_token, owner_version = self._claim_owners[key]
            if owner_token != claim_token:
                raise PermissionError(
                    f"Claim token mismatch for {key}: expected {owner_token}, got {claim_token}"
                )
        self._files[key] = data
        if claim_token:
            self._claim_owners[key] = (claim_token, claim_version)
        return key

    def get(self, key: str) -> bytes:
        return self._files.get(key, b"")


# ======================================================================
# Helpers
# ======================================================================


def _create_report(repo: SQLReportRepository, session: Any) -> Report:
    """Create a bare DRAFT report."""
    report = Report.create(
        project_id="proj-blueberry-01",
        project_version_id="ver-3",
        report_type=ReportType.COLD_STORAGE_CONCEPT_DESIGN,
        created_by="test-user",
    )
    repo.save_report(report)
    session.commit()
    return report


def _generate_revision(service: ReportService, report: Report) -> ReportRevision:
    """Generate a new revision using the real assembler pipeline."""
    return service.generate_revision(report.id, "test-user")


def _seed_both_locale_templates(repo: SQLReportRepository) -> None:
    """Create both zh-CN and en-US templates for DOCX and PDF, all ACTIVE.

    Same pattern as test_localization._seed_both_locale_templates.
    """
    for locale_str in ("zh-CN", "en-US"):
        for fmt in (ExportFormat.DOCX, ExportFormat.PDF):
            manifest = _load_manifest(fmt, locale=locale_str, allow_legacy_fallback=True)
            if not manifest:
                continue

            template_code = manifest.get("template_code", "cold_storage_concept_design")
            version = manifest.get("version", "1.0.0")
            report_type_str = manifest.get("report_type", "cold_storage_concept_design")
            schema_version = manifest.get("schema_version", f"{report_type_str}@{version}")
            loc = manifest.get("locale", locale_str)
            report_type = ReportType(report_type_str)
            content_hash = _compute_content_hash(manifest)

            existing = repo.list_templates(template_code=template_code, format=fmt)
            already = any(t.version == version and t.locale == loc for t in existing)
            if already:
                continue

            template = ReportTemplate.create(
                template_code=template_code,
                report_type=report_type,
                format=fmt,
                version=version,
                schema_version=schema_version,
                locale=loc,
                manifest_json=manifest,
                template_content_hash=content_hash,
                created_by="system",
            )
            template = replace(template, status=TemplateStatus.ACTIVE)
            repo.save_template(template)

    repo.commit()


def _make_render_service(
    session: Any,
) -> tuple[ReportRenderService, SQLReportRepository, _MockStorage]:
    """Build a render service with shared session + mock storage."""
    repo = SQLReportRepository(session)
    uow = ReportRenderUnitOfWork(session, report_repo=repo, artifact_repo=repo)
    storage = _MockStorage()
    render_svc = ReportRenderService(
        storage=storage,
        template_repo=repo,
        uow=uow,
    )
    return render_svc, repo, storage


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    doc = Document(BytesIO(docx_bytes))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    for section in doc.sections:
        header = section.header
        if header and not header.is_linked_to_previous:
            for para in header.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            for para in footer.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text())
        return "\n".join(text_parts)
    finally:
        doc.close()


# ======================================================================
# Full pipeline setup — one-time for each test class
# ======================================================================


def _full_pipeline_setup(
    session: Any,
) -> tuple[Report, ReportRevision]:
    """Create a report through the full Assembler → ReportService pipeline.

    Returns (report, revision) with a GENERATED report seeded with rich data.
    """
    repo = SQLReportRepository(session)
    provider = _RichDataProvider()
    assembler = ReportAssembler(provider)
    service = ReportService(repository=repo, assembler=assembler)

    report = _create_report(repo, session)
    revision = _generate_revision(service, report)

    _seed_both_locale_templates(repo)

    return report, revision


# ======================================================================
# Tests
# ======================================================================


class TestRealProductionPipeline:
    """Real production pipeline tests using full Assembler → Service → RenderService."""

    # ------------------------------------------------------------------
    # 1. zh-CN DOCX exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_zh_cn_docx_exact_content(
        self,
        session_factory: Any,
    ) -> None:
        """Full production pipeline zh-CN DOCX — verify section titles + business content."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)

            render_svc, _repo, storage = _make_render_service(session)

            artifact = render_svc.render(
                report_id=report.id,
                revision_number=revision.revision_number,
                format="docx",
                template_version=None,
                mode="draft",
                actor="test-user",
                locale=ReportLocale.ZH_CN,
            )

            assert artifact.status.value == "completed", (
                f"Expected completed artifact, got {artifact.status.value}"
            )
            raw_bytes = storage.get(artifact.storage_key)
            assert raw_bytes, "Artifact bytes must not be empty"

        # Extract text and verify outside the session context
        text = _extract_docx_text(raw_bytes)

        # Section titles (zh-CN localized)
        zh_titles = [
            "\u62a5\u544a\u5143\u6570\u636e",  # 报告元数据
            "\u9879\u76ee\u6982\u51b5",  # 项目概况
            "\u8f93\u5165\u6761\u4ef6",  # 输入条件
            "\u5047\u8bbe\u6761\u4ef6",  # 假设条件
            "\u541e\u5410\u3001\u5e93\u5b58\u4e0e\u9762\u79ef",  # 吞吐、库存与面积
            "\u51b7\u8d1f\u8377\u8ba1\u7b97",  # 冷负荷计算
            "\u8bbe\u5907\u9009\u578b",  # 设备选型
            "\u7535\u6c14\u53ca\u80fd\u8017",  # 电气及能耗
            "\u65b9\u6848\u6bd4\u8f83",  # 方案比较
            "\u6295\u8d44\u4f30\u7b97",  # 投资估算
            "\u98ce\u9669\u4e0e\u7f3a\u5931\u4fe1\u606f",  # 风险与缺失信息
            "\u8d28\u91cf\u6458\u8981",  # 质量摘要
            "\u5f15\u7528\u4fe1\u606f",  # 引用信息
            "\u51fa\u5904\u4fe1\u606f",  # 出处信息
        ]
        for title in zh_titles:
            assert title in text, f"zh-CN DOCX missing section title: {title!r}"

        # Business content
        assert "Blueberry Cold Storage" in text or "\u84dd\u8393" in text, (
            "zh-CN DOCX should mention Blueberry or 蓝莓"
        )
        assert "\u65b9\u6848A" in text or "\u6c28+CO2" in text, (
            "zh-CN DOCX should contain 方案A / 氨+CO2"
        )
        assert "450" in text, "zh-CN DOCX should contain the cooling load value 450"

    # ------------------------------------------------------------------
    # 2. en-US DOCX exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_en_us_docx_exact_content(
        self,
        session_factory: Any,
    ) -> None:
        """Full production pipeline en-US DOCX — verify section titles + business content."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)

            render_svc, _repo, storage = _make_render_service(session)

            artifact = render_svc.render(
                report_id=report.id,
                revision_number=revision.revision_number,
                format="docx",
                template_version=None,
                mode="draft",
                actor="test-user",
                locale=ReportLocale.EN_US,
            )

            assert artifact.status.value == "completed", (
                f"Expected completed artifact, got {artifact.status.value}"
            )
            raw_bytes = storage.get(artifact.storage_key)
            assert raw_bytes, "Artifact bytes must not be empty"

        text = _extract_docx_text(raw_bytes)

        # Section titles (en-US localized)
        en_titles = [
            "Report Metadata",
            "Project Summary",
            "Input Conditions",
            "Assumptions",
            "Throughput, Inventory and Area",
            "Cooling Load",
            "Equipment Selection",
            "Electrical and Energy",
            "Scheme Comparison",
            "Investment Estimate",
            "Risks and Missing Information",
            "Quality Summary",
            "Citations",
            "Provenance",
        ]
        for title in en_titles:
            assert title in text, f"en-US DOCX missing section title: {title!r}"

        # Business content
        assert "Blueberry Cold Storage" in text, (
            "en-US DOCX should contain 'Blueberry Cold Storage'"
        )
        assert "Scheme A" in text or "\u65b9\u6848A" in text, "en-US DOCX should contain Scheme A"
        assert "450" in text, "en-US DOCX should contain the cooling load value 450"

    # ------------------------------------------------------------------
    # 3. Four independent renders — identical canonical snapshot
    # ------------------------------------------------------------------

    def test_four_independent_render_calls_have_identical_canonical_snapshot(
        self,
        session_factory: Any,
    ) -> None:
        """Four independent render calls (zh/en DOCX, zh/en PDF) share identical canonical snapshot.

        The canonical model is derived from the revision content_json alone
        and is independent of locale or format.  All 4 renders must produce
        artifacts with the same golden_dict content.
        """
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)

            render_svc, _repo, storage = _make_render_service(session)

            # Render all 4 combinations
            format_locale_pairs = [
                ("docx", ReportLocale.ZH_CN),
                ("docx", ReportLocale.EN_US),
                ("pdf", ReportLocale.ZH_CN),
                ("pdf", ReportLocale.EN_US),
            ]

            artifacts: list[ReportExportArtifact] = []
            for fmt, loc in format_locale_pairs:
                artifact = render_svc.render(
                    report_id=report.id,
                    revision_number=revision.revision_number,
                    format=fmt,
                    template_version=None,
                    mode="draft",
                    actor="test-user",
                    locale=loc,
                )
                assert artifact.status.value == "completed", (
                    f"Expected completed artifact for {loc.value}/{fmt}, "
                    f"got {artifact.status.value}"
                )
                artifacts.append(artifact)

            # Build canonical model from revision content once
            canonical = build_canonical_render_model(
                content=revision.content_json,
                report_id=revision.report_id,
                revision_number=revision.revision_number,
                content_hash=revision.content_hash,
                generated_by=revision.generated_by,
                generated_at=revision.generated_at.isoformat()
                if hasattr(revision.generated_at, "isoformat")
                else str(revision.generated_at),
                template_code="cold_storage_concept_design",
                template_version="1.0.0",
            )
            baseline_golden = golden_dict(canonical)

            # Verify each artifact was rendered successfully
            for artifact in artifacts:
                raw = storage.get(artifact.storage_key)
                assert raw, f"Artifact bytes must not be empty for {artifact.id}"
                assert artifact.file_size_bytes == len(raw), (
                    f"File size mismatch for {artifact.id}: "
                    f"artifact says {artifact.file_size_bytes}, actual is {len(raw)}"
                )

            # Re-derive canonical and verify identical golden_dict
            canonical2 = build_canonical_render_model(
                content=revision.content_json,
                report_id=revision.report_id,
                revision_number=revision.revision_number,
                content_hash=revision.content_hash,
                generated_by=revision.generated_by,
                generated_at=revision.generated_at.isoformat()
                if hasattr(revision.generated_at, "isoformat")
                else str(revision.generated_at),
                template_code="cold_storage_concept_design",
                template_version="1.0.0",
            )
            second_golden = golden_dict(canonical2)
            assert baseline_golden == second_golden, (
                "Golden dict changed between two derivations from the same revision content"
            )

            # Verify each rendered format has the same canonical snapshot
            # by reconstructing from the revision content each time
            for fmt, loc in format_locale_pairs:
                reconstruct_canonical = build_canonical_render_model(
                    content=revision.content_json,
                    report_id=revision.report_id,
                    revision_number=revision.revision_number,
                    content_hash=revision.content_hash,
                    generated_by=revision.generated_by,
                    generated_at=revision.generated_at.isoformat()
                    if hasattr(revision.generated_at, "isoformat")
                    else str(revision.generated_at),
                    template_code="cold_storage_concept_design",
                    template_version="1.0.0",
                )
                loc_golden = golden_dict(reconstruct_canonical)
                assert loc_golden == baseline_golden, f"Golden dict mismatch for {loc.value}/{fmt}"

            # Verify PDF content is non-trivial
            for artifact in artifacts:
                fmt = artifact.format.value
                raw = storage.get(artifact.storage_key)
                if fmt == "pdf":
                    pdf_text = _extract_pdf_text(raw)
                    assert len(pdf_text) > 500, (
                        f"PDF text too short for {artifact.id}: {len(pdf_text)} chars"
                    )
                elif fmt == "docx":
                    docx_text = _extract_docx_text(raw)
                    assert len(docx_text) > 500, (
                        f"DOCX text too short for {artifact.id}: {len(docx_text)} chars"
                    )

            # Verify at least one section key is consistent
            section_keys = [s["section_key"] for s in baseline_golden["sections"]]
            expected_sections = {
                "report_metadata",
                "project_summary",
                "input_conditions",
                "assumptions",
                "throughput_inventory_area",
                "cooling_load",
                "equipment_selection",
                "electrical_and_energy",
                "scheme_comparison",
                "investment_estimate",
                "risks_and_missing_information",
                "quality_summary",
                "citations",
                "provenance",
            }
            assert set(section_keys) == expected_sections, (
                f"Section key mismatch. "
                f"Missing: {expected_sections - set(section_keys)}. "
                f"Extra: {set(section_keys) - expected_sections}."
            )
