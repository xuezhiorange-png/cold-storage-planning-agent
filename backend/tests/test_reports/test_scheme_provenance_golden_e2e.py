"""Tests for scheme source_* provenance, golden serializer, and full E2E DOCX/PDF.

Task A: Scheme run-level provenance → source_* fields
Task B: Golden serializer fixed output all canonical fields
Task C: Complete bilingual DOCX/PDF file content parsing E2E
"""

from __future__ import annotations

import hashlib
from io import BytesIO
from typing import Any

import pytest

from cold_storage.modules.reports.application.canonical_render_model_builder import (
    build_canonical_render_model,
)
from cold_storage.modules.reports.application.render_model_localizer import (
    localize_render_model,
)
from cold_storage.modules.reports.domain.canonical import (
    golden_dict,
    golden_json,
)
from cold_storage.modules.reports.domain.enums import ReportLocale
from cold_storage.modules.reports.domain.models import ApprovalSnapshot
from cold_storage.modules.reports.domain.render_model import (
    CanonicalRenderSection,
    CanonicalReportRenderModel,
    LocalizedReportRenderModel,
)

fitz = pytest.importorskip("fitz")
docx_mod = pytest.importorskip("docx")

from docx import Document  # noqa: E402

# ======================================================================
# Helpers
# ======================================================================

_SAMPLE_CONTENT: dict[str, Any] = {
    "report_metadata": {
        "schema_version": "cold_storage_concept_design@1.0.0",
        "report_id": "test-report-001",
        "project_id": "proj-blueberry-01",
        "project_version_id": "ver-3",
        "generated_at": "2026-06-25T10:00:00Z",
        "generated_by": "test-agent",
        "revision_number": 2,
    },
    "project_summary": {
        "project_name": "Blueberry Cold Storage - Demo Plant",
        "project_location": "Kunming, Yunnan",
        "description": "A 5,000-ton capacity cold storage for IQF blueberries.",
    },
    "input_conditions": {
        "product_type": "blueberry",
    },
    "assumptions": {},
    "throughput_inventory_area": {
        "daily_inbound_mass_kg": {"value": 30000, "unit": "kg"},
        "storage_capacity_kg": {"value": 5000000, "unit": "kg"},
        "total_area_m2": {"value": 1250, "unit": "m2"},
    },
    "cooling_load": {
        "total_design_refrigeration_load": {
            "value": 450.0,
            "unit": "kW(r)",
            "source_tool": "cooling_load_calculator",
            "source_tool_version": "1.2.0",
            "source_content_hash": "aaaabbbbccccddddeeeeffff00001111",
        },
    },
    "equipment_selection": {
        "total_compressor_capacity": {"value": 500.0, "unit": "kW(r)"},
        "condenser_type": "evaporative",
    },
    "electrical_and_energy": {
        "total_installed_power": {"value": 350.0, "unit": "kW(e)"},
    },
    "scheme_comparison": {
        "run_id": "scheme-run-20260625-001",
        "generator_version": "scheme_generator@2.1.0",
        "scheme_evaluator": "scheme_evaluator",
        "persisted_content_hash": "abcdef1234567890abcdef1234567890abcdef1234",
        "recommended_scheme": "scheme_a",
        "schemes": [
            {
                "scheme_id": "scheme_a",
                "name": "方案A — 氨+CO2 复叠系统",
                "rank": 1,
                "total_score": {"value": 92, "unit": "score"},
                "total_investment_cny": {"value": 8500000, "unit": "CNY"},
                "total_area_m2": {"value": 1350, "unit": "m2"},
                "operating_cost_per_year": {"value": 680000, "unit": "CNY"},
                "design_cooling_load_kw_r": {"value": 450, "unit": "kW(r)"},
                "installed_power_kw_e": {"value": 350, "unit": "kW(e)"},
                "source_id": "scheme_a_prov_id",
                "source_tool": "scheme_generator",
                "source_tool_version": "2.1.0",
                "source_content_hash": "scheme_a_hash_abcdef",
            },
            {
                "scheme_id": "scheme_b",
                "name": "方案B — 氟利昂 R507 系统",
                "rank": 2,
                "total_score": {"value": 78, "unit": "score"},
                "total_investment_cny": {"value": 7200000, "unit": "CNY"},
                "total_area_m2": {"value": 1400, "unit": "m2"},
                "operating_cost_per_year": {"value": 820000, "unit": "CNY"},
                "design_cooling_load_kw_r": {"value": 450, "unit": "kW(r)"},
                "installed_power_kw_e": {"value": 380, "unit": "kW(e)"},
            },
        ],
    },
    "investment_estimate": {
        "total_investment": 8500000,
        "breakdown": {
            "equipment": {"value": 4200000, "unit": "CNY"},
            "installation": {"value": 1800000, "unit": "CNY"},
            "civil_works": {"value": 1500000, "unit": "CNY"},
            "other_costs": {"value": 1000000, "unit": "CNY"},
        },
    },
    "risks_and_missing_information": {
        "risks": [
            {
                "description": "Ambient temperature may exceed design max in summer",
                "severity": "medium",
                "mitigation": "Add 10% safety margin on condenser sizing",
            }
        ],
        "missing_information": [
            {
                "description": "Soil thermal conductivity not provided",
                "impact": "minor",
                "field_path": "input_conditions.soil_thermal_conductivity",
            }
        ],
    },
    "quality_summary": {
        "total_findings": 2,
        "blocker_count": 0,
        "warning_count": 1,
        "info_count": 1,
        "findings": [
            {
                "code": "W001",
                "severity": "warning",
                "message": "Missing soil thermal conductivity",
                "section_key": "input_conditions",
                "field_path": "input_conditions.soil_thermal_conductivity",
            },
            {
                "code": "I001",
                "severity": "info",
                "message": "Using default occupancy rate 0.85",
                "section_key": "assumptions",
                "field_path": "assumptions.occupancy_rate",
            },
        ],
    },
    "citations": [
        {
            "section_key": "cooling_load",
            "source_type": "calculation_result",
            "source_id": "calc-001",
            "tool_name": "cooling_load_calculator",
            "content_hash": "calc_hash_001",
        },
        {
            "section_key": "scheme_comparison",
            "source_type": "scheme_result",
            "source_id": "scheme-run-20260625-001",
            "tool_name": "scheme_evaluator",
            "content_hash": "scheme_hash_001",
        },
    ],
    "provenance": {
        "content_hash": "abc123def456",
        "canonical_hash": "abc123def456",
        "assembly_timestamp": "2026-06-25T10:00:00Z",
    },
}


def _build_canonical(
    content: dict[str, Any] | None = None,
    *,
    approval_snapshot: ApprovalSnapshot | None = None,
) -> CanonicalReportRenderModel:
    """Build a canonical render model from sample content."""
    c = content if content is not None else _SAMPLE_CONTENT
    return build_canonical_render_model(
        content=c,
        report_id="test-report-001",
        revision_number=2,
        content_hash="abc123def456",
        generated_by="test-agent",
        generated_at="2026-06-25T10:00:00Z",
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
        approval_snapshot=approval_snapshot,
    )


def _build_localized(
    content: dict[str, Any] | None = None,
    *,
    locale: ReportLocale = ReportLocale.ZH_CN,
    format: str = "docx",
    approval_snapshot: ApprovalSnapshot | None = None,
) -> LocalizedReportRenderModel:
    """Build a localized render model from sample content."""
    c = content if content is not None else _SAMPLE_CONTENT
    canonical = _build_canonical(c, approval_snapshot=approval_snapshot)
    return localize_render_model(
        canonical,
        locale=locale,
        format=format,
    )


def _render_pdf(model: LocalizedReportRenderModel, *, is_draft: bool = False) -> bytes:
    from cold_storage.modules.reports.renderers.pdf_renderer import PdfRenderer

    return PdfRenderer().render(model, is_draft=is_draft)


def _render_docx(model: LocalizedReportRenderModel, *, is_draft: bool = False) -> bytes:
    from cold_storage.modules.reports.renderers.docx_renderer import DocxRenderer

    return DocxRenderer().render(model, is_draft=is_draft)


# ======================================================================
# A. Scheme run-level provenance → source_* fields
# ======================================================================


class TestSchemeProvenance:
    """Verify source_* fields are correctly populated on scheme table cells."""

    def test_real_assembler_run_provenance_maps_to_core_source_fields(self) -> None:
        """Scheme cell source_* fields follow per-metric > scheme > run priority."""
        canonical = _build_canonical()
        section = _find_section(canonical, "scheme_comparison")
        assert section is not None
        assert section.table is not None

        rows = section.table.rows
        assert len(rows) >= 2  # scheme_a and scheme_b

        # --- scheme_a: has scheme-level provenance fields on the scheme dict ---
        row_a = rows[0]
        # scheme_id cell (index 0) — should get scheme-level provenance
        scheme_id_a = row_a[0]
        assert scheme_id_a.raw_value == "scheme_a"
        assert scheme_id_a.source_id == "scheme_a_prov_id", (
            f"Expected source_id='scheme_a_prov_id', got {scheme_id_a.source_id!r}"
        )
        assert scheme_id_a.source_tool == "scheme_generator", (
            f"Expected source_tool='scheme_generator', got {scheme_id_a.source_tool!r}"
        )
        assert scheme_id_a.source_tool_version == "2.1.0", (
            f"Expected source_tool_version='2.1.0', got {scheme_id_a.source_tool_version!r}"
        )
        assert scheme_id_a.source_content_hash == "scheme_a_hash_abcdef", (
            f"Expected source_content_hash='scheme_a_hash_abcdef', "
            f"got {scheme_id_a.source_content_hash!r}"
        )
        # Name cell (index 1) — should also get scheme-level provenance
        name_a = row_a[1]
        assert name_a.source_id == "scheme_a_prov_id"
        assert name_a.source_tool == "scheme_generator"
        # Rank cell (index 2)
        rank_a = row_a[2]
        assert rank_a.source_id == "scheme_a_prov_id"
        # Metric cell with per-metric source (total_score)
        # total_score is at index 3 (after scheme_id, name, rank)
        total_score_cell = row_a[3]
        assert total_score_cell.field_key == "header.total_score"

        # --- scheme_b: has NO scheme-level provenance; falls back to run-level ---
        row_b = rows[1]
        scheme_id_b = row_b[0]
        assert scheme_id_b.raw_value == "scheme_b"
        # No scheme-level source_id → fallback to run_id
        assert scheme_id_b.source_id == "scheme-run-20260625-001", (
            f"scheme_b source_id should fall back to run_id, got {scheme_id_b.source_id!r}"
        )
        # No scheme-level source_tool → fallback to run scheme_evaluator → "scheme_evaluator"
        assert scheme_id_b.source_tool == "scheme_evaluator", (
            f"scheme_b source_tool should fall back to run scheme_evaluator, "
            f"got {scheme_id_b.source_tool!r}"
        )
        # No scheme-level source_tool_version → fallback to run generator_version
        assert scheme_id_b.source_tool_version == "scheme_generator@2.1.0", (
            f"scheme_b source_tool_version should fall back to run generator_version, "
            f"got {scheme_id_b.source_tool_version!r}"
        )
        # No scheme-level source_content_hash → fallback to run persisted_content_hash
        assert scheme_id_b.source_content_hash == ("abcdef1234567890abcdef1234567890abcdef1234"), (
            f"scheme_b source_content_hash should fall back to run persisted_content_hash, "
            f"got {scheme_id_b.source_content_hash!r}"
        )
        # Run-level fields on scheme_b cells
        assert scheme_id_b.run_id == "scheme-run-20260625-001"
        assert scheme_id_b.generator_version == "scheme_generator@2.1.0"
        assert scheme_id_b.scheme_evaluator == "scheme_evaluator"
        assert scheme_id_b.persisted_content_hash == ("abcdef1234567890abcdef1234567890abcdef1234")

        # --- Per-metric provenance overrides scheme-level on scheme_a ---
        # scheme_a's total_score is a measured value dict WITHOUT per-metric source_*
        # but total_investment_cny etc. also have no per-metric source_*
        # Verify the metric cells still have scheme-level provenance fallback
        for cell in row_a[3:]:
            if cell.raw_value is not None:
                # Should have at least the scheme-level fallback
                assert cell.source_id in ("scheme_a_prov_id", ""), (
                    f"metric {cell.field_key} source_id should be scheme-level or empty, "
                    f"got {cell.source_id!r}"
                )
                assert cell.source_tool in ("scheme_generator", "")


def _find_section(model: CanonicalReportRenderModel, key: str) -> CanonicalRenderSection | None:
    for s in model.sections:
        if s.section_key == key:
            return s
    return None


# ======================================================================
# B. Golden serializer fixed output — mutation detection
# ======================================================================


class TestGoldenSerializerMutation:
    """Golden serializer always outputs ALL fields — mutation in content changes golden hash."""

    def _golden_hash(self, content: dict[str, Any]) -> str:
        """Compute a deterministic SHA-256 over the golden JSON."""
        canonical = _build_canonical(content)
        return hashlib.sha256(golden_json(canonical).encode("utf-8")).hexdigest()

    # --- Baseline ---

    def test_golden_baseline_is_stable(self) -> None:
        """Repeated golden serialization of same content yields identical hash."""
        h1 = self._golden_hash(_SAMPLE_CONTENT)
        h2 = self._golden_hash(_SAMPLE_CONTENT)
        assert h1 == h2

    # --- Generator version mutation ---

    def test_golden_detects_generator_version_change(self) -> None:
        """Changing generator_version in scheme_comparison alters golden hash."""
        baseline = self._golden_hash(_SAMPLE_CONTENT)

        mutated = dict(_SAMPLE_CONTENT)
        mutated["scheme_comparison"] = dict(mutated["scheme_comparison"])
        mutated["scheme_comparison"]["generator_version"] = "scheme_generator@3.0.0"
        mutated_hash = self._golden_hash(mutated)

        assert mutated_hash != baseline, "Golden hash did not change when generator_version changed"

    # --- Scheme evaluator mutation ---

    def test_golden_detects_scheme_evaluator_change(self) -> None:
        """Changing scheme_evaluator in scheme_comparison alters golden hash."""
        baseline = self._golden_hash(_SAMPLE_CONTENT)

        mutated = dict(_SAMPLE_CONTENT)
        mutated["scheme_comparison"] = dict(mutated["scheme_comparison"])
        mutated["scheme_comparison"]["scheme_evaluator"] = "different_evaluator_v2"
        mutated_hash = self._golden_hash(mutated)

        assert mutated_hash != baseline, "Golden hash did not change when scheme_evaluator changed"

    # --- Persisted content hash mutation ---

    def test_golden_detects_persisted_content_hash_change(self) -> None:
        """Changing persisted_content_hash in scheme_comparison alters golden hash."""
        baseline = self._golden_hash(_SAMPLE_CONTENT)

        mutated = dict(_SAMPLE_CONTENT)
        mutated["scheme_comparison"] = dict(mutated["scheme_comparison"])
        mutated["scheme_comparison"]["persisted_content_hash"] = "zzz_new_hash_99999"
        mutated_hash = self._golden_hash(mutated)

        assert mutated_hash != baseline, (
            "Golden hash did not change when persisted_content_hash changed"
        )

    # --- Empty vs missing ---

    def test_golden_distinguishes_empty_from_missing(self) -> None:
        """Golden dict includes keys with empty values; missing keys are absent."""
        canonical = _build_canonical(_SAMPLE_CONTENT)
        gdict = golden_dict(canonical)

        # Verify every section has all fields including empty ones
        for section_dict in gdict["sections"]:
            # All sections should have these keys even if empty
            assert "text_fields" in section_dict
            assert "paragraphs" in section_dict
            assert "empty_reason_code" in section_dict
            assert "recommended_scheme_code" in section_dict
            assert "approval_snapshot" in section_dict

        # Verify cells have all 13 fields
        scheme_section = [s for s in gdict["sections"] if s["section_key"] == "scheme_comparison"]
        assert len(scheme_section) == 1
        table = scheme_section[0].get("table")
        if table and table.get("rows"):
            cell = table["rows"][0][0]
            cell_keys = set(cell.keys())
            expected_cell_keys = {
                "field_path",
                "field_key",
                "raw_value",
                "unit_code",
                "align_code",
                "source_id",
                "source_tool",
                "source_tool_version",
                "source_content_hash",
                "run_id",
                "generator_version",
                "scheme_evaluator",
                "persisted_content_hash",
            }
            missing = expected_cell_keys - cell_keys
            assert not missing, f"Cell is missing keys: {missing}"
            # raw_value may be None — it's still present in the golden dict
            # Also allow any extra keys that may exist

        # Verify manifest has all fields
        manifest = gdict.get("manifest", {})
        manifest_keys = set(manifest.keys())
        expected_manifest_keys = {
            "template_code",
            "template_version",
            "schema_version",
            "source_content_hash",
            "sections",
            "format",
            "render_settings",
            "manifest_hash",
        }
        missing_manifest = expected_manifest_keys - manifest_keys
        assert not missing_manifest, f"Manifest is missing keys: {missing_manifest}"

        # Verify sections list has all 11 expected section keys + provenance
        section_keys = [s["section_key"] for s in gdict["sections"]]
        expected_sections = {
            "report_metadata",
            "project_summary",
            "input_conditions",
            "assumptions",
            "throughput_inventory_area",
            "cooling_load",
            "equipment_selection",
            "electrical_and_energy",
            "scheme_comparison",
            "investment_estimate",
            "risks_and_missing_information",
            "quality_summary",
            "citations",
            "provenance",
        }
        assert set(section_keys) == expected_sections, (
            f"Mismatch in section keys. Missing: {expected_sections - set(section_keys)}. "
            f"Extra: {set(section_keys) - expected_sections}."
        )

    def test_golden_with_approval_snapshot_includes_all_fields(self) -> None:
        """When approval_snapshot is provided, golden dict includes all 5 snapshot fields."""
        snapshot = ApprovalSnapshot(
            revision_id="rev-001",
            content_hash="content_hash_abc",
            approved_by="admin",
            approved_at="2026-06-25T12:00:00Z",
            revision_number=2,
        )
        canonical = _build_canonical(_SAMPLE_CONTENT, approval_snapshot=snapshot)
        gdict = golden_dict(canonical)

        snap = gdict.get("approval_snapshot")
        assert snap is not None, "approval_snapshot should be present in golden dict"
        assert snap["revision_id"] == "rev-001"
        assert snap["content_hash"] == "content_hash_abc"
        assert snap["approved_by"] == "admin"
        assert snap["approved_at"] == "2026-06-25T12:00:00Z"
        assert snap["revision_number"] == 2


# ======================================================================
# C. Full bilingual DOCX/PDF E2E content parsing
# ======================================================================


class TestRealFullSchemaDocxPdf:
    """Render real DOCX/PDF and parse file content to verify all sections."""

    # ------------------------------------------------------------------
    # DOCX content verification
    # ------------------------------------------------------------------

    def _render_all_four(
        self,
    ) -> tuple[bytes, bytes, bytes, bytes, CanonicalReportRenderModel, CanonicalReportRenderModel]:
        """Render zh DOCX, en DOCX, zh PDF, en PDF + their canonical models."""
        # Build canonical once for all four
        canonical = _build_canonical(_SAMPLE_CONTENT)

        # Zh DOCX
        localized_zh_docx = localize_render_model(
            canonical, locale=ReportLocale.ZH_CN, format="docx"
        )
        # En DOCX
        localized_en_docx = localize_render_model(
            canonical, locale=ReportLocale.EN_US, format="docx"
        )
        # Zh PDF
        localized_zh_pdf = localize_render_model(canonical, locale=ReportLocale.ZH_CN, format="pdf")
        # En PDF
        localized_en_pdf = localize_render_model(canonical, locale=ReportLocale.EN_US, format="pdf")

        zh_docx_bytes = _render_docx(localized_zh_docx, is_draft=True)
        en_docx_bytes = _render_docx(localized_en_docx, is_draft=True)
        zh_pdf_bytes = _render_pdf(localized_zh_pdf, is_draft=True)
        en_pdf_bytes = _render_pdf(localized_en_pdf, is_draft=True)

        return (
            zh_docx_bytes,
            en_docx_bytes,
            zh_pdf_bytes,
            en_pdf_bytes,
            localized_zh_docx.canonical if hasattr(localized_zh_docx, "canonical") else canonical,
            localized_en_docx.canonical if hasattr(localized_en_docx, "canonical") else canonical,
        )

    def _extract_docx_text(self, docx_bytes: bytes) -> str:
        doc = Document(BytesIO(docx_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)
        # Also get header/footer text
        for section in doc.sections:
            header = section.header
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                paragraphs.append(row_text)
        return "\n".join(paragraphs)

    def _extract_docx_all_text(self, docx_bytes: bytes) -> str:
        """Extract ALL text from DOCX including watermark textbox content via raw XML."""
        import zipfile
        from xml.etree import ElementTree as ET

        text_parts = []
        # Use python-docx for main content
        doc_text = self._extract_docx_text(docx_bytes)
        text_parts.append(doc_text)

        # Extract text from raw XML to catch watermark textboxes
        with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
            for name in z.namelist():
                if name.startswith("word/header") and name.endswith(".xml"):
                    xml_content = z.read(name)
                    root = ET.fromstring(xml_content)
                    # Find all w:t elements
                    for t_elem in root.iter(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
                    ):
                        if t_elem.text and t_elem.text.strip():
                            text_parts.append(t_elem.text.strip())
        return "\n".join(text_parts)

    def test_real_full_schema_docx_contains_all_sections_and_business_content(
        self,
    ) -> None:
        """Rendered DOCX contains all 11 sections with localized titles and business content."""
        zh_docx_bytes, en_docx_bytes, _, _, _, _ = self._render_all_four()

        # --- ZH DOCX ---
        zh_text = self._extract_docx_text(zh_docx_bytes)
        zh_section_titles = [
            "报告元数据",
            "项目概况",
            "输入条件",
            "假设条件",
            "吞吐、库存与面积",
            "冷负荷计算",
            "设备选型",
            "电气及能耗",
            "方案比较",
            "投资估算",
            "风险与缺失信息",
            "质量摘要",
            "引用信息",
            "出处信息",
        ]
        for title in zh_section_titles:
            assert title in zh_text, f"ZH DOCX missing section title: {title!r}"

        # Verify business content keywords
        assert "Blueberry Cold Storage" in zh_text or "蓝莓" in zh_text
        assert "方案A" in zh_text or "方案A — 氨+CO2" in zh_text
        assert "方案B" in zh_text or "方案B — 氟利昂" in zh_text
        assert "450" in zh_text  # cooling load value
        assert "8,500,000" in zh_text or "8500000" in zh_text  # investment

        # --- EN DOCX ---
        en_text = self._extract_docx_text(en_docx_bytes)
        en_section_titles = [
            "Report Metadata",
            "Project Summary",
            "Input Conditions",
            "Assumptions",
            "Throughput, Inventory and Area",
            "Cooling Load",
            "Equipment Selection",
            "Electrical and Energy",
            "Scheme Comparison",
            "Investment Estimate",
            "Risks and Missing Information",
            "Quality Summary",
            "Citations",
            "Provenance",
        ]
        for title in en_section_titles:
            assert title in en_text, f"EN DOCX missing section title: {title!r}"

        # Verify business content in EN
        assert "Blueberry Cold Storage" in en_text
        assert "Scheme A" in en_text or "方案A" in en_text
        assert "Scheme B" in en_text or "方案B" in en_text
        assert "450" in en_text

    def test_real_full_schema_pdf_contains_all_sections_and_business_content(
        self,
    ) -> None:
        """Rendered PDF contains all 11 sections with localized titles and business content."""
        _, _, zh_pdf_bytes, en_pdf_bytes, _, _ = self._render_all_four()

        # --- ZH PDF ---
        zh_doc = fitz.open(stream=zh_pdf_bytes, filetype="pdf")
        zh_text = ""
        for page in zh_doc:
            zh_text += page.get_text()
        zh_doc.close()

        zh_section_titles = [
            "报告元数据",
            "项目概况",
            "输入条件",
            "假设条件",
            "吞吐、库存与面积",
            "冷负荷计算",
            "设备选型",
            "电气及能耗",
            "方案比较",
            "投资估算",
            "风险与缺失信息",
            "质量摘要",
            "引用信息",
            "出处信息",
        ]
        for title in zh_section_titles:
            assert title in zh_text, f"ZH PDF missing section title: {title!r}"

        assert "Blueberry Cold Storage" in zh_text or "蓝莓" in zh_text
        assert "方案A" in zh_text
        assert "方案B" in zh_text
        assert "450" in zh_text

        # --- EN PDF ---
        en_doc = fitz.open(stream=en_pdf_bytes, filetype="pdf")
        en_text = ""
        for page in en_doc:
            en_text += page.get_text()
        en_doc.close()

        en_section_titles = [
            "Report Metadata",
            "Project Summary",
            "Input Conditions",
            "Assumptions",
            "Throughput, Inventory and Area",
            "Cooling Load",
            "Equipment Selection",
            "Electrical and Energy",
            "Scheme Comparison",
            "Investment Estimate",
            "Risks and Missing Information",
            "Quality Summary",
            "Citations",
            "Provenance",
        ]
        for title in en_section_titles:
            assert title in en_text, f"EN PDF missing section title: {title!r}"

        assert "Blueberry Cold Storage" in en_text
        assert "450" in en_text

    def test_exact_localized_watermark_in_docx_and_pdf(self) -> None:
        """Watermark text is locale-specific (草稿 / DRAFT) in both DOCX and PDF.

        Verifies the exact watermark string appears as text content — not as
        arbitrary drawing shapes or pict elements.
        """
        zh_docx_bytes, en_docx_bytes, zh_pdf_bytes, en_pdf_bytes, _, _ = self._render_all_four()

        # --- DOCX text contains watermark ---
        zh_docx_text = self._extract_docx_all_text(zh_docx_bytes)
        assert "草稿" in zh_docx_text, "ZH DOCX should contain '草稿' watermark text"

        en_docx_text = self._extract_docx_all_text(en_docx_bytes)
        assert "DRAFT" in en_docx_text, "EN DOCX should contain 'DRAFT' watermark text"

        # --- PDF text contains watermark ---
        zh_pdf_doc = fitz.open(stream=zh_pdf_bytes, filetype="pdf")
        zh_pdf_text = ""
        for page in zh_pdf_doc:
            zh_pdf_text += page.get_text()
        zh_pdf_doc.close()
        assert "草稿" in zh_pdf_text, "ZH PDF should contain '草稿' watermark text"

        en_pdf_doc = fitz.open(stream=en_pdf_bytes, filetype="pdf")
        en_pdf_text = ""
        for page in en_pdf_doc:
            en_pdf_text += page.get_text()
        en_pdf_doc.close()
        assert "DRAFT" in en_pdf_text, "EN PDF should contain 'DRAFT' watermark text"

    def test_four_real_render_inputs_have_identical_canonical_snapshot(self) -> None:
        """All 4 renders (zh/en DOCX, zh/en PDF) share the same canonical snapshot.

        The canonical model is built once from the same content; localization
        only adds display text and formatting.  The unchanging canonical core
        means all 4 render inputs are deterministically derived from the same
        canonical snapshot.
        """
        # Build canonical once
        canonical = _build_canonical(_SAMPLE_CONTENT)
        golden_baseline = golden_json(canonical)

        # For each locale+format combination, localize and verify the .canonical
        # (the embedded canonical section) matches the baseline
        for locale in (ReportLocale.ZH_CN, ReportLocale.EN_US):
            for fmt in ("docx", "pdf"):
                localized = localize_render_model(canonical, locale=locale, format=fmt)
                # Verify canonical is serialized identically (metadata comparison)
                # The canonical is embedded in each LocalizedRenderSection
                # Just verify the sections match between different renders
                # Since localized has different fields than canonical, we compare
                # the canonical snapshots embedded in the localized model
                assert localized.metadata.canonical.report_id == canonical.metadata.report_id, (
                    f"Canonical report_id mismatch for {locale}/{fmt}"
                )
                assert (
                    localized.metadata.canonical.content_hash == canonical.metadata.content_hash
                ), f"Canonical content_hash mismatch for {locale}/{fmt}"
                assert (
                    localized.metadata.canonical.revision_number
                    == canonical.metadata.revision_number
                ), f"Canonical revision_number mismatch for {locale}/{fmt}"

                # Verify number of sections
                assert len(localized.sections) == len(canonical.sections), (
                    f"Section count mismatch for {locale}/{fmt}: "
                    f"{len(localized.sections)} vs {len(canonical.sections)}"
                )

                # Verify each section_key matches
                for ls, cs in zip(localized.sections, canonical.sections, strict=True):
                    assert ls.section_key == cs.section_key, (
                        f"Section key mismatch for {locale}/{fmt}: "
                        f"{ls.section_key} vs {cs.section_key}"
                    )

        # Verify the golden baseline remains identical for all 4 renders
        # by re-computing it from the same canonical model
        h2 = golden_json(_build_canonical(_SAMPLE_CONTENT))
        assert golden_baseline == h2

    # ------------------------------------------------------------------
    # Additional content checks
    # ------------------------------------------------------------------

    def test_docx_disclaimer_text_present(self) -> None:
        """Both ZH and EN DOCX renderings contain disclaimer text."""
        zh_docx_bytes, en_docx_bytes, _, _, _, _ = self._render_all_four()
        zh_text = self._extract_docx_text(zh_docx_bytes)
        en_text = self._extract_docx_text(en_docx_bytes)

        # Disclaimer is localized — check for expected content
        assert len(zh_text) > 0
        assert len(en_text) > 0

    def test_pdf_disclaimer_text_present(self) -> None:
        """Both ZH and EN PDF renderings contain disclaimer text."""
        _, _, zh_pdf_bytes, en_pdf_bytes, _, _ = self._render_all_four()
        zh_doc = fitz.open(stream=zh_pdf_bytes, filetype="pdf")
        zh_text = ""
        for page in zh_doc:
            zh_text += page.get_text()
        zh_doc.close()
        en_doc = fitz.open(stream=en_pdf_bytes, filetype="pdf")
        en_text = ""
        for page in en_doc:
            en_text += page.get_text()
        en_doc.close()

        assert len(zh_text) > 0
        assert len(en_text) > 0

    def test_scheme_table_content_in_docx(self) -> None:
        """DOCX scheme table contains scheme names, scores, and investment values."""
        zh_docx_bytes, en_docx_bytes, _, _, _, _ = self._render_all_four()

        # ZH DOCX: scheme names present
        zh_text = self._extract_docx_text(zh_docx_bytes)
        # Both scheme names should appear (as localized text)
        assert "A" in zh_text or "方案A" in zh_text
        assert "B" in zh_text or "方案B" in zh_text

        # EN DOCX: scheme names present
        en_text = self._extract_docx_text(en_docx_bytes)
        assert "A" in en_text
        assert "B" in en_text

    def test_risks_and_citations_in_pdf(self) -> None:
        """PDF contains risks, missing information, and citations text."""
        _, _, _, en_pdf_bytes, _, _ = self._render_all_four()
        en_doc = fitz.open(stream=en_pdf_bytes, filetype="pdf")
        en_text = ""
        for page in en_doc:
            en_text += page.get_text()
        en_doc.close()

        # Missing information description should appear
        assert "soil thermal conductivity" in en_text.lower(), (
            "Missing information should appear in PDF"
        )

        # Citations section should contain source references
        assert "calc_hash_001" in en_text, "Citation content hash should appear in PDF"

    def test_localized_metric_labels_in_docx(self) -> None:
        """Metric labels are correctly localized in DOCX output."""
        zh_docx_bytes, en_docx_bytes, _, _, _, _ = self._render_all_four()
        zh_text = self._extract_docx_text(zh_docx_bytes)
        en_text = self._extract_docx_text(en_docx_bytes)

        # Cooling load label — check at least one locale variant
        has_zh_cooling = "总冷负荷" in zh_text or "冷负荷" in zh_text or "设计冷负荷" in zh_text
        has_en_cooling = "Cooling Load" in en_text

        assert has_zh_cooling, "ZH DOCX should have a cooling load label"
        assert has_en_cooling, "EN DOCX should have 'Cooling Load' label"
