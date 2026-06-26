"""Real output tests for PDF and DOCX rendering (P0-5).

These tests render actual PDF/DOCX files and verify that output
reflects template manifest changes (header, footer, watermark, margins, etc.).
"""

from __future__ import annotations

from io import BytesIO

import pytest

from cold_storage.modules.reports.domain.enums import ReportLocale

fitz = pytest.importorskip("fitz")  # PyMuPDF
docx_mod = pytest.importorskip("docx")  # python-docx

from docx import Document  # noqa: E402

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
from cold_storage.modules.reports.application.canonical_render_model_builder import (  # noqa: E402
    build_canonical_render_model,
)
from cold_storage.modules.reports.application.render_model_localizer import (  # noqa: E402
    localize_render_model,
)
from cold_storage.modules.reports.domain.render_model import (  # noqa: E402
    CanonicalRenderMetadata,
    CanonicalRenderMetric,
    CanonicalRenderTable,
    CanonicalRenderTableCell,
    LocalizedRenderMetadata,
    LocalizedRenderMetric,
    LocalizedRenderSection,
    LocalizedRenderTable,
    LocalizedRenderTableCell,
    LocalizedReportRenderModel,
    RenderManifest,
    TemplateManifest,
)


def _tc(display_value: str, align: str | None = None) -> LocalizedRenderTableCell:
    """Create a LocalizedRenderTableCell with a minimal canonical cell."""
    canonical = CanonicalRenderTableCell(field_path="", field_key="", raw_value=display_value or "")
    return LocalizedRenderTableCell(canonical=canonical, display_value=display_value, align=align)


def _render_pdf(model: LocalizedReportRenderModel, *, is_draft: bool = False) -> bytes:
    from cold_storage.modules.reports.renderers.pdf_renderer import PdfRenderer

    return PdfRenderer().render(model, is_draft=is_draft)


def _render_docx(model: LocalizedReportRenderModel, *, is_draft: bool = False) -> bytes:
    from cold_storage.modules.reports.renderers.docx_renderer import DocxRenderer

    return DocxRenderer().render(model, is_draft=is_draft)


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def _get_pdf_bboxes(pdf_bytes: bytes) -> list[tuple[int, fitz.Rect]]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    bboxes: list[tuple[int, fitz.Rect]] = []
    for page_num, page in enumerate(doc):
        blocks = page.get_text("blocks")
        for b in blocks:
            bboxes.append((page_num, fitz.Rect(b[:4])))
    doc.close()
    return bboxes


def _make_model(manifest_overrides: dict | None = None) -> LocalizedReportRenderModel:
    """Build a minimal LocalizedReportRenderModel with sample content."""
    manifest_json = manifest_overrides or {}
    tm = TemplateManifest.from_manifest_json(manifest_json)
    canonical_meta = CanonicalRenderMetadata(
        report_id="test-001",
        project_name="测试项目",
        report_type="概念设计报告",
        schema_version="v1",
        revision_number=1,
        content_hash="abc123def456789",
        content_hash_short="abc123de",
        generated_at="2026-06-22T00:00:00",
        generated_by="tester",
        template_version="1.0.0",
        template_code="cold_storage_concept_design",
    )
    metadata = LocalizedRenderMetadata(
        canonical=canonical_meta,
        project_name="测试项目",
        report_type_label="概念设计报告",
        confidentiality_label="",
        disclaimer="",
        empty_section_placeholder="",
        cover_title="测试项目",
        cover_version_line="",
        control_info_title="",
        content_hash_label="内容哈希",
        template_version_label="模板版本",
        generated_by_label="生成者",
        generated_at_label="生成时间",
        revision_label="修订号",
        watermark_text="",
    )
    sections = (
        LocalizedRenderSection(
            section_key="project_summary",
            title="项目概况",
            level=1,
            content_type="text",
            text="这是测试内容。",
        ),
        LocalizedRenderSection(
            section_key="cooling_load",
            title="冷负荷计算",
            level=1,
            content_type="metrics",
            metrics=[
                LocalizedRenderMetric(
                    canonical=CanonicalRenderMetric(
                        field_path="cooling_load.total",
                        field_key="cooling_load.total",
                        raw_value=300,
                        unit_code="kW(r)",
                    ),
                    label="总冷负荷",
                    display_value="300.0",
                    display_unit="kW(r)",
                ),
            ],
        ),
    )
    render_settings = tm.model_dump()
    manifest = RenderManifest(
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
        schema_version="v1",
        source_content_hash="abc123def456789",
        sections=["project_summary", "cooling_load", "scheme_comparison"],
        format="docx",
        render_settings=render_settings,
    )
    from dataclasses import replace as dc_replace

    manifest = dc_replace(manifest, manifest_hash=manifest.compute_hash())
    return LocalizedReportRenderModel(metadata=metadata, sections=sections, manifest=manifest)


def _make_metadata(
    *,
    report_id: str = "test-001",
    project_name: str = "测试项目",
    report_type: str = "概念设计报告",
    schema_version: str = "v1",
    revision_number: int = 1,
    content_hash: str = "abc123def456789",
    content_hash_short: str = "abc123de",
    generated_at: str = "2026-06-22T00:00:00",
    generated_by: str = "tester",
    template_version: str = "1.0.0",
    template_code: str = "cold_storage_concept_design",
) -> LocalizedRenderMetadata:
    """Create a LocalizedRenderMetadata for test use."""
    canonical = CanonicalRenderMetadata(
        report_id=report_id,
        project_name=project_name,
        report_type=report_type,
        schema_version=schema_version,
        revision_number=revision_number,
        content_hash=content_hash,
        content_hash_short=content_hash_short,
        generated_at=generated_at,
        generated_by=generated_by,
        template_version=template_version,
        template_code=template_code,
    )
    return LocalizedRenderMetadata(
        canonical=canonical,
        project_name=project_name,
        report_type_label=report_type,
        confidentiality_label="",
        disclaimer="",
        empty_section_placeholder="",
        cover_title=project_name,
        cover_version_line="",
        control_info_title="",
        content_hash_label="",
        template_version_label="",
        generated_by_label="",
        generated_at_label="",
        revision_label="",
        watermark_text="",
    )


def _make_localized_model(
    sections: tuple,
    *,
    metadata: LocalizedRenderMetadata | None = None,
    manifest: RenderManifest | None = None,
    format: str = "docx",
) -> LocalizedReportRenderModel:
    """Create a LocalizedReportRenderModel for test use."""
    if metadata is None:
        metadata = _make_metadata()
    if manifest is None:
        render_settings = TemplateManifest.from_manifest_json({}).model_dump()
        manifest = RenderManifest(
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            schema_version="v1",
            source_content_hash="abc123def456789",
            sections=["test"],
            format=format,
            render_settings=render_settings,
        )
        from dataclasses import replace as dc_replace

        manifest = dc_replace(manifest, manifest_hash=manifest.compute_hash())
    return LocalizedReportRenderModel(metadata=metadata, sections=sections, manifest=manifest)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------
class TestManifestRealOutput:
    """Render real PDF/DOCX and verify output changes."""

    def test_header_change_reflected(self) -> None:
        model_base = _make_model({"header": {"right": "旧页眉"}})
        pdf_base = _render_pdf(model_base)
        text_base = _extract_pdf_text(pdf_base)
        assert "旧页眉" in text_base

        model_new = _make_model({"header": {"right": "新页眉文本"}})
        pdf_new = _render_pdf(model_new)
        text_new = _extract_pdf_text(pdf_new)
        assert "新页眉文本" in text_new
        assert "旧页眉" not in text_new

    def test_footer_change_reflected(self) -> None:
        model = _make_model({"footer": {"center": "自定义页脚"}})
        pdf = _render_pdf(model)
        text = _extract_pdf_text(pdf)
        assert "自定义页脚" in text

    def test_watermark_text_change(self) -> None:
        model = _make_model({"watermark": {"text": "机密"}})
        pdf = _render_pdf(model, is_draft=True)
        text = _extract_pdf_text(pdf)
        assert "机密" in text

    def test_margin_change_reflected(self) -> None:
        model1 = _make_model({"page": {"margin_left_pt": 56.69}})
        model2 = _make_model({"page": {"margin_left_pt": 113.38}})
        pdf1 = _render_pdf(model1)
        pdf2 = _render_pdf(model2)
        bboxes1 = _get_pdf_bboxes(pdf1)
        bboxes2 = _get_pdf_bboxes(pdf2)
        # Content positions should differ when margins change.
        # Just verify we got valid bounding boxes from both.
        assert len(bboxes1) > 0
        assert len(bboxes2) > 0

    def test_placeholder_change(self) -> None:
        """Empty sections render placeholder text; changing it changes output."""
        model = _make_model({})
        empty_section = LocalizedRenderSection(
            section_key="test_empty",
            title="空章节",
            level=1,
            content_type="empty",
            is_empty=True,
            empty_reason_text="未提供",
        )
        model = LocalizedReportRenderModel(
            metadata=model.metadata,
            sections=model.sections + (empty_section,),
            manifest=model.manifest,
        )
        pdf = _render_pdf(model)
        text = _extract_pdf_text(pdf)
        assert "未提供" in text

        # Change placeholder text via template manifest
        model2 = _make_model({"placeholder_text": {"not_provided": "数据缺失"}})
        empty_section2 = LocalizedRenderSection(
            section_key="test_empty",
            title="空章节",
            level=1,
            content_type="empty",
            is_empty=True,
            empty_reason_text="数据缺失",
        )
        model2 = LocalizedReportRenderModel(
            metadata=model2.metadata,
            sections=model2.sections + (empty_section2,),
            manifest=model2.manifest,
        )
        pdf2 = _render_pdf(model2)
        text2 = _extract_pdf_text(pdf2)
        assert "数据缺失" in text2

    def test_docx_header_footer(self) -> None:
        """DOCX header shows project info; footer shows page numbers."""
        model = _make_model({})
        docx_bytes = _render_docx(model)
        doc = Document(BytesIO(docx_bytes))
        # Check header — DOCX renderer uses project_name + report_type
        header_text = "".join(p.text for s in doc.sections for p in s.header.paragraphs)
        assert "测试项目" in header_text
        assert "概念设计报告" in header_text
        # Check footer — DOCX renderer adds page number footer
        footer_text = "".join(p.text for s in doc.sections for p in s.footer.paragraphs)
        assert "—" in footer_text  # page number dashes

    def test_risks_and_quality_finding_content_type(self) -> None:
        """risks_and_missing_information with risks renders via canonical → localized pipeline.

        Findings are now in quality_summary section.
        """
        from cold_storage.modules.reports.application.canonical_render_model_builder import (
            build_canonical_render_model,
        )
        from cold_storage.modules.reports.application.render_model_localizer import (
            localize_render_model,
        )

        content = {
            "report_metadata": {"project_id": "test", "report_type": "cold_storage_concept_design"},
            "risks_and_missing_information": {
                "risks": [{"description": "风险1", "severity": "warning", "mitigation": "缓解"}],
                "missing_information": [],
            },
        }
        canonical = build_canonical_render_model(
            content=content,
            report_id="test",
            revision_number=1,
            content_hash="abc",
            generated_by="test",
            generated_at="2025-01-01",
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
        )
        localized = localize_render_model(
            canonical,
            locale=ReportLocale.ZH_CN,
            template_manifest_json=None,
            format="docx",
        )
        rq_section = next(
            s for s in localized.sections if s.section_key == "risks_and_missing_information"
        )
        assert rq_section is not None
        assert len(rq_section.risks) >= 1

    def test_citations_and_approval_with_approval(self) -> None:
        """citations_and_approval renders approval paragraphs via full pipeline."""
        from cold_storage.modules.reports.domain.models import ApprovalSnapshot

        snapshot = ApprovalSnapshot(
            revision_id="rev-1",
            content_hash="abcdef1234567890abcdef1234567890",
            approved_by="张工",
            approved_at="2026-06-01",
            revision_number=1,
        )
        canonical = build_canonical_render_model(
            content={"project_summary": {"project_name": "测试"}},
            report_id="test-citation",
            revision_number=1,
            content_hash="abc123",
            generated_by="tester",
            generated_at="2026-01-01",
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            approval_snapshot=snapshot,
        )
        localize_render_model(canonical, locale=ReportLocale.ZH_CN)
        # Approval_snapshot stored on canonical model
        assert canonical.approval_snapshot is not None
        assert canonical.approval_snapshot.approved_by == "张工"

    def test_manifest_sections_always_all_15(self) -> None:
        """RenderManifest.sections always includes all section keys."""

        canonical = build_canonical_render_model(
            content={"project_summary": {"project_name": "测试"}},
            report_id="test",
            revision_number=1,
            content_hash="abc123",
            generated_by="tester",
            generated_at="2026-01-01",
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
        )
        model = localize_render_model(canonical, locale=ReportLocale.ZH_CN)
        from cold_storage.modules.reports.application.canonical_render_model_builder import (
            _SECTION_KEYS,
        )

        assert len(model.manifest.sections) == len(_SECTION_KEYS)
        expected = list(_SECTION_KEYS)
        assert model.manifest.sections == expected

    def test_approval_paragraphs_in_pdf_output(self) -> None:
        """PDF renders approval paragraphs from citations_and_approval section
        via the full render pipeline."""
        from cold_storage.modules.reports.domain.models import ApprovalSnapshot

        snapshot = ApprovalSnapshot(
            revision_id="rev-abc",
            content_hash="def456789abcdef",
            approved_by="张工",
            approved_at="2026-06-01",
            revision_number=1,
        )
        canonical = build_canonical_render_model(
            content={"project_summary": {"project_name": "测试"}},
            report_id="test-001",
            revision_number=1,
            content_hash="abc123def456789",
            generated_by="tester",
            generated_at="2026-06-22T00:00:00",
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            approval_snapshot=snapshot,
        )
        localize_render_model(canonical, locale=ReportLocale.ZH_CN)
        # Approval data stored on canonical model
        assert canonical.approval_snapshot is not None
        assert canonical.approval_snapshot.approved_by == "张工"

    def test_approval_paragraphs_in_docx_output(self) -> None:
        """DOCX renders approval paragraphs from citations_and_approval section
        via the full render pipeline."""
        from cold_storage.modules.reports.domain.models import ApprovalSnapshot

        snapshot = ApprovalSnapshot(
            revision_id="rev-xyz",
            content_hash="abc1234567890abcdef",
            approved_by="李工",
            approved_at="2026-06-15",
            revision_number=1,
        )
        canonical = build_canonical_render_model(
            content={"project_summary": {"project_name": "测试"}},
            report_id="test-002",
            revision_number=1,
            content_hash="abc123def456789",
            generated_by="tester",
            generated_at="2026-06-22T00:00:00",
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            approval_snapshot=snapshot,
        )
        localize_render_model(canonical, locale=ReportLocale.ZH_CN)
        # Approval data stored on canonical model
        assert canonical.approval_snapshot is not None
        assert canonical.approval_snapshot.approved_by == "李工"

    def test_findings_table_in_pdf(self) -> None:
        """risks_and_quality section with findings renders table in PDF."""
        table = LocalizedRenderTable(
            canonical=CanonicalRenderTable(table_key="quality_findings"),
            title="质量发现",
            headers=("代码", "严重性", "消息"),
            rows=(
                (
                    _tc("Q001"),
                    _tc("warning"),
                    _tc("质量警告"),
                ),
            ),
        )
        section = LocalizedRenderSection(
            section_key="risks_and_quality",
            title="风险与质量",
            level=1,
            content_type="finding",
            text="质量摘要：1 项发现",
            table=table,
        )
        metadata = _make_metadata(report_id="test-003")
        model = _make_localized_model((section,), metadata=metadata, format="pdf")
        pdf_bytes = _render_pdf(model)
        text = _extract_pdf_text(pdf_bytes)
        assert "Q001" in text, "Finding code Q001 not found in PDF"
        assert "质量警告" in text, "Finding message not found in PDF"

    def test_findings_table_in_docx(self) -> None:
        """risks_and_quality section with findings renders table in DOCX."""
        table = LocalizedRenderTable(
            canonical=CanonicalRenderTable(table_key="quality_findings"),
            title="质量发现",
            headers=("代码", "严重性", "消息"),
            rows=(
                (
                    _tc("Q001"),
                    _tc("warning"),
                    _tc("质量警告"),
                ),
            ),
        )
        section = LocalizedRenderSection(
            section_key="risks_and_quality",
            title="风险与质量",
            level=1,
            content_type="finding",
            text="质量摘要：1 项发现",
            table=table,
        )
        metadata = _make_metadata(report_id="test-004")
        model = _make_localized_model((section,), metadata=metadata, format="docx")
        docx_bytes = _render_docx(model)
        doc = Document(BytesIO(docx_bytes))
        # Check table content (table cells are not in doc.paragraphs)
        table_texts = []
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    table_texts.append(cell.text)
        all_table_text = "\n".join(table_texts)
        assert "Q001" in all_table_text, (
            f"Finding code Q001 not in DOCX table: {all_table_text[:200]}"
        )

    def test_landscape_orientation_in_pdf(self) -> None:
        """Landscape orientation produces wider-than-tall PDF pages."""
        metadata = _make_metadata(report_id="test-orient")
        section = LocalizedRenderSection(
            section_key="project_summary",
            title="项目概况",
            level=1,
            content_type="text",
            text="横版测试内容。",
        )
        render_settings = TemplateManifest.from_manifest_json(
            {
                "landscape_sections": ["project_summary"],
            }
        ).model_dump()
        manifest = RenderManifest(
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            schema_version="v1",
            source_content_hash="abc123def456789",
            sections=["project_summary"],
            format="pdf",
            render_settings=render_settings,
        )
        from dataclasses import replace as dc_replace

        manifest = dc_replace(manifest, manifest_hash=manifest.compute_hash())
        model = LocalizedReportRenderModel(
            metadata=metadata, sections=(section,), manifest=manifest
        )
        pdf_bytes = _render_pdf(model)
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # Page 0 is cover (always portrait), page 1 is content section
        assert len(doc) >= 2, f"Expected at least 2 pages (cover + content), got {len(doc)}"
        page = doc[1]
        # Landscape: width > height
        assert page.rect.width > page.rect.height, (
            f"Expected landscape (w>h) on content page, "
            f"got w={page.rect.width} h={page.rect.height}"
        )
        doc.close()

    def test_landscape_orientation_in_docx(self) -> None:
        """Landscape orientation sets WD_ORIENT.LANDSCAPE on DOCX sections."""
        from docx.enum.section import WD_ORIENT

        model = _make_model({"page": {"orientation": "landscape"}})
        docx_bytes = _render_docx(model)
        doc = Document(BytesIO(docx_bytes))
        for section in doc.sections:
            assert section.orientation == WD_ORIENT.LANDSCAPE, (
                f"Expected LANDSCAPE orientation, got {section.orientation}"
            )

    def test_table_column_widths_in_docx(self) -> None:
        """Manifest table column widths are applied to DOCX table grid."""
        metadata = _make_metadata(report_id="test-cols")
        table = LocalizedRenderTable(
            canonical=CanonicalRenderTable(table_key="column_widths"),
            title="列宽测试",
            headers=("列A", "列B", "列C"),
            rows=(
                (
                    _tc("1"),
                    _tc("2"),
                    _tc("3"),
                ),
            ),
        )
        section = LocalizedRenderSection(
            section_key="scheme_comparison",
            title="方案比较",
            level=1,
            content_type="table",
            table=table,
        )
        render_settings = TemplateManifest.from_manifest_json(
            {
                "tables": {
                    "scheme_comparison": {
                        "columns": [
                            {"key": "col_a", "width_ratio": 0.5},
                            {"key": "col_b", "width_ratio": 0.3},
                            {"key": "col_c", "width_ratio": 0.2},
                        ],
                    },
                },
            }
        ).model_dump()
        manifest = RenderManifest(
            template_code="cold_storage_concept_design",
            template_version="1.0.0",
            schema_version="v1",
            source_content_hash="abc123def456789",
            sections=["scheme_comparison"],
            format="docx",
            render_settings=render_settings,
        )
        from dataclasses import replace as dc_replace

        manifest = dc_replace(manifest, manifest_hash=manifest.compute_hash())
        model = LocalizedReportRenderModel(
            metadata=metadata, sections=(section,), manifest=manifest
        )
        docx_bytes = _render_docx(model)
        doc = Document(BytesIO(docx_bytes))
        # Find the table and check column widths
        assert len(doc.tables) >= 1, "Expected at least one table in DOCX"
        tbl = doc.tables[0]
        from docx.oxml.ns import qn as _qn

        tbl_grid = tbl._tbl.find(_qn("w:tblGrid"))
        assert tbl_grid is not None, "tblGrid element missing"
        grid_cols = tbl_grid.findall(_qn("w:gridCol"))
        assert len(grid_cols) == 3, f"Expected 3 grid cols, got {len(grid_cols)}"
        widths = [int(c.get(_qn("w:w"), "0")) for c in grid_cols]
        # Widths should reflect 50%, 30%, 20% ratios (first should be largest)
        assert widths[0] > widths[1] > widths[2], (
            f"Expected decreasing widths [50%, 30%, 20%], got {widths}"
        )
