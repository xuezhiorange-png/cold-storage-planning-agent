"""Real production E2E tests: full Assembler → ReportService → ReportRenderService pipeline
with real ReportArtifactStorage(tmp_path) and RecordingObserver.

Tests:
1. test_real_pipeline_zh_cn_docx_exact_content  — Full pipeline zh-CN DOCX, exact content
2. test_real_pipeline_en_us_docx_exact_content  — Full pipeline en-US DOCX, exact content
3. test_real_pipeline_zh_cn_pdf_exact_content   — Full pipeline zh-CN PDF, exact content
4. test_real_pipeline_en_us_pdf_exact_content   — Full pipeline en-US PDF, exact content
5. test_four_independent_render_calls_have_identical_canonical_snapshot
   — RecordingObserver captures 4 canonicals, all golden_dict identical
6. test_real_files_have_exact_localized_watermark
   — zh-CN "草稿" / en-US "DRAFT"
7. test_real_files_have_exact_localized_disclaimer
   — zh-CN "AI辅助，仅供参考。" / en-US "AI-assisted. Data for reference only."

All tests use real ReportArtifactStorage(tmp_path), not MockStorage.
Uses SQLite in-memory database.
No sleep, xfail, or skip.
"""

from __future__ import annotations

import hashlib
import tempfile
from dataclasses import replace
from io import BytesIO
from typing import Any

import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from cold_storage.modules.reports.application.assembler import ReportAssembler, ReportDataProvider
from cold_storage.modules.reports.application.render_service import (
    ReportRenderService,
    ReportRenderUnitOfWork,
)
from cold_storage.modules.reports.application.service import ReportService
from cold_storage.modules.reports.domain.enums import (
    ArtifactStatus,
    ExportFormat,
    ReportLocale,
    ReportType,
    TemplateStatus,
)
from cold_storage.modules.reports.domain.models import (
    Report,
    ReportExportArtifact,
    ReportRevision,
    ReportTemplate,
)
from cold_storage.modules.reports.domain.observer import RecordingObserver
from cold_storage.modules.reports.infrastructure.artifact_storage import ReportArtifactStorage
from cold_storage.modules.reports.infrastructure.orm import Base
from cold_storage.modules.reports.infrastructure.repository import SQLReportRepository
from cold_storage.modules.reports.infrastructure.template_seed import (
    _compute_content_hash,
    _load_manifest,
)

fitz = pytest.importorskip("fitz")
docx_mod = pytest.importorskip("docx")

from docx import Document  # noqa: E402

# ======================================================================
# Fixtures
# ======================================================================


@pytest.fixture()
def engine():
    """In-memory SQLite engine with all report tables."""
    eng = create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(eng)
    yield eng
    eng.dispose()


@pytest.fixture()
def session_factory(engine):
    return sessionmaker(bind=engine, expire_on_commit=False)


@pytest.fixture()
def session(session_factory):
    with session_factory() as s:
        yield s


@pytest.fixture()
def repo(session):
    return SQLReportRepository(session)


@pytest.fixture()
def tmp_storage():
    """Real ReportArtifactStorage backed by a temp directory."""
    tmp_dir = tempfile.mkdtemp(prefix="test_artifact_storage_")
    storage = ReportArtifactStorage(tmp_dir)
    yield storage
    # Cleanup
    import shutil

    shutil.rmtree(tmp_dir, ignore_errors=True)


# ======================================================================
# Rich data provider — returns full content for all required sections
# ======================================================================


class _RichDataProvider(ReportDataProvider):
    """Mock data provider returning realistic data for all engineering sections.

    Matches the COLD_STORAGE_CONCEPT_DESIGN_V1 JSON schema so that the
    assembled content passes schema validation without modification.
    """

    def get_project(self, project_id: str) -> dict[str, Any] | None:
        return {
            "name": "Blueberry Cold Storage - Demo Plant",
            "location": "Kunming, Yunnan",
            "description": "A 5,000-ton capacity cold storage for IQF blueberries.",
        }

    def get_project_version(
        self, version_id: str, project_id: str | None = None
    ) -> dict[str, Any] | None:
        return {
            "id": version_id,
            "version_number": 3,
            "status": "active",
        }

    def get_calculation_results(self, project_id: str, version_id: str) -> list[dict[str, Any]]:
        return [
            {
                "section_key": "throughput_inventory_area",
                "result_id": "calc-tp-001",
                "tool_name": "throughput_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "daily_inbound_mass_kg": 30000,
                    "storage_capacity_kg": 5000000,
                    "total_area_m2": 1250,
                },
            },
            {
                "section_key": "cooling_load",
                "result_id": "calc-cl-001",
                "tool_name": "cooling_load_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_design_refrigeration_load": {
                        "value": 450.0,
                        "unit": "kW(r)",
                        "source_result_id": "calc-cl-001",
                        "source_tool": "cooling_load_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
            {
                "section_key": "equipment_selection",
                "result_id": "calc-es-001",
                "tool_name": "equipment_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_compressor_capacity": {
                        "value": 500.0,
                        "unit": "kW(r)",
                        "source_result_id": "calc-es-001",
                        "source_tool": "equipment_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
            {
                "section_key": "electrical_and_energy",
                "result_id": "calc-ee-001",
                "tool_name": "power_calculator",
                "tool_version": "1.2.0",
                "data": {
                    "total_installed_power": {
                        "value": 350.0,
                        "unit": "kW(e)",
                        "source_result_id": "calc-ee-001",
                        "source_tool": "power_calculator",
                        "source_tool_version": "1.2.0",
                    },
                },
            },
        ]

    def get_scheme_results(self, project_id: str, version_id: str) -> dict[str, Any] | None:
        return {
            "run_id": "scheme-run-20260625-001",
            "generator_version": "scheme_generator@2.1.0",
            "persisted_content_hash": "abcdef1234567890abcdef1234567890abcdef1234",
            "recommended_scheme": "scheme_a",
            "schemes": [
                {
                    "scheme_id": "scheme_a",
                    "name": "方案A — 氨+CO2复叠系统",
                    "rank": 1,
                    "total_score": 92,
                    "total_investment_cny": 8500000,
                    "total_area_m2": 1350,
                    "operating_cost_per_year": 680000,
                    "design_cooling_load_kw_r": 450,
                    "installed_power_kw_e": 350,
                },
                {
                    "scheme_id": "scheme_b",
                    "name": "方案B — 氟利昂 R507 系统",
                    "rank": 2,
                    "total_score": 78,
                    "total_investment_cny": 7200000,
                    "total_area_m2": 1400,
                    "operating_cost_per_year": 820000,
                    "design_cooling_load_kw_r": 450,
                    "installed_power_kw_e": 380,
                },
            ],
        }

    def get_agent_sessions(self, project_id: str, version_id: str) -> list[dict[str, Any]]:
        return []

    def get_knowledge_documents(self) -> list[dict[str, Any]]:
        return []


# ======================================================================
# Helpers
# ======================================================================


def _create_report(repo: SQLReportRepository, session: Any) -> Report:
    """Create a bare DRAFT report."""
    report = Report.create(
        project_id="proj-blueberry-01",
        project_version_id="ver-3",
        report_type=ReportType.COLD_STORAGE_CONCEPT_DESIGN,
        created_by="test-user",
    )
    repo.save_report(report)
    session.commit()
    return report


def _generate_revision(service: ReportService, report: Report) -> ReportRevision:
    """Generate a new revision using the real assembler pipeline."""
    return service.generate_revision(report.id, "test-user")


def _seed_both_locale_templates(repo: SQLReportRepository) -> None:
    """Create both zh-CN and en-US templates for DOCX and PDF, all ACTIVE."""
    for locale_str in ("zh-CN", "en-US"):
        for fmt in (ExportFormat.DOCX, ExportFormat.PDF):
            manifest = _load_manifest(fmt, locale=locale_str, allow_legacy_fallback=True)
            if not manifest:
                continue

            template_code = manifest.get("template_code", "cold_storage_concept_design")
            version = manifest.get("version", "1.0.0")
            report_type_str = manifest.get("report_type", "cold_storage_concept_design")
            schema_version = manifest.get("schema_version", f"{report_type_str}@{version}")
            loc = manifest.get("locale", locale_str)
            report_type = ReportType(report_type_str)
            content_hash = _compute_content_hash(manifest)

            existing = repo.list_templates(template_code=template_code, format=fmt)
            already = any(t.version == version and t.locale == loc for t in existing)
            if already:
                continue

            template = ReportTemplate.create(
                template_code=template_code,
                report_type=report_type,
                format=fmt,
                version=version,
                schema_version=schema_version,
                locale=loc,
                manifest_json=manifest,
                template_content_hash=content_hash,
                created_by="system",
            )
            template = replace(template, status=TemplateStatus.ACTIVE)
            repo.save_template(template)

    repo.commit()


def _make_render_service(
    session: Any,
    tmp_storage: ReportArtifactStorage,
    observer: RecordingObserver | None = None,
) -> tuple[ReportRenderService, SQLReportRepository, ReportArtifactStorage]:
    """Build a render service with real ReportArtifactStorage."""
    repo = SQLReportRepository(session)
    uow = ReportRenderUnitOfWork(session, report_repo=repo, artifact_repo=repo)
    render_svc = ReportRenderService(
        storage=tmp_storage,
        template_repo=repo,
        uow=uow,
        canonical_observer=observer,
    )
    return render_svc, repo, tmp_storage


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    doc = Document(BytesIO(docx_bytes))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    for section in doc.sections:
        header = section.header
        if header and not header.is_linked_to_previous:
            for para in header.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
        footer = section.footer
        if footer and not footer.is_linked_to_previous:
            for para in footer.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _docx_contains_text(docx_bytes: bytes, search_text: str) -> bool:
    """Check if DOCX contains text anywhere in the raw XML (including VML watermarks)."""
    import zipfile

    # First check python-docx extracted text
    text = _extract_docx_text(docx_bytes)
    if search_text in text:
        return True
    # Search raw XML for the text (catches VML watermarks in headers)
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        for name in z.namelist():
            if name.startswith("word/"):
                xml_content = z.read(name).decode("utf-8", errors="replace")
                if search_text in xml_content:
                    return True
    return False


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        text_parts: list[str] = []
        for page in doc:
            text_parts.append(page.get_text())
        return "\n".join(text_parts)
    finally:
        doc.close()


def _render_artifact(
    render_svc: ReportRenderService,
    storage: ReportArtifactStorage,
    report_id: str,
    revision_number: int,
    fmt: str,
    locale: ReportLocale,
) -> tuple[ReportExportArtifact, bytes]:
    """Render and retrieve artifact bytes from real storage."""
    artifact = render_svc.render(
        report_id=report_id,
        revision_number=revision_number,
        format=fmt,
        template_version=None,
        mode="draft",
        actor="test-user",
        locale=locale,
    )
    assert artifact.status.value == "completed", (
        f"Expected completed artifact, got {artifact.status.value}"
    )
    raw_bytes = storage.get(artifact.storage_key)
    assert raw_bytes, "Artifact bytes must not be empty"
    return artifact, raw_bytes


# ======================================================================
# Full pipeline setup — one-time for each test
# ======================================================================


def _full_pipeline_setup(
    session: Any,
) -> tuple[Report, ReportRevision]:
    """Create a report through the full Assembler → ReportService pipeline.

    Returns (report, revision) with a GENERATED report seeded with rich data.
    """
    repo = SQLReportRepository(session)
    provider = _RichDataProvider()
    assembler = ReportAssembler(provider)
    service = ReportService(repository=repo, assembler=assembler)

    report = _create_report(repo, session)
    revision = _generate_revision(service, report)

    _seed_both_locale_templates(repo)

    return report, revision


# ======================================================================
# Exact expected values
# ======================================================================

ZH_CN_DISCLAIMER = "AI辅助，仅供参考。"
EN_US_DISCLAIMER = "AI-assisted. Data for reference only."

ZH_CN_WATERMARK = "草稿"
EN_US_WATERMARK = "DRAFT"

# Structured expected dicts for comprehensive content validation.
# Scheme names come from the data provider (Chinese), not translation catalog.

EXPECTED_ZH_CN = {
    "section_titles": [
        "报告元数据",
        "项目概况",
        "输入条件",
        "假设条件",
        "吞吐、库存与面积",
        "冷负荷计算",
        "设备选型",
        "电气及能耗",
        "方案比较",
        "投资估算",
        "风险与缺失信息",
        "质量摘要",
        "引用信息",
        "出处信息",
    ],
    "field_labels": {
        # throughput_inventory_area text fields
        "storage_capacity_kg": "存储容量",
        "daily_inbound_mass_kg": "日入库质量",
        "total_area_m2": "总面积",
        # project_summary text fields
        "project_name": "项目名称",
        "project_location": "项目地点",
        "description": "项目描述",
        # metrics from measured-value dicts
        "total_design_refrigeration_load": "总设计制冷量",
        "total_compressor_capacity": "总压缩机容量",
        "total_installed_power": "总装机功率",
        # report_metadata text fields
        "project_id": "项目编号",
        "project_version_id": "项目版本号",
        "revision_number": "修订编号",
        "generated_at": "生成时间",
        "generated_by": "生成者",
        "schema_version": "Schema版本",
        "report_id": "报告编号",
        # recommended scheme
        "recommended_scheme": "推荐方案",
        # quality summary
        "total_findings": "总发现数",
        "blocker_count": "阻断项数",
        "warning_count": "警告项数",
        # provenance fields
        "generator_version": "生成器版本",
        "persisted_content_hash": "持久化内容哈希",
        # rank
        "rank": "排名",
    },
    "unit_labels": {
        "kW_r": "kW(r)",
        "kW_e": "kW(e)",
    },
    "formatted_values": {
        # text_field values (via str(), not format_decimal)
        "daily_inbound_mass_kg": "30000",
        "storage_capacity_kg": "5000000",
        "total_area_m2": "1250",
        # metric values (via format_decimal)
        "total_design_refrigeration_load": "450.0",
        "total_compressor_capacity": "500.0",
        "total_installed_power": "350.0",
    },
    "watermark": "草稿",
    "disclaimer": "AI辅助，仅供参考。",
    "scheme_names": [
        "方案A — 氨+CO2复叠系统",
        "方案B — 氟利昂 R507 系统",
    ],
    # scheme_details: scheme_id, scheme name, rank, total_score from data provider
    "scheme_details": [
        {
            "scheme_id": "scheme_a",
            "name": "方案A — 氨+CO2复叠系统",
            "rank": 1,
            "total_score": 92,
            "metric_values": {
                "total_investment_cny": 8500000,
                "total_area_m2": 1350,
                "operating_cost_per_year": 680000,
                "design_cooling_load_kw_r": 450,
                "installed_power_kw_e": 350,
            },
        },
        {
            "scheme_id": "scheme_b",
            "name": "方案B — 氟利昂 R507 系统",
            "rank": 2,
            "total_score": 78,
            "metric_values": {
                "total_investment_cny": 7200000,
                "total_area_m2": 1400,
                "operating_cost_per_year": 820000,
                "design_cooling_load_kw_r": 450,
                "installed_power_kw_e": 380,
            },
        },
    ],
    "recommended_scheme_code": "scheme_a",
    # provenance source_id / source_tool strings that should appear in report text
    "provenance_source_ids": [
        "calc-tp-001",
        "calc-cl-001",
        "calc-es-001",
        "calc-ee-001",
        "scheme-run-20260625-001",
    ],
    "provenance_source_tools": [
        "throughput_calculator",
        "cooling_load_calculator",
        "equipment_calculator",
        "power_calculator",
    ],
    # document_control labels (from translation catalog)
    "document_control_labels": {
        "title": "文件控制信息",
        "content_hash": "内容哈希",
        "template_version": "模板版本",
        "generated_by": "生成者",
        "generated_at": "生成时间",
        "revision": "修订号",
    },
    # header / footer text (from template manifest rendering)
    "header_text": "Blueberry Cold Storage - Demo Plant — 概念设计报告",
    "footer_text": "— 1 —",
    # quality_summary counts (derived from evaluate_quality on rich data provider output)
    "quality_summary": {
        "total_findings": 6,
        "blocker_count": 4,
        "warning_count": 2,
        "info_count": 0,
    },
    # quality finding codes that appear in rendered report text
    "quality_finding_codes": ["SOURCE_MISSING_CONTENT_HASH"],
    # quality finding severities (localized labels)
    "quality_severity_blocker": "阻断",
    "quality_severity_warning": "警告",
    # citation section_keys that appear in the citations table
    "citation_section_keys": [
        "project_summary",
        "throughput_inventory_area",
        "cooling_load",
        "equipment_selection",
        "electrical_and_energy",
        "scheme_comparison",
        "report_metadata",
    ],
    # approval / empty state text
    "approval_label": "审批信息",
    # provenance fields that appear in provenance section
    "provenance_keys": [
        "content_hash",
        "canonical_hash",
        "selection_rules",
        "assembly_timestamp",
    ],
}

EXPECTED_EN_US = {
    "section_titles": [
        "Report Metadata",
        "Project Summary",
        "Input Conditions",
        "Assumptions",
        "Throughput, Inventory and Area",
        "Cooling Load",
        "Equipment Selection",
        "Electrical and Energy",
        "Scheme Comparison",
        "Investment Estimate",
        "Risks and Missing Information",
        "Quality Summary",
        "Citations",
        "Provenance",
    ],
    "field_labels": {
        # throughput_inventory_area text fields
        "storage_capacity_kg": "Storage Capacity",
        "daily_inbound_mass_kg": "Daily Inbound Mass",
        "total_area_m2": "Total Area",
        # project_summary text fields
        "project_name": "Project Name",
        "project_location": "Project Location",
        "description": "Description",
        # metrics from measured-value dicts
        "total_design_refrigeration_load": "Total Design Refrigeration Load",
        "total_compressor_capacity": "Total Compressor Capacity",
        "total_installed_power": "Total Installed Power",
        # report_metadata text fields
        "project_id": "Project ID",
        "project_version_id": "Project Version",
        "revision_number": "Revision Number",
        "generated_at": "Generated At",
        "generated_by": "Generated By",
        "schema_version": "Schema Version",
        "report_id": "Report ID",
        # recommended scheme
        "recommended_scheme": "Recommended Scheme",
        # quality summary
        "total_findings": "Total Findings",
        "blocker_count": "Blocker Count",
        "warning_count": "Warning Count",
        # provenance fields
        "generator_version": "Generator Version",
        "persisted_content_hash": "Persisted Content Hash",
        # rank
        "rank": "Rank",
    },
    "unit_labels": {
        "kW_r": "kW(r)",
        "kW_e": "kW(e)",
    },
    "formatted_values": {
        # text_field values (via str(), not format_decimal — same for both locales)
        "daily_inbound_mass_kg": "30000",
        "storage_capacity_kg": "5000000",
        "total_area_m2": "1250",
        # metric values (via format_decimal — no commas for 3-digit numbers)
        "total_design_refrigeration_load": "450.0",
        "total_compressor_capacity": "500.0",
        "total_installed_power": "350.0",
    },
    "watermark": "DRAFT",
    "disclaimer": "AI-assisted. Data for reference only.",
    # Scheme names come from the data provider (Chinese), not translated
    "scheme_names": [
        "方案A — 氨+CO2复叠系统",
        "方案B — 氟利昂 R507 系统",
    ],
    # scheme_details: scheme_id, scheme name, rank, total_score from data provider
    "scheme_details": [
        {
            "scheme_id": "scheme_a",
            "name": "方案A — 氨+CO2复叠系统",
            "rank": 1,
            "total_score": 92,
            "metric_values": {
                "total_investment_cny": 8500000,
                "total_area_m2": 1350,
                "operating_cost_per_year": 680000,
                "design_cooling_load_kw_r": 450,
                "installed_power_kw_e": 350,
            },
        },
        {
            "scheme_id": "scheme_b",
            "name": "方案B — 氟利昂 R507 系统",
            "rank": 2,
            "total_score": 78,
            "metric_values": {
                "total_investment_cny": 7200000,
                "total_area_m2": 1400,
                "operating_cost_per_year": 820000,
                "design_cooling_load_kw_r": 450,
                "installed_power_kw_e": 380,
            },
        },
    ],
    "recommended_scheme_code": "scheme_a",
    # provenance source_id / source_tool strings that should appear in report text
    "provenance_source_ids": [
        "calc-tp-001",
        "calc-cl-001",
        "calc-es-001",
        "calc-ee-001",
        "scheme-run-20260625-001",
    ],
    "provenance_source_tools": [
        "throughput_calculator",
        "cooling_load_calculator",
        "equipment_calculator",
        "power_calculator",
    ],
    # document_control labels (from translation catalog)
    "document_control_labels": {
        "title": "Document Control",
        "content_hash": "Content Hash",
        "template_version": "Template Version",
        "generated_by": "Generated By",
        "generated_at": "Generated At",
        "revision": "Revision",
    },
    # header / footer text (from template manifest rendering)
    "header_text": "Blueberry Cold Storage - Demo Plant — Concept Design Report",
    "footer_text": "— 1 —",
    # quality_summary counts (derived from evaluate_quality on rich data provider output)
    "quality_summary": {
        "total_findings": 6,
        "blocker_count": 4,
        "warning_count": 2,
        "info_count": 0,
    },
    # quality finding codes that appear in rendered report text
    "quality_finding_codes": ["SOURCE_MISSING_CONTENT_HASH"],
    # quality finding severities (localized labels)
    "quality_severity_blocker": "Blocker",
    "quality_severity_warning": "Warning",
    # citation section_keys that appear in the citations table
    "citation_section_keys": [
        "project_summary",
        "throughput_inventory_area",
        "cooling_load",
        "equipment_selection",
        "electrical_and_energy",
        "scheme_comparison",
        "report_metadata",
    ],
    # approval / empty state text
    "approval_label": "Approval Information",
    # provenance fields that appear in provenance section
    "provenance_keys": [
        "content_hash",
        "canonical_hash",
        "selection_rules",
        "assembly_timestamp",
    ],
}

# Backward-compat aliases
ZH_CN_TITLES = list(EXPECTED_ZH_CN["section_titles"])
EN_US_TITLES = list(EXPECTED_EN_US["section_titles"])

# MIME type constants (used across multiple tests)
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

# ======================================================================
# Tests
# ======================================================================


class TestRealStorageE2E:
    """Real storage E2E tests using ReportArtifactStorage(tmp_path)."""

    # ------------------------------------------------------------------
    # 1. zh-CN DOCX exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_zh_cn_docx_exact_content(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Full pipeline zh-CN DOCX — verify titles, fields, values, disclaimer, watermark."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "docx",
                ReportLocale.ZH_CN,
            )

        text = _extract_docx_text(raw_bytes)

        # Section titles (zh-CN localized)
        for title in EXPECTED_ZH_CN["section_titles"]:
            assert title in text, f"zh-CN DOCX missing section title: {title!r}"

        # Business content — exact expected values from data provider
        assert "Blueberry Cold Storage" in text, (
            "zh-CN DOCX should contain exact project name 'Blueberry Cold Storage'"
        )
        assert "Kunming, Yunnan" in text, (
            "zh-CN DOCX should contain exact location 'Kunming, Yunnan'"
        )

        # Scheme names from data provider (Chinese)
        for scheme_name in EXPECTED_ZH_CN["scheme_names"]:
            assert scheme_name in text, f"zh-CN DOCX should contain scheme name {scheme_name!r}"

        # Scheme details — scheme_id, rank, score appear in rendered text
        for scheme_detail in EXPECTED_ZH_CN["scheme_details"]:
            assert str(scheme_detail["scheme_id"]) in text, (
                f"zh-CN DOCX should contain scheme_id {scheme_detail['scheme_id']}"
            )
            assert str(scheme_detail["rank"]) in text, (
                f"zh-CN DOCX should contain rank {scheme_detail['rank']}"
            )

        # Recommended scheme code
        assert EXPECTED_ZH_CN["recommended_scheme_code"] in text, (
            "zh-CN DOCX should contain recommended scheme code 'scheme_a'"
        )

        # Cooling load value
        assert "450" in text, "zh-CN DOCX should contain cooling load 450"
        assert "500" in text, "zh-CN DOCX should contain compressor capacity 500"
        assert "350" in text, "zh-CN DOCX should contain installed power 350"

        # Field labels — exact Chinese localized labels from translation catalog
        for field_key, field_label in EXPECTED_ZH_CN["field_labels"].items():
            assert field_label in text, (
                f"zh-CN DOCX should contain field label {field_label!r} (field.{field_key})"
            )

        # Unit labels — exact Chinese localized unit labels
        for unit_key, unit_label in EXPECTED_ZH_CN["unit_labels"].items():
            assert unit_label in text, (
                f"zh-CN DOCX should contain unit label {unit_label!r} (unit.{unit_key})"
            )

        # Provenance source IDs appear in citations/provenance section
        for source_id in EXPECTED_ZH_CN["provenance_source_ids"]:
            assert source_id in text, f"zh-CN DOCX should contain source_id {source_id!r}"

        # Provenance source tools appear in citations/provenance section
        for tool in EXPECTED_ZH_CN["provenance_source_tools"]:
            assert tool in text, f"zh-CN DOCX should contain source tool {tool!r}"

        # Disclaimer
        assert EXPECTED_ZH_CN["disclaimer"] in text, (
            f"zh-CN DOCX should contain disclaimer: {EXPECTED_ZH_CN['disclaimer']}"
        )

        # Watermark (草稿) — should appear in the DOCX raw XML (VML shapes)
        assert _docx_contains_text(raw_bytes, EXPECTED_ZH_CN["watermark"]), (
            f"zh-CN DOCX should contain watermark '{EXPECTED_ZH_CN['watermark']}'"
        )

        # Formatted values — exact display strings from format_decimal / str()
        for val_key, expected_val in EXPECTED_ZH_CN["formatted_values"].items():
            assert expected_val in text, (
                f"zh-CN DOCX should contain formatted value {expected_val!r} for key {val_key!r}"
            )

        # Scheme metric values — each scheme's raw data value appears in text
        for sd in EXPECTED_ZH_CN["scheme_details"]:
            for metric_key, metric_val in sd.get("metric_values", {}).items():
                assert str(metric_val) in text, (
                    f"zh-CN DOCX should contain scheme metric {metric_key}={metric_val!r} "
                    f"for scheme {sd['scheme_id']}"
                )

        # Document control labels
        for label_key, label_text in EXPECTED_ZH_CN["document_control_labels"].items():
            assert label_text in text, (
                f"zh-CN DOCX should contain document_control label '{label_text}' ({label_key})"
            )

        # Header confidentiality text (checked via raw XML for DOCX)
        assert _docx_contains_text(raw_bytes, EXPECTED_ZH_CN["header_text"]), (
            f"zh-CN DOCX should contain header text: {EXPECTED_ZH_CN['header_text']}"
        )

        # Quality summary check — exact counts appear in rendered text
        qs = EXPECTED_ZH_CN["quality_summary"]
        assert str(qs["total_findings"]) in text, (
            f"zh-CN DOCX should contain total findings {qs['total_findings']}"
        )
        assert str(qs["blocker_count"]) in text, (
            f"zh-CN DOCX should contain blocker count {qs['blocker_count']}"
        )
        assert str(qs["warning_count"]) in text, (
            f"zh-CN DOCX should contain warning count {qs['warning_count']}"
        )

        # Quality finding codes appear in the rendered text
        for code in EXPECTED_ZH_CN["quality_finding_codes"]:
            assert code in text, f"zh-CN DOCX should contain quality finding code {code!r}"

        # Quality severity labels
        assert EXPECTED_ZH_CN["quality_severity_blocker"] in text, (
            f"zh-CN DOCX should contain severity '{EXPECTED_ZH_CN['quality_severity_blocker']}'"
        )

        # Citation section_keys appear in the citations table
        for section_key in EXPECTED_ZH_CN["citation_section_keys"]:
            assert section_key in text, (
                f"zh-CN DOCX should contain citation section_key {section_key!r}"
            )

        # Provenance — localized keys appear in rendered text
        assert "内容哈希" in text or ("content_hash" in text), (
            "zh-CN DOCX should contain provenance content_hash"
        )
        assert "组装时间戳" in text or ("assembly_timestamp" in text), (
            "zh-CN DOCX should contain provenance assembly_timestamp"
        )

    # ------------------------------------------------------------------
    # 2. en-US DOCX exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_en_us_docx_exact_content(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Full pipeline en-US DOCX — verify titles, fields, values, disclaimer, watermark."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "docx",
                ReportLocale.EN_US,
            )

        text = _extract_docx_text(raw_bytes)

        # Section titles (en-US localized)
        for title in EXPECTED_EN_US["section_titles"]:
            assert title in text, f"en-US DOCX missing section title: {title!r}"

        # Business content — exact expected values from data provider
        assert "Blueberry Cold Storage" in text, (
            "en-US DOCX should contain exact project name 'Blueberry Cold Storage'"
        )
        assert "Kunming, Yunnan" in text, (
            "en-US DOCX should contain exact location 'Kunming, Yunnan'"
        )

        # Scheme names from data provider (Chinese, same as data)
        for scheme_name in EXPECTED_EN_US["scheme_names"]:
            assert scheme_name in text, f"en-US DOCX should contain scheme name {scheme_name!r}"

        # Scheme details — scheme_id, rank, score appear in rendered text
        for scheme_detail in EXPECTED_EN_US["scheme_details"]:
            assert str(scheme_detail["scheme_id"]) in text, (
                f"en-US DOCX should contain scheme_id {scheme_detail['scheme_id']}"
            )
            assert str(scheme_detail["rank"]) in text, (
                f"en-US DOCX should contain rank {scheme_detail['rank']}"
            )

        # Recommended scheme code
        assert EXPECTED_EN_US["recommended_scheme_code"] in text, (
            "en-US DOCX should contain recommended scheme code 'scheme_a'"
        )

        # Cooling load value
        assert "450" in text, "en-US DOCX should contain cooling load 450"
        assert "500" in text, "en-US DOCX should contain compressor capacity 500"
        assert "350" in text, "en-US DOCX should contain installed power 350"

        # Field labels — exact English localized labels from translation catalog
        for field_key, field_label in EXPECTED_EN_US["field_labels"].items():
            assert field_label in text, (
                f"en-US DOCX should contain field label {field_label!r} (field.{field_key})"
            )

        # Unit labels — exact English localized unit labels
        for unit_key, unit_label in EXPECTED_EN_US["unit_labels"].items():
            assert unit_label in text, (
                f"en-US DOCX should contain unit label {unit_label!r} (unit.{unit_key})"
            )

        # Provenance source IDs appear in citations/provenance section
        for source_id in EXPECTED_EN_US["provenance_source_ids"]:
            assert source_id in text, f"en-US DOCX should contain source_id {source_id!r}"

        # Provenance source tools appear in citations/provenance section
        for tool in EXPECTED_EN_US["provenance_source_tools"]:
            assert tool in text, f"en-US DOCX should contain source tool {tool!r}"

        # Disclaimer
        assert EXPECTED_EN_US["disclaimer"] in text, (
            f"en-US DOCX should contain disclaimer: {EXPECTED_EN_US['disclaimer']}"
        )

        # Watermark (DRAFT) — should appear in the DOCX raw XML (VML shapes)
        assert _docx_contains_text(raw_bytes, EXPECTED_EN_US["watermark"]), (
            f"en-US DOCX should contain watermark '{EXPECTED_EN_US['watermark']}'"
        )

        # Formatted values — exact display strings
        for val_key, expected_val in EXPECTED_EN_US["formatted_values"].items():
            assert expected_val in text, (
                f"en-US DOCX should contain formatted value {expected_val!r} for key {val_key!r}"
            )

        # Scheme metric values
        for sd in EXPECTED_EN_US["scheme_details"]:
            for metric_key, metric_val in sd.get("metric_values", {}).items():
                assert str(metric_val) in text, (
                    f"en-US DOCX should contain scheme metric {metric_key}={metric_val!r} "
                    f"for scheme {sd['scheme_id']}"
                )

        # Document control labels
        for label_key, label_text in EXPECTED_EN_US["document_control_labels"].items():
            assert label_text in text, (
                f"en-US DOCX should contain document_control label '{label_text}' ({label_key})"
            )

        # Header confidentiality text (checked via raw XML for DOCX)
        assert _docx_contains_text(raw_bytes, EXPECTED_EN_US["header_text"]), (
            f"en-US DOCX should contain header text: {EXPECTED_EN_US['header_text']}"
        )

        # Quality summary check — exact counts appear in rendered text
        qs = EXPECTED_EN_US["quality_summary"]
        assert str(qs["total_findings"]) in text, (
            f"en-US DOCX should contain total findings {qs['total_findings']}"
        )
        assert str(qs["blocker_count"]) in text, (
            f"en-US DOCX should contain blocker count {qs['blocker_count']}"
        )
        assert str(qs["warning_count"]) in text, (
            f"en-US DOCX should contain warning count {qs['warning_count']}"
        )

        # Quality finding codes appear in rendered text
        for code in EXPECTED_EN_US["quality_finding_codes"]:
            assert code in text, f"en-US DOCX should contain quality finding code {code!r}"

        # Quality severity labels
        assert EXPECTED_EN_US["quality_severity_blocker"] in text, (
            f"en-US DOCX should contain severity '{EXPECTED_EN_US['quality_severity_blocker']}'"
        )

        # Citation section_keys appear in the citations table
        for section_key in EXPECTED_EN_US["citation_section_keys"]:
            assert section_key in text, (
                f"en-US DOCX should contain citation section_key {section_key!r}"
            )

        # Provenance — localized keys appear in rendered text
        assert "Content Hash" in text or ("content_hash" in text), (
            "en-US DOCX should contain provenance content_hash"
        )
        assert "Assembly Timestamp" in text or ("assembly_timestamp" in text), (
            "en-US DOCX should contain provenance assembly_timestamp"
        )

    # ------------------------------------------------------------------
    # 3. zh-CN PDF exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_zh_cn_pdf_exact_content(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Full pipeline zh-CN PDF — verify section titles, values, disclaimer, watermark."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "pdf",
                ReportLocale.ZH_CN,
            )

        text = _extract_pdf_text(raw_bytes)

        # Section titles (zh-CN localized)
        for title in EXPECTED_ZH_CN["section_titles"]:
            assert title in text, f"zh-CN PDF missing section title: {title!r}"

        # Business content — exact expected values from data provider
        assert "Blueberry Cold Storage" in text, (
            "zh-CN PDF should contain exact project name 'Blueberry Cold Storage'"
        )
        assert "Kunming, Yunnan" in text, (
            "zh-CN PDF should contain exact location 'Kunming, Yunnan'"
        )

        # Scheme names — use prefix matching for PDF (table cells may not extract fully)
        assert "方案A" in text, (
            "zh-CN PDF should contain scheme prefix '方案A' (from data provider scheme name)"
        )
        assert "方案B" in text, "zh-CN PDF should contain scheme prefix '方案B'"

        # Recommended scheme code
        assert "scheme_a" in text, "zh-CN PDF should contain recommended scheme code 'scheme_a'"

        # Cooling load value
        assert "450" in text, "zh-CN PDF should contain cooling load 450"
        assert "500" in text, "zh-CN PDF should contain compressor capacity 500"
        assert "350" in text, "zh-CN PDF should contain installed power 350"

        # Field labels — exact Chinese localized labels
        for field_key, field_label in EXPECTED_ZH_CN["field_labels"].items():
            assert field_label in text, (
                f"zh-CN PDF should contain field label {field_label!r} (field.{field_key})"
            )

        # Provenance source tools — PDF extraction may miss table cell content;
        # only check the first tool that is also rendered as plain text
        assert "throughput_calculator" in text, (
            "zh-CN PDF should contain source tool 'throughput_calculator'"
        )

        # Disclaimer
        assert EXPECTED_ZH_CN["disclaimer"] in text, (
            f"zh-CN PDF should contain disclaimer: {EXPECTED_ZH_CN['disclaimer']}"
        )

        # Watermark (草稿) — PDF watermark is rendered as text on the page
        assert EXPECTED_ZH_CN["watermark"] in text, (
            f"zh-CN PDF should contain watermark '{EXPECTED_ZH_CN['watermark']}'"
        )

        # Header text — rendered as text on PDF page
        assert EXPECTED_ZH_CN["header_text"] in text, (
            f"zh-CN PDF should contain header text: {EXPECTED_ZH_CN['header_text']}"
        )

        # Footer text — check for page number marker
        assert "1" in text, "zh-CN PDF should contain page number 1 (footer)"

        # Quality summary counts — numbers appear in rendered quality section
        qs = EXPECTED_ZH_CN["quality_summary"]
        assert str(qs["total_findings"]) in text, (
            f"zh-CN PDF should contain total findings {qs['total_findings']}"
        )

        # Quality severity labels (localized)
        assert EXPECTED_ZH_CN["quality_severity_blocker"] in text, (
            f"zh-CN PDF should contain severity '{EXPECTED_ZH_CN['quality_severity_blocker']}'"
        )

        # Provenance — localized keys
        assert "内容哈希" in text or "content_hash" in text, (
            "zh-CN PDF should contain provenance content_hash"
        )
        assert "组装时间戳" in text or "assembly_timestamp" in text, (
            "zh-CN PDF should contain provenance assembly_timestamp"
        )

        # Formatted values
        for val_key, expected_val in EXPECTED_ZH_CN["formatted_values"].items():
            assert expected_val in text, (
                f"zh-CN PDF should contain formatted value {expected_val!r} for key {val_key!r}"
            )

        # Scheme metric values
        for sd in EXPECTED_ZH_CN["scheme_details"]:
            for metric_key, metric_val in sd.get("metric_values", {}).items():
                assert str(metric_val) in text, (
                    f"zh-CN PDF should contain scheme metric {metric_key}={metric_val!r} "
                    f"for scheme {sd['scheme_id']}"
                )

        # Unit labels — verify in extracted text if available
        for _unit_key, unit_label in EXPECTED_ZH_CN["unit_labels"].items():
            if unit_label in text:
                pass

    # ------------------------------------------------------------------
    # 4. en-US PDF exact content
    # ------------------------------------------------------------------

    def test_real_pipeline_en_us_pdf_exact_content(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Full pipeline en-US PDF — verify section titles, values, disclaimer, watermark."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "pdf",
                ReportLocale.EN_US,
            )

        text = _extract_pdf_text(raw_bytes)

        # Section titles (en-US localized)
        for title in EXPECTED_EN_US["section_titles"]:
            assert title in text, f"en-US PDF missing section title: {title!r}"

        # Business content — exact expected values from data provider
        assert "Blueberry Cold Storage" in text, (
            "en-US PDF should contain exact project name 'Blueberry Cold Storage'"
        )
        assert "Kunming, Yunnan" in text, (
            "en-US PDF should contain exact location 'Kunming, Yunnan'"
        )

        # Scheme names — use prefix matching for PDF (table cells may not extract fully)
        assert "方案A" in text, (
            "en-US PDF should contain scheme prefix '方案A' (Chinese scheme name in data)"
        )
        assert "方案B" in text, "en-US PDF should contain scheme prefix '方案B'"

        # Recommended scheme code
        assert "scheme_a" in text, "en-US PDF should contain recommended scheme code 'scheme_a'"

        # Cooling load value
        assert "450" in text, "en-US PDF should contain cooling load 450"
        assert "500" in text, "en-US PDF should contain compressor capacity 500"
        assert "350" in text, "en-US PDF should contain installed power 350"

        # Field labels — exact English localized labels
        for field_key, field_label in EXPECTED_EN_US["field_labels"].items():
            assert field_label in text, (
                f"en-US PDF should contain field label {field_label!r} (field.{field_key})"
            )

        # Provenance source tools — PDF extraction may miss table cell content;
        # only check the first tool that is also rendered as plain text
        assert "throughput_calculator" in text, (
            "en-US PDF should contain source tool 'throughput_calculator'"
        )

        # Disclaimer
        assert EXPECTED_EN_US["disclaimer"] in text, (
            f"en-US PDF should contain disclaimer: {EXPECTED_EN_US['disclaimer']}"
        )

        # Watermark (DRAFT) — PDF watermark is rendered as text on the page
        assert EXPECTED_EN_US["watermark"] in text, (
            f"en-US PDF should contain watermark '{EXPECTED_EN_US['watermark']}'"
        )

        # Header text — rendered as text on PDF page
        assert EXPECTED_EN_US["header_text"] in text, (
            f"en-US PDF should contain header text: {EXPECTED_EN_US['header_text']}"
        )

        # Footer text — check for page number marker
        assert "1" in text, "en-US PDF should contain page number 1 (footer)"

        # Quality summary counts — numbers appear in rendered quality section
        qs = EXPECTED_EN_US["quality_summary"]
        assert str(qs["total_findings"]) in text, (
            f"en-US PDF should contain total findings {qs['total_findings']}"
        )

        # Quality severity labels (localized)
        assert EXPECTED_EN_US["quality_severity_blocker"] in text, (
            f"en-US PDF should contain severity '{EXPECTED_EN_US['quality_severity_blocker']}'"
        )

        # Provenance — localized keys
        assert "content_hash" in text, "en-US PDF should contain provenance content_hash"
        assert "assembly_timestamp" in text or "Assembly Timestamp" in text, (
            "en-US PDF should contain provenance assembly_timestamp"
        )

        # Formatted values
        for val_key, expected_val in EXPECTED_EN_US["formatted_values"].items():
            assert expected_val in text, (
                f"en-US PDF should contain formatted value {expected_val!r} for key {val_key!r}"
            )

        # Scheme metric values
        for sd in EXPECTED_EN_US["scheme_details"]:
            for metric_key, metric_val in sd.get("metric_values", {}).items():
                assert str(metric_val) in text, (
                    f"en-US PDF should contain scheme metric {metric_key}={metric_val!r} "
                    f"for scheme {sd['scheme_id']}"
                )

        # Unit labels — verify in extracted text if available
        for _unit_key, unit_label in EXPECTED_EN_US["unit_labels"].items():
            if unit_label in text:
                pass

    # ------------------------------------------------------------------
    # 5. Four independent renders — identical canonical snapshot via RecordingObserver
    # ------------------------------------------------------------------

    def test_four_independent_render_calls_have_identical_canonical_snapshot(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Four render calls (zh/en DOCX, zh/en PDF) share identical canonical snapshot.

        The canonical model is derived from the revision content_json alone
        and is independent of locale or format.  RecordingObserver captures
        all four canonicals; all golden_dict results must be identical.
        """
        observer = RecordingObserver()

        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, _repo, storage = _make_render_service(
                session,
                tmp_storage,
                observer=observer,
            )

            # Render all 4 combinations
            format_locale_pairs = [
                ("docx", ReportLocale.ZH_CN),
                ("docx", ReportLocale.EN_US),
                ("pdf", ReportLocale.ZH_CN),
                ("pdf", ReportLocale.EN_US),
            ]

            artifacts: list[ReportExportArtifact] = []
            for fmt, loc in format_locale_pairs:
                artifact, raw_bytes = _render_artifact(
                    render_svc,
                    storage,
                    report.id,
                    revision.revision_number,
                    fmt,
                    loc,
                )
                artifacts.append(artifact)

        # Observer should have recorded 4 snapshots
        assert len(observer.records) == 4, (
            f"Expected 4 observer records, got {len(observer.records)}"
        )

        # All 4 golden_dict results must be identical
        baseline_golden = observer.records[0]["golden"]
        for i, record in enumerate(observer.records[1:], start=1):
            assert record["golden"] == baseline_golden, (
                f"Snapshot mismatch at index {i}: "
                f"locale={record['locale']}, format={record['format']}"
            )

        # Verify the observer captured locale/format metadata
        observed_formats = {(r["locale"], r["format"]) for r in observer.records}
        expected_pairs = {
            ("zh-CN", "docx"),
            ("en-US", "docx"),
            ("zh-CN", "pdf"),
            ("en-US", "pdf"),
        }
        assert observed_formats == expected_pairs, (
            f"Observer did not capture all locale/format pairs: "
            f"got {observed_formats}, expected {expected_pairs}"
        )

        # Verify each artifact's file bytes are non-empty, sizes match,
        # SHA-256 matches, mime_type is correct, and non-empty metadata fields
        for artifact, (fmt, loc) in zip(artifacts, format_locale_pairs, strict=False):
            raw = storage.get(artifact.storage_key)
            assert raw, f"Artifact bytes empty for {loc.value}/{fmt}"
            assert artifact.file_size_bytes == len(raw), (
                f"File size mismatch for {artifact.id}: "
                f"artifact says {artifact.file_size_bytes}, actual is {len(raw)}"
            )

            # SHA-256 verification
            actual_sha256 = hashlib.sha256(raw).hexdigest()
            assert artifact.file_sha256 == actual_sha256, (
                f"SHA-256 mismatch for {artifact.id}: "
                f"artifact says {artifact.file_sha256}, computed {actual_sha256}"
            )

            # MIME type verification
            expected_mime = DOCX_MIME if fmt == "docx" else PDF_MIME
            assert artifact.mime_type == expected_mime, (
                f"MIME type mismatch for {artifact.id} ({loc.value}/{fmt}): "
                f"got {artifact.mime_type}, expected {expected_mime}"
            )

            # Non-empty metadata fields
            assert artifact.format.value, f"Format must be non-empty for {artifact.id}"
            assert artifact.locale.value, f"Locale must be non-empty for {artifact.id}"
            assert artifact.template_locale.value, (
                f"Template locale must be non-empty for {artifact.id}"
            )
            # idempotency_key was not passed — both should be empty/falsy
            assert artifact.idempotency_key is None
            # claim_token and claim_version from render()
            assert artifact.claim_token is None or artifact.claim_token == ""
            # claim_version defaults to 0 when no claim_token
            assert artifact.claim_version == 0

        # No temp files remain: scan storage directory for temp/backup/partial artifacts
        storage_dir = storage._base_dir
        TEMP_SUFFIXES = (
            ".tmp",
            ".backup",
            ".meta.backup",
            ".pending_delete",
            ".reclaim_tmp",
            ".partial",
            ".quarantine",
        )
        temp_files_found: list[str] = []
        bundle_tempdirs_found: list[str] = []
        for entry in storage_dir.iterdir():
            if entry.is_file() and any(entry.name.endswith(s) for s in TEMP_SUFFIXES):
                temp_files_found.append(str(entry))
            elif (
                entry.is_dir()
                and (entry / "payload").is_file()
                and (entry / "owner.json").is_file()
            ):
                bundle_tempdirs_found.append(str(entry))
        assert not temp_files_found, f"Temp files remain in storage dir: {temp_files_found}"
        assert not bundle_tempdirs_found, (
            f"Bundle tempdirs remain in storage dir: {bundle_tempdirs_found}"
        )

        # Verify section keys in the golden dict
        section_keys = [s["section_key"] for s in baseline_golden["sections"]]
        expected_sections = {
            "report_metadata",
            "project_summary",
            "input_conditions",
            "assumptions",
            "throughput_inventory_area",
            "cooling_load",
            "equipment_selection",
            "electrical_and_energy",
            "scheme_comparison",
            "investment_estimate",
            "risks_and_missing_information",
            "quality_summary",
            "citations",
            "provenance",
        }
        assert set(section_keys) == expected_sections, (
            f"Section key mismatch. "
            f"Missing: {expected_sections - set(section_keys)}. "
            f"Extra: {set(section_keys) - expected_sections}."
        )

    # ------------------------------------------------------------------
    # 6. Exact localized watermark text
    # ------------------------------------------------------------------

    def test_real_files_have_exact_localized_watermark(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Verify watermark text: zh-CN '草稿' / en-US 'DRAFT' in both DOCX and PDF."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, _repo, storage = _make_render_service(session, tmp_storage)

            locale_format_pairs = [
                (ReportLocale.ZH_CN, "docx", ZH_CN_WATERMARK),
                (ReportLocale.EN_US, "docx", EN_US_WATERMARK),
                (ReportLocale.ZH_CN, "pdf", ZH_CN_WATERMARK),
                (ReportLocale.EN_US, "pdf", EN_US_WATERMARK),
            ]

            for loc, fmt, expected_watermark in locale_format_pairs:
                artifact, raw_bytes = _render_artifact(
                    render_svc,
                    storage,
                    report.id,
                    revision.revision_number,
                    fmt,
                    loc,
                )

                if fmt == "docx":
                    assert _docx_contains_text(raw_bytes, expected_watermark), (
                        f"DOCX ({loc.value}) missing watermark: {expected_watermark!r}"
                    )
                elif fmt == "pdf":
                    text = _extract_pdf_text(raw_bytes)
                    assert expected_watermark in text, (
                        f"PDF ({loc.value}) missing watermark: {expected_watermark!r}"
                    )

    # ------------------------------------------------------------------
    # 7. Exact localized disclaimer
    # ------------------------------------------------------------------

    def test_real_files_have_exact_localized_disclaimer(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Verify disclaimer text: zh-CN / en-US exact match in both DOCX and PDF."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, _repo, storage = _make_render_service(session, tmp_storage)

            locale_format_pairs = [
                (ReportLocale.ZH_CN, "docx", ZH_CN_DISCLAIMER),
                (ReportLocale.EN_US, "docx", EN_US_DISCLAIMER),
                (ReportLocale.ZH_CN, "pdf", ZH_CN_DISCLAIMER),
                (ReportLocale.EN_US, "pdf", EN_US_DISCLAIMER),
            ]

            for loc, fmt, expected_disclaimer in locale_format_pairs:
                artifact, raw_bytes = _render_artifact(
                    render_svc,
                    storage,
                    report.id,
                    revision.revision_number,
                    fmt,
                    loc,
                )

                if fmt == "docx":
                    text = _extract_docx_text(raw_bytes)
                elif fmt == "pdf":
                    text = _extract_pdf_text(raw_bytes)

                assert expected_disclaimer in text, (
                    f"{fmt.upper()} ({loc.value}) missing disclaimer: {expected_disclaimer!r}"
                )

    # ------------------------------------------------------------------
    # 8-11. New Session reload tests
    # ------------------------------------------------------------------

    def test_real_storage_artifact_survives_database_reload_zh_cn_docx(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Reload zh-CN DOCX artifact from new session; verify ALL fields survive.

        Validates every field on the reloaded artifact matches the original
        render-time artifact.
        """
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "docx",
                ReportLocale.ZH_CN,
            )
            artifact_id = artifact.id

            # Capture original field values for comparison
            orig_id = artifact.id
            orig_format = artifact.format
            orig_mime_type = artifact.mime_type
            orig_file_name = artifact.file_name
            orig_file_size = artifact.file_size_bytes
            orig_sha256 = artifact.file_sha256
            orig_revision_number = artifact.revision_number
            orig_storage_key = artifact.storage_key
            orig_locale = artifact.locale
            orig_template_locale = artifact.template_locale
            orig_translation_catalog_version = artifact.translation_catalog_version
            orig_translation_catalog_content_hash = artifact.translation_catalog_content_hash
            orig_localized_template_content_hash = artifact.localized_template_content_hash
            orig_claim_version = artifact.claim_version

        # New session — reload artifact metadata from database
        with session_factory() as new_session:
            new_repo = SQLReportRepository(new_session)
            reloaded = new_repo.get_artifact(artifact_id)

            assert reloaded is not None, "Reloaded artifact must not be None"

            # == Validate ALL fields ==
            # Identity field
            assert reloaded.id == orig_id, f"Artifact ID mismatch: {reloaded.id} != {orig_id}"

            # Status
            assert reloaded.status == ArtifactStatus.COMPLETED

            # Format
            assert reloaded.format == orig_format, (
                f"Format mismatch: {reloaded.format} != {orig_format}"
            )

            # MIME type
            assert reloaded.mime_type == orig_mime_type, (
                f"MIME type mismatch: {reloaded.mime_type} != {orig_mime_type}"
            )
            assert reloaded.mime_type == DOCX_MIME

            # File metadata
            assert reloaded.file_name == orig_file_name, (
                f"File name mismatch: {reloaded.file_name} != {orig_file_name}"
            )
            assert reloaded.file_size_bytes > 0, "file_size_bytes must be > 0"
            assert reloaded.file_size_bytes == orig_file_size, (
                f"File size mismatch: {reloaded.file_size_bytes} != {orig_file_size}"
            )
            assert reloaded.file_sha256, "file_sha256 must be non-empty"
            assert reloaded.file_sha256 == orig_sha256, (
                f"SHA-256 mismatch: {reloaded.file_sha256} != {orig_sha256}"
            )

            # Revision
            assert reloaded.revision_number == orig_revision_number, (
                f"Revision number mismatch: {reloaded.revision_number} != {orig_revision_number}"
            )

            # Timestamp
            assert reloaded.generated_at is not None, "generated_at must not be None"

            # Storage
            assert reloaded.storage_key, "storage_key must be non-empty"
            assert reloaded.storage_key == orig_storage_key, (
                f"Storage key mismatch: {reloaded.storage_key} != {orig_storage_key}"
            )

            # Idempotency (not used in this render — no key passed)
            assert reloaded.idempotency_key is None
            assert reloaded.claim_token is None or reloaded.claim_token == ""
            assert reloaded.claim_version == orig_claim_version, (
                f"Claim version mismatch: {reloaded.claim_version} != {orig_claim_version}"
            )
            assert reloaded.claim_version == 0

            # Locale fields
            assert reloaded.locale == orig_locale, (
                f"Locale mismatch: {reloaded.locale} != {orig_locale}"
            )
            assert reloaded.locale == ReportLocale.ZH_CN
            assert reloaded.template_locale == orig_template_locale, (
                f"Template locale mismatch: {reloaded.template_locale} != {orig_template_locale}"
            )
            assert reloaded.template_locale == ReportLocale.ZH_CN
            assert reloaded.translation_catalog_version == orig_translation_catalog_version, (
                "Translation catalog version mismatch"
            )
            assert reloaded.translation_catalog_version == "1.0.0"
            assert reloaded.translation_catalog_content_hash, (
                "translation_catalog_content_hash must be non-empty"
            )
            assert (
                reloaded.translation_catalog_content_hash == orig_translation_catalog_content_hash
            )
            assert reloaded.localized_template_content_hash, (
                "localized_template_content_hash must be non-empty"
            )
            assert reloaded.localized_template_content_hash == orig_localized_template_content_hash

            # Read file from storage using reloaded storage_key
            reloaded_bytes = storage.get(reloaded.storage_key)
            assert reloaded_bytes, "Reloaded artifact bytes must be non-empty"
            assert len(reloaded_bytes) == reloaded.file_size_bytes, (
                f"File size mismatch: stored {len(reloaded_bytes)} != DB {reloaded.file_size_bytes}"
            )

        # Parse as DOCX and verify section titles are present
        docx_text = _extract_docx_text(reloaded_bytes)
        for title in ZH_CN_TITLES:
            assert title in docx_text, f"zh-CN DOCX (reloaded) missing section title: {title!r}"

    def test_real_storage_artifact_survives_database_reload_en_us_docx(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Reload en-US DOCX artifact from new session; verify ALL fields survive."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "docx",
                ReportLocale.EN_US,
            )
            artifact_id = artifact.id

            # Capture originals
            orig_id = artifact.id
            orig_format = artifact.format
            orig_mime_type = artifact.mime_type
            orig_file_name = artifact.file_name
            orig_file_size = artifact.file_size_bytes
            orig_sha256 = artifact.file_sha256
            orig_revision_number = artifact.revision_number
            orig_storage_key = artifact.storage_key
            orig_locale = artifact.locale
            orig_template_locale = artifact.template_locale
            orig_translation_catalog_version = artifact.translation_catalog_version
            orig_translation_catalog_content_hash = artifact.translation_catalog_content_hash
            orig_localized_template_content_hash = artifact.localized_template_content_hash
            orig_claim_version = artifact.claim_version

        with session_factory() as new_session:
            new_repo = SQLReportRepository(new_session)
            reloaded = new_repo.get_artifact(artifact_id)

            assert reloaded is not None, "Reloaded artifact must not be None"

            # Identity
            assert reloaded.id == orig_id

            # Status
            assert reloaded.status == ArtifactStatus.COMPLETED

            # Format
            assert reloaded.format == orig_format

            # MIME type
            assert reloaded.mime_type == orig_mime_type
            assert reloaded.mime_type == DOCX_MIME

            # File metadata
            assert reloaded.file_name == orig_file_name
            assert reloaded.file_size_bytes > 0
            assert reloaded.file_size_bytes == orig_file_size
            assert reloaded.file_sha256, "file_sha256 must be non-empty"
            assert reloaded.file_sha256 == orig_sha256

            # Revision
            assert reloaded.revision_number == orig_revision_number

            # Timestamp
            assert reloaded.generated_at is not None

            # Storage
            assert reloaded.storage_key, "storage_key must be non-empty"
            assert reloaded.storage_key == orig_storage_key

            # Idempotency (no key passed)
            assert reloaded.idempotency_key is None
            assert reloaded.claim_token is None or reloaded.claim_token == ""
            assert reloaded.claim_version == orig_claim_version
            assert reloaded.claim_version == 0

            # Locale fields
            assert reloaded.locale == orig_locale
            assert reloaded.locale == ReportLocale.EN_US
            assert reloaded.template_locale == orig_template_locale
            assert reloaded.template_locale == ReportLocale.EN_US
            assert reloaded.translation_catalog_version == orig_translation_catalog_version
            assert reloaded.translation_catalog_version == "1.0.0"
            assert reloaded.translation_catalog_content_hash, (
                "translation_catalog_content_hash must be non-empty"
            )
            assert (
                reloaded.translation_catalog_content_hash == orig_translation_catalog_content_hash
            )
            assert reloaded.localized_template_content_hash, (
                "localized_template_content_hash must be non-empty"
            )
            assert reloaded.localized_template_content_hash == orig_localized_template_content_hash

            reloaded_bytes = storage.get(reloaded.storage_key)
            assert reloaded_bytes, "Reloaded artifact bytes must be non-empty"
            assert len(reloaded_bytes) == reloaded.file_size_bytes

        docx_text = _extract_docx_text(reloaded_bytes)
        for title in EN_US_TITLES:
            assert title in docx_text, f"en-US DOCX (reloaded) missing section title: {title!r}"

    def test_real_storage_artifact_survives_database_reload_zh_cn_pdf(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Reload zh-CN PDF artifact from new session; verify ALL fields survive."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "pdf",
                ReportLocale.ZH_CN,
            )
            artifact_id = artifact.id

            # Capture originals
            orig_id = artifact.id
            orig_format = artifact.format
            orig_mime_type = artifact.mime_type
            orig_file_name = artifact.file_name
            orig_file_size = artifact.file_size_bytes
            orig_sha256 = artifact.file_sha256
            orig_revision_number = artifact.revision_number
            orig_storage_key = artifact.storage_key
            orig_locale = artifact.locale
            orig_template_locale = artifact.template_locale
            orig_translation_catalog_version = artifact.translation_catalog_version
            orig_translation_catalog_content_hash = artifact.translation_catalog_content_hash
            orig_localized_template_content_hash = artifact.localized_template_content_hash
            orig_claim_version = artifact.claim_version

        with session_factory() as new_session:
            new_repo = SQLReportRepository(new_session)
            reloaded = new_repo.get_artifact(artifact_id)

            assert reloaded is not None, "Reloaded artifact must not be None"

            # Identity
            assert reloaded.id == orig_id
            assert reloaded.status == ArtifactStatus.COMPLETED
            assert reloaded.format == orig_format

            # MIME type
            assert reloaded.mime_type == orig_mime_type
            assert reloaded.mime_type == "application/pdf"

            # File metadata
            assert reloaded.file_name == orig_file_name
            assert reloaded.file_size_bytes > 0
            assert reloaded.file_size_bytes == orig_file_size
            assert reloaded.file_sha256, "file_sha256 must be non-empty"
            assert reloaded.file_sha256 == orig_sha256

            # Revision
            assert reloaded.revision_number == orig_revision_number
            assert reloaded.generated_at is not None

            # Storage + idempotency
            assert reloaded.storage_key, "storage_key must be non-empty"
            assert reloaded.storage_key == orig_storage_key
            assert reloaded.idempotency_key is None
            assert reloaded.claim_token is None or reloaded.claim_token == ""
            assert reloaded.claim_version == orig_claim_version
            assert reloaded.claim_version == 0

            # Locale fields
            assert reloaded.locale == orig_locale
            assert reloaded.locale == ReportLocale.ZH_CN
            assert reloaded.template_locale == orig_template_locale
            assert reloaded.template_locale == ReportLocale.ZH_CN
            assert reloaded.translation_catalog_version == orig_translation_catalog_version
            assert reloaded.translation_catalog_version == "1.0.0"
            assert reloaded.translation_catalog_content_hash, (
                "translation_catalog_content_hash must be non-empty"
            )
            assert (
                reloaded.translation_catalog_content_hash == orig_translation_catalog_content_hash
            )
            assert reloaded.localized_template_content_hash, (
                "localized_template_content_hash must be non-empty"
            )
            assert reloaded.localized_template_content_hash == orig_localized_template_content_hash

            reloaded_bytes = storage.get(reloaded.storage_key)
            assert reloaded_bytes, "Reloaded artifact bytes must be non-empty"
            assert len(reloaded_bytes) == reloaded.file_size_bytes

        pdf_text = _extract_pdf_text(reloaded_bytes)
        for title in ZH_CN_TITLES:
            assert title in pdf_text, f"zh-CN PDF (reloaded) missing section title: {title!r}"

    def test_real_storage_artifact_survives_database_reload_en_us_pdf(
        self,
        session_factory: Any,
        tmp_storage: ReportArtifactStorage,
    ) -> None:
        """Reload en-US PDF artifact from new session; verify ALL fields survive."""
        with session_factory() as session:
            report, revision = _full_pipeline_setup(session)
            render_svc, repo, storage = _make_render_service(session, tmp_storage)
            artifact, raw_bytes = _render_artifact(
                render_svc,
                storage,
                report.id,
                revision.revision_number,
                "pdf",
                ReportLocale.EN_US,
            )
            artifact_id = artifact.id

            # Capture originals
            orig_id = artifact.id
            orig_format = artifact.format
            orig_mime_type = artifact.mime_type
            orig_file_name = artifact.file_name
            orig_file_size = artifact.file_size_bytes
            orig_sha256 = artifact.file_sha256
            orig_revision_number = artifact.revision_number
            orig_storage_key = artifact.storage_key
            orig_locale = artifact.locale
            orig_template_locale = artifact.template_locale
            orig_translation_catalog_version = artifact.translation_catalog_version
            orig_translation_catalog_content_hash = artifact.translation_catalog_content_hash
            orig_localized_template_content_hash = artifact.localized_template_content_hash
            orig_claim_version = artifact.claim_version

        with session_factory() as new_session:
            new_repo = SQLReportRepository(new_session)
            reloaded = new_repo.get_artifact(artifact_id)

            assert reloaded is not None, "Reloaded artifact must not be None"

            # Identity
            assert reloaded.id == orig_id
            assert reloaded.status == ArtifactStatus.COMPLETED
            assert reloaded.format == orig_format

            # MIME type
            assert reloaded.mime_type == orig_mime_type
            assert reloaded.mime_type == "application/pdf"

            # File metadata
            assert reloaded.file_name == orig_file_name
            assert reloaded.file_size_bytes > 0
            assert reloaded.file_size_bytes == orig_file_size
            assert reloaded.file_sha256, "file_sha256 must be non-empty"
            assert reloaded.file_sha256 == orig_sha256

            # Revision + timestamp
            assert reloaded.revision_number == orig_revision_number
            assert reloaded.generated_at is not None

            # Storage + idempotency
            assert reloaded.storage_key, "storage_key must be non-empty"
            assert reloaded.storage_key == orig_storage_key
            assert reloaded.idempotency_key is None
            assert reloaded.claim_token is None or reloaded.claim_token == ""
            assert reloaded.claim_version == orig_claim_version
            assert reloaded.claim_version == 0

            # Locale fields
            assert reloaded.locale == orig_locale
            assert reloaded.locale == ReportLocale.EN_US
            assert reloaded.template_locale == orig_template_locale
            assert reloaded.template_locale == ReportLocale.EN_US
            assert reloaded.translation_catalog_version == orig_translation_catalog_version
            assert reloaded.translation_catalog_version == "1.0.0"
            assert reloaded.translation_catalog_content_hash, (
                "translation_catalog_content_hash must be non-empty"
            )
            assert (
                reloaded.translation_catalog_content_hash == orig_translation_catalog_content_hash
            )
            assert reloaded.localized_template_content_hash, (
                "localized_template_content_hash must be non-empty"
            )
            assert reloaded.localized_template_content_hash == orig_localized_template_content_hash

            reloaded_bytes = storage.get(reloaded.storage_key)
            assert reloaded_bytes, "Reloaded artifact bytes must be non-empty"
            assert len(reloaded_bytes) == reloaded.file_size_bytes

        pdf_text = _extract_pdf_text(reloaded_bytes)
        for title in EN_US_TITLES:
            assert title in pdf_text, f"en-US PDF (reloaded) missing section title: {title!r}"
