"""Multilingual report pilot verifier for the frozen TASK-011 Slice 1 contract."""

from __future__ import annotations

import hashlib
import io
import re
from collections.abc import Callable, Mapping
from dataclasses import dataclass
from datetime import UTC, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.oxml.ns import qn as _docx_qn

from cold_storage.evaluation.artifact_io import (
    assert_no_managed_artifacts,
    atomic_write_bytes,
    atomic_write_json,
)
from cold_storage.evaluation.errors import EvaluationRunnerError
from cold_storage.modules.reports.application.canonical_render_model_builder import (
    build_canonical_render_model,
)
from cold_storage.modules.reports.application.render_model_localizer import (
    localize_render_model,
)
from cold_storage.modules.reports.domain.enums import (
    ArtifactStatus,
    ExportFormat,
    ReportLocale,
    ReportType,
)
from cold_storage.modules.reports.domain.render_model import (
    CanonicalRenderMetric,
    CanonicalRenderTableCell,
    CanonicalReportRenderModel,
)
from cold_storage.modules.reports.localization.catalog import (
    compute_catalog_content_hash,
    get_catalog,
)

PILOT_RESULT_SCHEMA_VERSION = "task11-pilot-report.v1"
PILOT_CHECK_ID = "multilingual_report_same_revision"
_SHA256_LENGTH = 64
_RENDER_MATRIX: tuple[tuple[ReportLocale, ExportFormat], ...] = (
    (ReportLocale.ZH_CN, ExportFormat.DOCX),
    (ReportLocale.ZH_CN, ExportFormat.PDF),
    (ReportLocale.EN_US, ExportFormat.DOCX),
    (ReportLocale.EN_US, ExportFormat.PDF),
)
DownloadArtifact = Callable[[str, str, str], tuple[bytes, Mapping[str, str]]]


class PilotVerificationError(EvaluationRunnerError):
    """Typed fail-closed error for a pilot acceptance mismatch."""

    def __init__(
        self,
        code: str,
        message: str,
        *,
        details: Mapping[str, Any] | None = None,
    ) -> None:
        super().__init__(message, details=dict(details or {}))
        self.code = code


def _fail(code: str, message: str, **details: Any) -> None:
    raise PilotVerificationError(code, message, details=details)


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _is_sha256(value: object) -> bool:
    if not isinstance(value, str) or len(value) != _SHA256_LENGTH:
        return False
    return all(char in "0123456789abcdef" for char in value)


# ---------------------------------------------------------------------------
# P1-3: Structured artifact observation + field-bound verification
# ---------------------------------------------------------------------------
#
# The P1-3 fix removes the previous global substring search
# (metric.display_value in extracted_text) and replaces it with:
#
#   DOWNLOADED_ARTIFACT
#   → STRUCTURED_OBSERVATION  (_observe_docx / _observe_pdf)
#   → SECTION/FIELD/ROW_BINDING  (_find_metric_binding / _find_number_binding
#                                 / _find_table_cell_binding)
#   → OBSERVED_VALUE_AND_UNIT    (_ObservedNumericField)
#   → EXPECTED_LOCALIZED_VALUE_AND_UNIT_COMPARISON  (whitespace fold only)
#   → FAIL_CLOSED   (8 typed codes, no false PASS)
#
# The numeric comparison NEVER operates on a flattened global string. It
# only operates on the structured observation scoped to the target
# section/field/row/column, and the observed value/unit come from the
# downloaded artifact bytes, NOT from the localized expected model.


_BINDING_KIND_METRIC = "metric"
_BINDING_KIND_NUMBER = "number"
_BINDING_KIND_TABLE_CELL = "table_cell"

_BINDING_KINDS: tuple[str, ...] = (
    _BINDING_KIND_METRIC,
    _BINDING_KIND_NUMBER,
    _BINDING_KIND_TABLE_CELL,
)


@dataclass(frozen=True, slots=True)
class _ObservedNumericField:
    """Real artifact observation for a single canonical numeric field.

    The fields are populated from the downloaded artifact (DOCX/PDF
    bytes), NOT from the localized expected model. The
    ``binding_kind`` and ``section_key`` describe the structural
    position the observation was taken from.
    """

    field_path: str
    section_key: str
    binding_kind: str  # "metric" | "number" | "table_cell"
    display_value: str
    display_unit: str
    row_index: int | None = None
    column_index: int | None = None
    page_number: int | None = None


@dataclass(frozen=True, slots=True)
class _BindingResult:
    """Outcome of a field-level binding lookup.

    Semantics:
    - On success: ``observed`` is set (artifact-derived), ``failure_code`` is None,
      ``candidates`` contains exactly the same single observation (audit trail).
    - On MISSING: ``observed`` is None, ``failure_code`` is a typed code,
      ``candidates`` is empty.
    - On AMBIGUOUS: ``observed`` is None, ``failure_code`` is a typed code,
      ``candidates`` contains the artifact-derived candidate observations
      (the audit consumer can see what WAS in the document, so the failure
      is not silent).
    """

    observed: _ObservedNumericField | None
    failure_code: str | None
    candidates: tuple[_ObservedNumericField, ...] = ()


def _fold_whitespace(text: str) -> str:
    """Collapse runs of internal whitespace to a single space, strip ends.

    Permitted per the P1-3 contract:
        - remove leading/trailing whitespace
        - collapse pure-typographic runs of whitespace
    NOT permitted:
        - rewriting decimal/thousands separators
        - fuzzy numeric tolerance
        - dropping sign or unit
    """

    return " ".join(text.split())


def _strings_equal_folded(a: str, b: str) -> bool:
    """Compare two strings after whitespace folding only."""

    return _fold_whitespace(a) == _fold_whitespace(b)


# ── DOCX observation ──────────────────────────────────────────────────────


@dataclass(frozen=True, slots=True)
class _DocxBlock:
    kind: str  # "paragraph" | "table"
    text: str
    paragraph_index: int | None
    table_index: int | None
    cells: tuple[tuple[str, ...], ...] | None  # only for tables
    heading_text: str | None  # for paragraphs that are headings
    heading_level: int | None  # 1..6 for heading paragraphs


@dataclass(frozen=True, slots=True)
class _DocxObservation:
    body_blocks: tuple[_DocxBlock, ...]
    # section_key -> (start_block_idx_inclusive, end_block_idx_exclusive)
    section_scopes: dict[str, tuple[int, int]]


def _observe_docx(data: bytes) -> _DocxObservation:
    """Walk a downloaded DOCX in body-order to produce structured blocks.

    The walk visits ``w:body`` children in original XML order. A
    ``w:p`` (paragraph) is recorded with its text + heading style
    detection; a ``w:tbl`` (table) is recorded as a 2-D cell grid.
    ``w:sectPr`` (the trailing section properties) is ignored.

    Section scope is determined by Heading 1 paragraphs: a Heading 1
    starts a new section, and the section extends until the next
    Heading 1 or end-of-document.
    """
    document = Document(io.BytesIO(data))
    body = document.element.body

    # First pass: identify Heading 1 paragraphs in body-order to build
    # section boundaries. We use the localized title text the renderer
    # would emit (Heading 1, ``w:pStyle w:val="Heading 1"``).
    blocks: list[_DocxBlock] = []
    para_counter = 0
    table_counter = 0
    for child in body:
        tag = child.tag.split("}", 1)[-1] if "}" in child.tag else child.tag
        if tag == "p":
            p_elem = child
            # Detect pStyle
            heading_text: str | None = None
            heading_level: int | None = None
            pPr = p_elem.find(_docx_qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(_docx_qn("w:pStyle"))
                if pStyle is not None:
                    style_val = pStyle.get(_docx_qn("w:val"), "")
                    if style_val.startswith("Heading"):
                        # Accept ``Heading1``, ``Heading 1``, ``Heading 1.0``,
                        # ``heading1`` (case-insensitive). The numeric
                        # suffix may or may not be separated by a space.
                        suffix = style_val[len("Heading") :].strip()
                        try:
                            heading_level = int(suffix.split(".")[0])
                        except (ValueError, IndexError):
                            heading_level = None
                        # Pull all text from this paragraph.
                        text = "".join(t.text or "" for t in p_elem.iter(_docx_qn("w:t")))
                        heading_text = text
            # Normal paragraph text (always)
            text = "".join(t.text or "" for t in p_elem.iter(_docx_qn("w:t")))
            blocks.append(
                _DocxBlock(
                    kind="paragraph",
                    text=text,
                    paragraph_index=para_counter,
                    table_index=None,
                    cells=None,
                    heading_text=heading_text,
                    heading_level=heading_level,
                )
            )
            para_counter += 1
        elif tag == "tbl":
            tbl_elem = child
            rows: list[tuple[str, ...]] = []
            for tr in tbl_elem.findall(_docx_qn("w:tr")):
                cells: list[str] = []
                for tc in tr.findall(_docx_qn("w:tc")):
                    cell_text = "".join(t.text or "" for t in tc.iter(_docx_qn("w:t")))
                    cells.append(cell_text)
                rows.append(tuple(cells))
            # Combined text of the table (one row per line).
            combined = "\n".join("|".join(row) for row in rows)
            blocks.append(
                _DocxBlock(
                    kind="table",
                    text=combined,
                    paragraph_index=None,
                    table_index=table_counter,
                    cells=tuple(rows),
                    heading_text=None,
                    heading_level=None,
                )
            )
            table_counter += 1
        elif tag == "sectPr":
            # Trailing sectPr — skip.
            continue
        # Other elements (e.g. sdt) are ignored for binding purposes.

    # Second pass: build section scopes from Heading 1 paragraphs.
    # We don't have a heading→section_key map at this layer; the
    # caller (P1-3 _semantic_checks) supplies the heading→section_key
    # map from the localized model. We expose heading_text on each
    # paragraph and let the caller resolve.
    heading_indices: list[int] = []
    for idx, block in enumerate(blocks):
        if block.kind == "paragraph" and block.heading_level == 1 and block.heading_text:
            heading_indices.append(idx)
    section_scopes: dict[str, tuple[int, int]] = {}
    return _DocxObservation(
        body_blocks=tuple(blocks),
        section_scopes=section_scopes,
    )


# ── PDF observation ───────────────────────────────────────────────────────


@dataclass(frozen=True, slots=True)
class _PdfLine:
    page_number: int
    block_index: int
    line_index: int
    text: str
    bbox: tuple[float, float, float, float]
    # Maximum font size of any span in this line (used for heading
    # detection in ``_resolve_pdf_section_scopes``).
    max_font_size: float = 10.0


@dataclass(frozen=True, slots=True)
class _PdfTableGrid:
    """A 2-D grid reconstructed from PDF text lines (rows clustered by y).

    For real-renderer PDF tables, the renderer draws each cell as a
    separate text line positioned at a specific (x, y). This is a
    best-effort reconstruction: lines with overlapping y-coordinates
    form a row; within a row, lines are sorted by x to get column
    order. The unit row is identified as the row whose leftmost cell
    text matches one of the canonical unit codes.
    """

    page_number: int
    rows: tuple[tuple[_PdfLine, ...], ...]
    # unit_row_idx: index of the row that contains the unit labels.
    unit_row_idx: int | None


@dataclass(frozen=True, slots=True)
class _PdfSectionTable:
    """A table reconstructed from a single section's coordinate scope.

    Per the P1-3 corrective contract, PDF tables are section-local,
    not page-global. Each table belongs to one section; multiple
    tables on the same page are distinguished by their y-band ranges
    (top-to-bottom in document order). The table is row-aware and
    page-aware: each row carries the page number and the y-bands
    so that ``_find_table_cell_binding`` can map (row_index,
    column_index) to a specific PDF line.
    """

    section_key: str
    page_number: int
    # A table is a sequence of "visual rows". Each visual row is a
    # tuple of PDF lines that share the same y-band (within the
    # section's y tolerance). A visual row may have 0..N cells
    # (an empty unit cell produces NO PDF line and is therefore
    # absent from the row tuple). Cell identity is by x-band,
    # aligned against the header row's x centers.
    rows: tuple[tuple[_PdfLine, ...], ...]
    # Bounding box of the table on the page (used for table identity).
    bbox: tuple[float, float, float, float]
    # The first row's column centers, used to map a unit/data line
    # to its column index.
    column_centers: tuple[float, ...]


@dataclass(frozen=True, slots=True)
class _PdfObservation:
    all_lines: tuple[_PdfLine, ...]
    # Per Corrective 4, tables are no longer stored on the page-global
    # observation. They are reconstructed on demand per section using
    # the section's coordinate scope (see ``_build_section_local_tables``).
    section_scopes: dict[str, tuple[int, int]]
    # ── Corrective 2 ──
    # Text spans with explicit (page, bbox, text) triples. Each
    # span is a single text-run emitted by the renderer; a wrapped
    # header / wrapped data cell becomes 2+ spans that share the
    # same column bbox.
    text_spans: tuple[_PdfTextSpan, ...] = ()
    # ── Corrective 2 ──
    # Vector drawing segments (horizontal + vertical lines) emitted
    # by the renderer. Used to reconstruct row/column grid geometry
    # for real-renderer PDF tables.
    grid_segments: tuple[_PdfGridSegment, ...] = ()


# ── Corrective 2 ────────────────────────────────────────────────────────
# Grid-geometry primitives for real-renderer PDF tables.


@dataclass(frozen=True, slots=True)
class _PdfTextSpan:
    page_number: int
    text: str
    bbox: tuple[float, float, float, float]


@dataclass(frozen=True, slots=True)
class _PdfGridSegment:
    page_number: int
    orientation: str  # "horizontal" | "vertical"
    x0: float
    y0: float
    x1: float
    y1: float


@dataclass(frozen=True, slots=True)
class _PdfLogicalCell:
    page_number: int
    row_index: int
    column_index: int
    bbox: tuple[float, float, float, float]
    text: str


@dataclass(frozen=True, slots=True)
class _PdfLogicalRow:
    page_number: int
    cells: tuple[_PdfLogicalCell, ...]
    row_kind: str  # "header" | "unit" | "data" | "unknown"


# ── Corrective 4 ────────────────────────────────────────────────────────
# Table segment + logical table (cross-page continuation).


@dataclass(frozen=True, slots=True)
class _PdfTableSegment:
    section_key: str
    page_number: int
    header: _PdfLogicalRow
    unit_row: _PdfLogicalRow | None
    data_rows: tuple[_PdfLogicalRow, ...]
    bbox: tuple[float, float, float, float]


@dataclass(frozen=True, slots=True)
class _PdfLogicalTable:
    section_key: str
    segments: tuple[_PdfTableSegment, ...]
    data_rows: tuple[_PdfLogicalRow, ...]  # flat across all segments
    header: _PdfLogicalRow  # the first segment's header (canonical)


_Y_TOLERANCE = 2.0  # pixels for clustering lines into rows
_GRID_Y_TOLERANCE = 1.5  # grid line clustering (sub-pixel)
_GRID_X_TOLERANCE = 1.5


def _cluster_lines_into_rows(
    lines: tuple[_PdfLine, ...], *, y_tolerance: float = _Y_TOLERANCE
) -> tuple[tuple[_PdfLine, ...], ...]:
    """Group PDF lines into rows by y-coordinate, sorted top-to-bottom."""

    if not lines:
        return ()
    sorted_lines = sorted(lines, key=lambda ln: (ln.bbox[1], ln.bbox[0]))
    rows: list[list[_PdfLine]] = []
    current_row: list[_PdfLine] = [sorted_lines[0]]
    for line in sorted_lines[1:]:
        if abs(line.bbox[1] - current_row[0].bbox[1]) <= y_tolerance:
            current_row.append(line)
        else:
            rows.append(current_row)
            current_row = [line]
    rows.append(current_row)
    # Sort each row by x (left-to-right).
    return tuple(tuple(sorted(row, key=lambda ln: ln.bbox[0])) for row in rows)


def _observe_pdf(data: bytes) -> _PdfObservation:
    """Walk a downloaded PDF to produce structured lines (per-line spatial layout).

    Uses ``page.get_text("dict")`` to extract blocks → lines → spans
    while preserving the spatial layout. Each line is recorded with
    its page number, block index, line index, text, and bounding box.

    Per Corrective 2, the observation also collects:

      * ``text_spans`` — individual text spans (page, bbox, text) so
        wrapped headers / wrapped data cells are visible as 2+ spans
        sharing the same column bbox.
      * ``grid_segments`` — vector drawing segments (horizontal +
        vertical lines) emitted by the real ``PdfRenderer``. These
        drive the structural grid-geometry reconstruction in
        ``_build_section_local_tables``.

    Per Corrective 4, table extraction is NOT performed here. Tables
    are reconstructed on demand per section from the section's
    coordinate scope (see ``_build_section_local_tables``). The old
    page-global heuristic that merged all multi-row clusters on a
    page into a single table is gone.
    """

    all_lines: list[_PdfLine] = []
    text_spans: list[_PdfTextSpan] = []
    grid_segments: list[_PdfGridSegment] = []
    with fitz.open(stream=data, filetype="pdf") as document:
        for page_index, page in enumerate(document):
            page_number = page_index + 1
            d = page.get_text("dict")
            for block_index, block in enumerate(d.get("blocks", [])):
                if block.get("type") == 1:
                    # Image block — skip text extraction.
                    continue
                for line_index, line in enumerate(block.get("lines", [])):
                    line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                    line_bbox_tuple = line.get("bbox", (0.0, 0.0, 0.0, 0.0))
                    line_max_font_size = max(
                        (float(span.get("size", 0.0)) for span in line.get("spans", [])),
                        default=10.0,
                    )
                    if line_text:
                        all_lines.append(
                            _PdfLine(
                                page_number=page_number,
                                block_index=block_index,
                                line_index=line_index,
                                text=line_text,
                                bbox=(
                                    float(line_bbox_tuple[0]),
                                    float(line_bbox_tuple[1]),
                                    float(line_bbox_tuple[2]),
                                    float(line_bbox_tuple[3]),
                                ),
                                max_font_size=line_max_font_size,
                            )
                        )
                    for span in line.get("spans", []):
                        span_text = span.get("text", "")
                        if not span_text:
                            continue
                        span_bbox = span.get("bbox", (0.0, 0.0, 0.0, 0.0))
                        text_spans.append(
                            _PdfTextSpan(
                                page_number=page_number,
                                text=span_text,
                                bbox=(
                                    float(span_bbox[0]),
                                    float(span_bbox[1]),
                                    float(span_bbox[2]),
                                    float(span_bbox[3]),
                                ),
                            )
                        )
            # ── Corrective 2: collect vector drawings ──
            try:
                drawings = page.get_drawings()
            except Exception:  # pragma: no cover -- defensive
                drawings = []
            for drawing in drawings:
                for item in drawing.get("items", []):
                    item_type = item[0]
                    if item_type == "l":  # line
                        p0, p1 = item[1], item[2]
                        x0 = float(p0.x)
                        y0 = float(p0.y)
                        x1 = float(p1.x)
                        y1 = float(p1.y)
                    elif item_type == "re":  # rectangle → 4 edges
                        rect = item[1]
                        x0 = float(rect.x0)
                        y0 = float(rect.y0)
                        x1 = float(rect.x1)
                        y1 = float(rect.y1)
                        # Emit the 4 edges as line segments.
                        for ex0, ey0, ex1, ey1 in (
                            (x0, y0, x1, y0),  # top
                            (x0, y1, x1, y1),  # bottom
                            (x0, y0, x0, y1),  # left
                            (x1, y0, x1, y1),  # right
                        ):
                            if abs(ex0 - ex1) < _GRID_X_TOLERANCE:
                                orientation = "vertical"
                            elif abs(ey0 - ey1) < _GRID_Y_TOLERANCE:
                                orientation = "horizontal"
                            else:
                                # Diagonal line — ignore.
                                continue
                            grid_segments.append(
                                _PdfGridSegment(
                                    page_number=page_number,
                                    orientation=orientation,
                                    x0=ex0,
                                    y0=ey0,
                                    x1=ex1,
                                    y1=ey1,
                                )
                            )
                        continue
                    else:
                        # Curve / other — ignore for grid geometry.
                        continue
                    if abs(x0 - x1) < _GRID_X_TOLERANCE:
                        orientation = "vertical"
                    elif abs(y0 - y1) < _GRID_Y_TOLERANCE:
                        orientation = "horizontal"
                    else:
                        continue
                    grid_segments.append(
                        _PdfGridSegment(
                            page_number=page_number,
                            orientation=orientation,
                            x0=x0,
                            y0=y0,
                            x1=x1,
                            y1=y1,
                        )
                    )

    return _PdfObservation(
        all_lines=tuple(all_lines),
        section_scopes={},
        text_spans=tuple(text_spans),
        grid_segments=tuple(grid_segments),
    )


def _build_section_local_tables(
    *,
    pdf_observation: _PdfObservation,
    section_scopes: Mapping[str, tuple[int, int]],
    section_table_headers: Mapping[str, tuple[str, ...]],
) -> tuple[_PdfSectionTable, ...]:
    """Reconstruct tables for each section from the section's coordinate scope.

    Per Corrective 3+4, tables are recognized by **structural
    identity**: the section-local header row is the one whose
    text matches the localized table headers (folded-exact match
    on column count + per-cell match). A section MAY contain
    multiple structurally identical tables (e.g. a comparison
    table repeated in the same section); in that case the
    cell-binding MUST fail-closed as
    ``AMBIGUOUS_FIELD_BINDING`` (see Corrective 4). This
    reconstruction therefore emits ONE ``_PdfSectionTable`` per
    matching header row, so the cell-binding can detect
    multi-table ambiguity and not silently pick the first.

    The reconstruction algorithm:

      1. For each section, walk the section's lines in
         (page_number, y, x) order, and group consecutive lines
         (within ``_Y_TOLERANCE`` on y) into y-bands.
      2. For each y-band with ``len(band) == len(expected_headers)``,
         folded-exact compare the band's text to
         ``expected_headers``. On match, mark the band as a
         HEADER row.
      3. All lines between one header row and the next
         (within the section) form ONE table's body. The table
         bbox is the union of header + body line bboxes.
      4. Body lines are clustered by (page, y) into visual rows.

    If no header row is found, no table is emitted for the
    section. If multiple header rows are found, multiple
    ``_PdfSectionTable`` records are emitted (and the
    cell-binding will return ``AMBIGUOUS_FIELD_BINDING``).
    """

    if not section_scopes:
        return ()
    tables: list[_PdfSectionTable] = []
    for section_key, (start, end) in section_scopes.items():
        expected_headers = section_table_headers.get(section_key, ())
        if not expected_headers:
            continue
        section_lines = pdf_observation.all_lines[start:end]
        if not section_lines:
            continue
        # Group section lines into y-bands (rows), page-aware.
        sorted_section = sorted(
            section_lines, key=lambda ln: (ln.page_number, ln.bbox[1], ln.bbox[0])
        )
        bands: list[list[_PdfLine]] = []
        current_band: list[_PdfLine] = [sorted_section[0]]
        for ln in sorted_section[1:]:
            if (
                ln.page_number == current_band[0].page_number
                and abs(ln.bbox[1] - current_band[0].bbox[1]) <= _Y_TOLERANCE
            ):
                current_band.append(ln)
            else:
                bands.append(sorted(current_band, key=lambda x: x.bbox[0]))
                current_band = [ln]
        bands.append(sorted(current_band, key=lambda x: x.bbox[0]))
        # Find all band indices that match expected_headers.
        header_band_indices: list[int] = []
        for i, band in enumerate(bands):
            if len(band) != len(expected_headers):
                continue
            folded = tuple(_fold_whitespace(ln.text) for ln in band)
            if folded == tuple(_fold_whitespace(h) for h in expected_headers):
                header_band_indices.append(i)
        if not header_band_indices:
            continue
        # For each header band, build a _PdfSectionTable with
        # body = bands in (header_idx, next_header_idx).
        # Per Corrective 4: a band on a NEW page is the start of
        # a NEW physical table segment (a repeated header is a
        # continuation, but a body band on a new page without an
        # intervening header is a new table).
        for table_idx, header_band_i in enumerate(header_band_indices):
            body_bands = bands[header_band_i + 1 :]
            # Truncate at next header band first.
            if table_idx + 1 < len(header_band_indices):
                next_header_i = header_band_indices[table_idx + 1]
                body_bands = bands[header_band_i + 1 : next_header_i]
            # Per Corrective 4: also truncate when a band starts
            # on a different page than the header (a body band on
            # a new page without an intervening header band is a
            # new table's header).
            header_page = bands[header_band_i][0].page_number
            filtered_body_bands: list[list[_PdfLine]] = []
            for body_band in body_bands:
                if body_band[0].page_number != header_page:
                    break
                filtered_body_bands.append(body_band)
            body_bands = filtered_body_bands
            header_band = bands[header_band_i]
            all_lines_for_table: list[_PdfLine] = list(header_band)
            for body_band in body_bands:
                all_lines_for_table.extend(body_band)
            x0 = min(ln.bbox[0] for ln in all_lines_for_table)
            y0 = min(ln.bbox[1] for ln in all_lines_for_table)
            x1 = max(ln.bbox[2] for ln in all_lines_for_table)
            y1 = max(ln.bbox[3] for ln in all_lines_for_table)
            column_centers = tuple((ln.bbox[0] + ln.bbox[2]) / 2 for ln in header_band)
            # The table's rows = (header_band, *body_bands).
            rows: list[tuple[_PdfLine, ...]] = [tuple(header_band)]
            rows.extend(tuple(b) for b in body_bands)
            tables.append(
                _PdfSectionTable(
                    section_key=section_key,
                    page_number=header_band[0].page_number,
                    rows=tuple(rows),
                    bbox=(x0, y0, x1, y1),
                    column_centers=column_centers,
                )
            )
    return tuple(tables)


# ── Corrective 2/3/4/5: grid-geometry logical table reconstruction ────────


def _cluster_grid_segments_by_coord(
    segments: tuple[_PdfGridSegment, ...],
    *,
    axis: str,
    tolerance: float,
) -> list[float]:
    """Cluster colinear grid segments by their perpendicular axis.

    Returns the list of cluster centers (floats), sorted in
    ascending order. Used to derive row boundaries (from
    horizontal lines) and column boundaries (from vertical lines).
    """

    if axis == "x":
        coords = sorted(s.x0 for s in segments if abs(s.x0 - s.x1) < 1e-3)
    else:
        coords = sorted(s.y0 for s in segments if abs(s.y0 - s.y1) < 1e-3)
    if not coords:
        return []
    clusters: list[list[float]] = [[coords[0]]]
    for c in coords[1:]:
        if abs(c - clusters[-1][-1]) <= tolerance:
            clusters[-1].append(c)
        else:
            clusters.append([c])
    return [sum(cluster) / len(cluster) for cluster in clusters]


def _row_bbox(row: _PdfLogicalRow) -> tuple[float, float, float, float]:
    """Return the (x0, y0, x1, y1) bbox spanning all cells of a row."""
    cells = row.cells
    if not cells:
        return (0.0, 0.0, 0.0, 0.0)
    x0 = min(c.bbox[0] for c in cells)
    y0 = min(c.bbox[1] for c in cells)
    x1 = max(c.bbox[2] for c in cells)
    y1 = max(c.bbox[3] for c in cells)
    return (x0, y0, x1, y1)


def _row_column_xs(row: _PdfLogicalRow) -> tuple[float, ...]:
    """Sorted cell x-centers of a row (used for column-band check)."""
    return tuple(sorted((c.bbox[0] + c.bbox[2]) / 2 for c in row.cells))


def _body_line_band_for(
    line: _PdfLine,
    column_centers: tuple[float, ...],
    half_band: float | None = None,
) -> int | None:
    """Best-fit column band index for a body line, or None if no column
    center fits within the band (used by text-only structural check).
    """
    if not column_centers:
        return None
    line_center = (line.bbox[0] + line.bbox[2]) / 2
    sorted_centers = sorted(column_centers)
    if half_band is None:
        if len(sorted_centers) < 2:
            half_band = 100.0
        else:
            half_band = max(
                (sorted_centers[i + 1] - sorted_centers[i]) / 2
                for i in range(len(sorted_centers) - 1)
            )
    # Use a tolerant half-band so x-jitter from the renderer doesn't
    # cause a band mismatch when columns are clearly the same.
    band = half_band
    best_idx = -1
    best_dist = float("inf")
    for i, c in enumerate(sorted_centers):
        d = abs(line_center - c)
        if d <= band and d < best_dist:
            best_dist = d
            best_idx = i
    return best_idx if best_idx >= 0 else None


def _pdf_page_y_range(
    pdf_observation: _PdfObservation,
    page_number: int,
) -> tuple[float, float] | None:
    """Return (min_y, max_y) observed on the given page from artifact coords."""
    ys: list[float] = []
    for line in pdf_observation.all_lines:
        if line.page_number != page_number:
            continue
        ys.append(line.bbox[1])
        ys.append(line.bbox[3])
    if not ys:
        return None
    return (min(ys), max(ys))


def _has_intervening_marker(
    pdf_observation: _PdfObservation,
    prev_page: int,
    prev_bot_y: float,
    curr_page: int,
    curr_top_y: float,
) -> bool:
    """True if a substantive text block sits between prev_bot_y (on
    prev_page) and curr_top_y (on curr_page), suggesting an
    intervening heading / independent table marker.
    """
    # Only same-page post-segment text counts: page transitions are
    # already constrained by the page-bottom / page-top region tests.
    for line in pdf_observation.all_lines:
        if line.page_number != prev_page:
            continue
        if line.bbox[3] <= prev_bot_y + 2.0:
            continue
        text = _fold_whitespace(line.text)
        if not text or len(text) < 2:
            continue
        if _is_renderer_unit_token(text):
            continue
        # Distinct paragraph-like content between segments suggests
        # an intervening heading / text block, which would break
        # the strict continuation evidence.
        return True
    return False


def _is_pdf_table_continuation(
    *,
    previous: _PdfTableSegment,
    current: _PdfTableSegment,
    pdf_observation: _PdfObservation | None,
    section_line_range: tuple[int, int],
) -> bool:
    """Strict continuation predicate for cross-page PDF tables.

    Per P1-3.2 of the fourth corrective. Returns True only when
    ALL of the following structural predicates hold:

      SAME_SECTION=YES
      SAME_HEADER_TEXT=YES (folded-exact)
      SAME_COLUMN_COUNT=YES
      SAME_COLUMN_X_BANDS=YES (within tolerance)
      CURRENT_PAGE=PREVIOUS_PAGE+1 (immediately adjacent)
      PREVIOUS_SEGMENT_REACHES_PAGE_BOTTOM_REGION=YES
      CURRENT_SEGMENT_STARTS_NEAR_PAGE_TOP_TABLE_REGION=YES
      NO_INTERVENING_SECTION_HEADING=YES
      NO_INTERVENING_INDEPENDENT_TABLE_MARKER=YES

    page-bottom / page-top region are derived from artifact
    coordinates (per-page observed y-extents), NOT hard-coded
    PDF page heights.
    """
    if previous.section_key != current.section_key:
        return False
    if previous.page_number + 1 != current.page_number:
        return False
    prev_header_cells = previous.header.cells
    curr_header_cells = current.header.cells
    if len(prev_header_cells) != len(curr_header_cells):
        return False
    if not prev_header_cells:
        return False
    if tuple(_fold_whitespace(c.text) for c in prev_header_cells) != tuple(
        _fold_whitespace(c.text) for c in curr_header_cells
    ):
        return False
    if pdf_observation is None:
        return False
    prev_data_rows = previous.data_rows
    if not prev_data_rows:
        return False
    last_prev_data_row = prev_data_rows[-1]
    prev_row_bot_y = _row_bbox(last_prev_data_row)[3]
    curr_header_top_y = _row_bbox(current.header)[1]
    prev_page_y_range = _pdf_page_y_range(pdf_observation, previous.page_number)
    curr_page_y_range = _pdf_page_y_range(pdf_observation, current.page_number)
    if prev_page_y_range is None or curr_page_y_range is None:
        return False
    prev_min_y, prev_max_y = prev_page_y_range
    curr_min_y, curr_max_y = curr_page_y_range
    prev_page_height = max(prev_max_y - prev_min_y, 1.0)
    curr_page_height = max(curr_max_y - curr_min_y, 1.0)
    prev_row_in_bottom_region = (prev_row_bot_y - prev_min_y) >= 0.75 * prev_page_height
    curr_header_in_top_region = (curr_header_top_y - curr_min_y) <= 0.25 * curr_page_height
    if not (prev_row_in_bottom_region and curr_header_in_top_region):
        return False
    # SAME_COLUMN_X_BANDS: column x-centers must match within tolerance.
    prev_col_xs = _row_column_xs(previous.header)
    curr_col_xs = _row_column_xs(current.header)
    if len(prev_col_xs) != len(curr_col_xs):
        return False
    max_x_drift = max(abs(a - b) for a, b in zip(prev_col_xs, curr_col_xs, strict=False))
    if max_x_drift > _GRID_X_TOLERANCE * 4:
        return False
    return not _has_intervening_marker(
        pdf_observation,
        previous.page_number,
        prev_row_bot_y,
        current.page_number,
        curr_header_top_y,
    )


def _build_logical_tables_for_section(
    *,
    pdf_observation: _PdfObservation,
    section_key: str,
    section_line_range: tuple[int, int],
    expected_headers: tuple[str, ...],
) -> tuple[_PdfLogicalTable, ...]:
    """Build ``_PdfLogicalTable``s for one section from grid geometry.

    Per Correctives 2/3/4/5, the table's row/column structure is
    derived from the renderer's vector drawing segments:

      * horizontal grid lines → row boundaries
      * vertical grid lines → column boundaries
      * text spans whose bbox intersects a (row, column) cell →
        that cell's text

    A ``_PdfLogicalRow`` is constructed per (row, column) cell. The
    header row is the first row whose cells' folded text matches
    ``expected_headers`` (column count + per-cell text match).

    The unit row (when present) is the row immediately after the
    header whose cells are structurally a unit row:
      * row is an independent grid row
      * each non-empty cell is a renderer unit token (e.g.
        ``(kW(e))``, ``(CNY)``, ``m²``, ``个``, ``-``, or empty)

    Per Corrective 4, cross-page continuation is detected when
    a later segment in the same section starts with the same
    localized headers on a later page, the previous segment's
    last row was on the page-bottom continuation region, and the
    two segments have no intervening canonical-section heading.
    Such continuations are merged into ONE ``_PdfLogicalTable``
    with data_rows flat across all segments.
    """

    if not expected_headers:
        return ()
    start, end = section_line_range
    section_lines = pdf_observation.all_lines[start:end]
    if not section_lines:
        return ()
    # Identify the (page, y_top, y_bottom) of every grid row by
    # scanning horizontal grid lines that lie within the section's
    # line y-range. The horizontal lines define row boundaries.
    # The first line's y_top is the start of the first row.
    line_y0 = min(ln.bbox[1] for ln in section_lines)
    line_y1 = max(ln.bbox[3] for ln in section_lines)
    h_grid = [
        g
        for g in pdf_observation.grid_segments
        if g.orientation == "horizontal" and line_y0 - 50.0 <= (g.y0 + g.y1) / 2 <= line_y1 + 50.0
    ]
    v_grid = [
        g
        for g in pdf_observation.grid_segments
        if g.orientation == "vertical" and line_y0 - 50.0 <= (g.y0 + g.y1) / 2 <= line_y1 + 50.0
    ]
    if not h_grid or not v_grid:
        # No grid geometry available — fall back to text-only path
        # (the existing _build_section_local_tables handles this).
        return ()
    row_boundaries = _cluster_grid_segments_by_coord(
        tuple(h_grid), axis="y", tolerance=_GRID_Y_TOLERANCE
    )
    column_boundaries = _cluster_grid_segments_by_coord(
        tuple(v_grid), axis="x", tolerance=_GRID_X_TOLERANCE
    )
    if len(row_boundaries) < 2 or len(column_boundaries) < len(expected_headers) + 1:
        return ()
    # Row strips: row i is (row_boundaries[i], row_boundaries[i+1]).
    # Column strips: column j is (column_boundaries[j], column_boundaries[j+1]).
    segments: list[_PdfTableSegment] = []
    section_spans = [
        s
        for s in pdf_observation.text_spans
        if any(ln.page_number == s.page_number for ln in section_lines[:1])
    ]
    # Group section_spans by page.
    pages: dict[int, list[_PdfTextSpan]] = {}
    for span in section_spans:
        pages.setdefault(span.page_number, []).append(span)

    # Identify candidate header rows by looking for rows whose
    # per-column folded text matches the expected headers.
    header_rows_by_page: dict[int, list[int]] = {}
    for page_number, page_spans in pages.items():
        for row_idx in range(len(row_boundaries) - 1):
            y_top = row_boundaries[row_idx]
            y_bot = row_boundaries[row_idx + 1]
            row_cells = _collect_row_cells(
                page_spans,
                row_y_top=y_top,
                row_y_bot=y_bot,
                column_boundaries=column_boundaries,
                page_number=page_number,
                row_index=row_idx,
            )
            if not row_cells or len(row_cells) != len(expected_headers):
                continue
            folded = tuple(_fold_whitespace(c.text) for c in row_cells)
            if folded == tuple(_fold_whitespace(h) for h in expected_headers):
                header_rows_by_page.setdefault(page_number, []).append(row_idx)

    if not header_rows_by_page:
        return ()
    # Build segments: each header row defines one segment.
    for page_number, header_indices in header_rows_by_page.items():
        for header_idx in header_indices:
            seg = _build_segment_for_header(
                page_spans=pages[page_number],
                header_row_idx=header_idx,
                row_boundaries=row_boundaries,
                column_boundaries=column_boundaries,
                page_number=page_number,
                section_key=section_key,
            )
            if seg is not None:
                segments.append(seg)
    if not segments:
        return ()
    # Cross-page continuation: a later segment with same headers
    # and on a later page is a continuation of the previous if the
    # previous segment's last data row is on the page-bottom
    # region AND there is no canonical-section heading between
    # them. Here we merge sequential segments with identical
    # headers and ascending page numbers into a single logical
    # table.
    logical_tables: list[_PdfLogicalTable] = []
    current: list[_PdfTableSegment] = [segments[0]]
    for seg in segments[1:]:
        prev = current[-1]
        # Per P1-3.2 of the fourth corrective: cross-page continuation
        # requires structural evidence (page-bottom / page-top region
        # + same column bands + ascending page + same headers), NOT
        # just header text + ascending page number. Independent tables
        # are NOT merged.
        if _is_pdf_table_continuation(
            previous=prev,
            current=seg,
            pdf_observation=pdf_observation,
            section_line_range=section_line_range,
        ):
            current.append(seg)
            continue
        logical_tables.append(
            _PdfLogicalTable(
                section_key=section_key,
                segments=tuple(current),
                data_rows=tuple(row for seg in current for row in seg.data_rows),
                header=current[0].header,
            )
        )
        current = [seg]
    if current:
        logical_tables.append(
            _PdfLogicalTable(
                section_key=section_key,
                segments=tuple(current),
                data_rows=tuple(row for seg in current for row in seg.data_rows),
                header=current[0].header,
            )
        )
    return tuple(logical_tables)


def _collect_row_cells(
    spans: list[_PdfTextSpan],
    *,
    row_y_top: float,
    row_y_bot: float,
    column_boundaries: list[float],
    page_number: int,
    row_index: int,
) -> tuple[_PdfLogicalCell, ...]:
    """Collect text spans inside a single grid row into logical cells.

    Each span is assigned to a column based on the column
    boundary whose center is closest to the span's bbox center.
    Multiple spans in the same column are concatenated into one
    logical cell (wrapped text → one logical cell). Spans with
    empty text are ignored.
    """
    if len(column_boundaries) < 2:
        return ()
    centers = [
        (column_boundaries[i] + column_boundaries[i + 1]) / 2
        for i in range(len(column_boundaries) - 1)
    ]
    cells_by_col: dict[int, list[_PdfTextSpan]] = {}
    for span in spans:
        cy = (span.bbox[1] + span.bbox[3]) / 2
        if not (row_y_top - 1.0 <= cy <= row_y_bot + 1.0):
            continue
        cx = (span.bbox[0] + span.bbox[2]) / 2
        col_idx = min(
            range(len(centers)),
            key=lambda j: abs(centers[j] - cx),
        )
        cells_by_col.setdefault(col_idx, []).append(span)
    cells: list[_PdfLogicalCell] = []
    for col_idx in range(len(centers)):
        col_spans = sorted(
            cells_by_col.get(col_idx, []),
            key=lambda s: (s.bbox[1], s.bbox[0]),
        )
        if not col_spans:
            cells.append(
                _PdfLogicalCell(
                    page_number=page_number,
                    row_index=row_index,
                    column_index=col_idx,
                    bbox=(0.0, 0.0, 0.0, 0.0),
                    text="",
                )
            )
            continue
        text = "".join(s.text for s in col_spans)
        bbox = (
            min(s.bbox[0] for s in col_spans),
            min(s.bbox[1] for s in col_spans),
            max(s.bbox[2] for s in col_spans),
            max(s.bbox[3] for s in col_spans),
        )
        cells.append(
            _PdfLogicalCell(
                page_number=page_number,
                row_index=row_index,
                column_index=col_idx,
                bbox=bbox,
                text=text,
            )
        )
    return tuple(cells)


# Generic unit-token patterns: per Corrective 3, the unit-row
# parser MUST support parenthesized tokens (including tokens with
# nested parens like ``(kW(e))`` and ``(kW(r))``), bare tokens,
# ``m²`` / ``个``, the literal ``-``, and the empty string.

# Match an outer paren wrapping, allowing ONE level of nesting
# inside (so ``(kW(e))`` is matched as a single token). This
# mirrors the renderer's paren-wrapping for unit tokens.
_UNIT_TOKEN_PAREN_RE = re.compile(r"^\([^()]*(?:\([^()]*\)[^()]*)*\)$")


def _is_renderer_unit_token(text: str) -> bool:
    """Heuristic: does this cell text look like a renderer unit token?

    Accepts:
      * parenthesized tokens: (kW(e)), (CNY), (-), ()
      * bare tokens: kW(e), CNY, m², 个, -
      * the empty string (allowed empty unit cell)
    Rejects pure data-row text (numbers / letters that are not
    known unit forms).
    """
    folded = _fold_whitespace(text)
    if folded == "":
        return True
    if _UNIT_TOKEN_PAREN_RE.match(folded):
        return True
    if folded == "-":
        return True
    # Known short tokens used by the renderer / localization.
    known_short = {"m²", "个", "kW(e)", "kW(r)", "CNY", "USD"}
    return folded in known_short


def _build_segment_for_header(
    *,
    page_spans: list[_PdfTextSpan],
    header_row_idx: int,
    row_boundaries: list[float],
    column_boundaries: list[float],
    page_number: int,
    section_key: str,
) -> _PdfTableSegment | None:
    """Build a single ``_PdfTableSegment`` starting at ``header_row_idx``."""

    if header_row_idx >= len(row_boundaries) - 1:
        return None
    header_y_top = row_boundaries[header_row_idx]
    header_y_bot = row_boundaries[header_row_idx + 1]
    header_cells = _collect_row_cells(
        page_spans,
        row_y_top=header_y_top,
        row_y_bot=header_y_bot,
        column_boundaries=column_boundaries,
        page_number=page_number,
        row_index=header_row_idx,
    )
    if not header_cells:
        return None
    header = _PdfLogicalRow(
        page_number=page_number,
        cells=header_cells,
        row_kind="header",
    )
    # Unit row: the row immediately after the header, if its
    # cells structurally look like a unit row. Per Corrective 3
    # the unit-row detection is STRUCTURAL, NOT based on row
    # count alone. We detect by checking each non-empty cell's
    # text matches the renderer unit-token grammar.
    unit_row: _PdfLogicalRow | None = None
    data_rows: list[_PdfLogicalRow] = []
    body_idx = header_row_idx + 1
    if body_idx < len(row_boundaries) - 1:
        body_y_top = row_boundaries[body_idx]
        body_y_bot = row_boundaries[body_idx + 1]
        body_cells = _collect_row_cells(
            page_spans,
            row_y_top=body_y_top,
            row_y_bot=body_y_bot,
            column_boundaries=column_boundaries,
            page_number=page_number,
            row_index=body_idx,
        )
        non_empty_cells = [c for c in body_cells if _fold_whitespace(c.text)]
        if non_empty_cells and all(_is_renderer_unit_token(c.text) for c in non_empty_cells):
            unit_row = _PdfLogicalRow(
                page_number=page_number,
                cells=body_cells,
                row_kind="unit",
            )
            data_start_idx = body_idx + 1
        else:
            data_start_idx = body_idx
        # Data rows: all subsequent rows until end of section.
        for idx in range(data_start_idx, len(row_boundaries) - 1):
            y_top = row_boundaries[idx]
            y_bot = row_boundaries[idx + 1]
            cells = _collect_row_cells(
                page_spans,
                row_y_top=y_top,
                row_y_bot=y_bot,
                column_boundaries=column_boundaries,
                page_number=page_number,
                row_index=idx,
            )
            data_rows.append(
                _PdfLogicalRow(
                    page_number=page_number,
                    cells=cells,
                    row_kind="data",
                )
            )
    x0 = column_boundaries[0]
    y0 = row_boundaries[header_row_idx]
    x1 = column_boundaries[-1]
    y1 = (
        row_boundaries[data_start_idx]
        if data_start_idx < len(row_boundaries)
        else row_boundaries[-1]
    )
    return _PdfTableSegment(
        section_key=section_key,
        page_number=page_number,
        header=header,
        unit_row=unit_row,
        data_rows=tuple(data_rows),
        bbox=(x0, y0, x1, y1),
    )


# ── Section-scope resolution ──────────────────────────────────────────────


@dataclass(frozen=True, slots=True)
class _SectionScope:
    section_key: str
    heading_text: str


def _build_section_scopes(
    *,
    localized_sections: tuple[Any, ...],
) -> tuple[_SectionScope, ...]:
    """Build (section_key, localized_heading_text) tuples in document order."""

    return tuple(
        _SectionScope(section_key=section.section_key, heading_text=section.title)
        for section in localized_sections
    )


def _resolve_docx_section_scopes(
    *,
    observation: _DocxObservation,
    section_scopes: tuple[_SectionScope, ...],
) -> dict[str, tuple[int, int]]:
    """Map each section_key to (start_block_idx, end_block_idx) in DOCX order.

    A section's heading paragraph MUST match the localized section
    title exactly (after whitespace folding). If a section's heading
    cannot be found, the section is omitted from the scope map (the
    caller MUST treat the field as MISSING_FIELD_BINDING).
    """

    heading_to_key: dict[str, str] = {
        _fold_whitespace(s.heading_text): s.section_key for s in section_scopes
    }
    used: set[str] = set()
    resolved: dict[str, tuple[int, int]] = {}
    blocks = observation.body_blocks
    # Find the first block of each section (by heading match).
    starts: list[tuple[int, str]] = []  # (block_idx, section_key)
    for idx, block in enumerate(blocks):
        if block.kind == "paragraph" and block.heading_level == 1 and block.heading_text:
            folded = _fold_whitespace(block.heading_text)
            if folded in heading_to_key:
                section_key = heading_to_key[folded]
                if section_key not in used:
                    starts.append((idx, section_key))
                    used.add(section_key)
    for i, (start_idx, section_key) in enumerate(starts):
        end_idx = starts[i + 1][0] if i + 1 < len(starts) else len(blocks)
        resolved[section_key] = (start_idx, end_idx)
    return resolved


def _resolve_pdf_section_scopes(
    *,
    observation: _PdfObservation,
    section_scopes: tuple[_SectionScope, ...],
) -> dict[str, tuple[int, int]]:
    """Map each section_key to (start_line_idx, end_line_idx) in PDF order.

    Uses the y-coordinate of the heading line: a heading is a line
    whose text matches the localized section title (whitespace-folded)
    and whose y-coordinate is above a small threshold (i.e. it is a
    section divider, not a body line).
    """

    heading_to_key: dict[str, str] = {
        _fold_whitespace(s.heading_text): s.section_key for s in section_scopes
    }
    used: set[str] = set()
    starts: list[tuple[int, str]] = []
    for idx, line in enumerate(observation.all_lines):
        folded = _fold_whitespace(line.text)
        if folded in heading_to_key:
            section_key = heading_to_key[folded]
            if section_key not in used:
                starts.append((idx, section_key))
                used.add(section_key)
        # A line whose font is significantly larger than the
        # default body size is a visual heading, even if its text
        # does not match a canonical section heading. We use it
        # to terminate the previous section's range at this
        # heading's position (e.g. when the artifact contains a
        # second table in a section that is NOT in the canonical
        # model — the verifier must not extend the canonical
        # section's scope across that second heading).
        elif line.max_font_size >= 13.0 and folded:
            # Find the most recent canonical-section start; insert
            # a ``__SEAM__`` marker just before this heading to
            # truncate the previous section.
            for i in range(len(starts) - 1, -1, -1):
                if starts[i][0] < idx:
                    starts.insert(i + 1, (idx, "__SEAM__"))
                    break
    resolved: dict[str, tuple[int, int]] = {}
    last_end = len(observation.all_lines)
    for start_idx, section_key in reversed(starts):
        if section_key == "__SEAM__":
            last_end = start_idx
            continue
        resolved[section_key] = (start_idx, last_end)
        last_end = start_idx
    return resolved


# ── Field-level binding (structured) ──────────────────────────────────────


def _split_metric_paragraph(text: str) -> tuple[str, str, str] | None:
    """Parse a renderer-emitted metric line.

    The real DocxRenderer / PdfRenderer emits metrics as
    ``"{label}: {display_value} {display_unit}"`` (whitespace-folded).
    Returns (label, value, unit) or None if the line does not match
    the expected pattern.
    """

    folded = _fold_whitespace(text)
    if ":" not in folded:
        return None
    label_part, rest_part = folded.split(":", 1)
    label = label_part.strip()
    rest = rest_part.strip()
    if not label or not rest:
        return None
    # The value and unit are whitespace-separated tokens at the end.
    # The value can contain digits + . + , and an optional sign; the
    # unit is the last token.
    tokens = rest.split(" ")
    if len(tokens) < 1:
        return None
    unit = tokens[-1]
    value = " ".join(tokens[:-1]).strip()
    if not value:
        return None
    return (label, value, unit)


def _split_number_paragraph(text: str) -> tuple[str, str] | None:
    """Parse a renderer-emitted number line.

    The real renderer emits a number line as either:
      - ``"{display_value} {display_unit}"`` (value + unit), or
      - ``"{display_value}"`` (value only, no unit).

    Returns (value, unit) where ``unit`` is empty string when the
    number has no unit. Returns None if the line is empty or the
    value token is not a recognizable number.

    The value MUST be a recognizable number (digits with optional
    decimal point, comma thousands separator, and sign). This
    prevents plain prose like "Heading 1" from being mistakenly
    classified as a number record.
    """

    folded = _fold_whitespace(text)
    if not folded:
        return None
    tokens = folded.split(" ")
    if not tokens:
        return None
    # The value MUST be the first token and MUST be a recognizable
    # number. The unit (if any) is the remaining tokens joined.
    value_token = tokens[0]
    if not _looks_like_number(value_token):
        return None
    unit = " ".join(tokens[1:]).strip()
    return (value_token, unit)


def _looks_like_number(token: str) -> bool:
    """Heuristic: does this token look like a number?

    Accepts integers, decimals, signed values, and thousands-separated
    values like "1,000" or "1,000.5". The renderer uses
    ``format_decimal`` which emits thousands separators in en-US.
    Returns False for plain prose like "Heading" or "总".
    """

    if not token:
        return False
    # Strip leading sign and thousands separators; require at least one digit.
    stripped = token.lstrip("+-")
    if not stripped:
        return False
    digits = stripped.replace(",", "").replace(".", "")
    return digits.isdigit() and any(ch.isdigit() for ch in stripped)


def _find_metric_binding(
    *,
    docx_observation: _DocxObservation | None,
    pdf_observation: _PdfObservation | None,
    section_key: str,
    section_scopes: Mapping[str, tuple[int, int]],
    expected_label: str,
    expected_value: str,
    expected_unit: str,
) -> _BindingResult:
    """Find the unique paragraph in the target section that binds the metric.

    The metric binding is based on the localized label, not on the
    expected value/unit. The expected value/unit are passed only for
    downstream comparison, NOT for candidate filtering. This prevents
    the "search artifact for expected value" pattern.

    Returns ``_BindingResult(observed, None, ())`` on success.
    Returns ``_BindingResult(None, "MISSING_FIELD_BINDING", ())`` when
    no candidate paragraph was found.
    Returns ``_BindingResult(None, "AMBIGUOUS_FIELD_BINDING", candidates)``
    when more than one candidate was found; the artifact-derived
    candidates are exposed for audit.
    Returns ``_BindingResult(None, "MISSING_SECTION", ())`` when the
    section's heading is not in the artifact.
    """

    if docx_observation is not None:
        if section_key not in section_scopes:
            return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())
        start, end = section_scopes[section_key]
        expected_label_folded = _fold_whitespace(expected_label)
        candidates: list[_ObservedNumericField] = []
        for idx in range(start, end):
            block = docx_observation.body_blocks[idx]
            if block.kind != "paragraph":
                continue
            # Heading paragraphs are skipped.
            if block.heading_level is not None:
                continue
            parsed = _split_metric_paragraph(block.text)
            if parsed is None:
                continue
            label, value, unit = parsed
            if _fold_whitespace(label) != expected_label_folded:
                continue
            candidates.append(
                _ObservedNumericField(
                    field_path="",  # populated by caller
                    section_key=section_key,
                    binding_kind=_BINDING_KIND_METRIC,
                    display_value=value,
                    display_unit=unit,
                    row_index=None,
                    column_index=None,
                )
            )
        if len(candidates) == 0:
            return _BindingResult(
                observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
            )
        if len(candidates) > 1:
            return _BindingResult(
                observed=None,
                failure_code="AMBIGUOUS_FIELD_BINDING",
                candidates=tuple(candidates),
            )
        return _BindingResult(observed=candidates[0], failure_code=None, candidates=())

    if pdf_observation is not None:
        if section_key not in section_scopes:
            return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())
        start, end = section_scopes[section_key]
        expected_label_folded = _fold_whitespace(expected_label)
        candidates = []
        for idx in range(start, end):
            line = pdf_observation.all_lines[idx]
            parsed = _split_metric_paragraph(line.text)
            if parsed is None:
                continue
            label, value, unit = parsed
            if _fold_whitespace(label) != expected_label_folded:
                continue
            candidates.append(
                _ObservedNumericField(
                    field_path="",
                    section_key=section_key,
                    binding_kind=_BINDING_KIND_METRIC,
                    display_value=value,
                    display_unit=unit,
                    row_index=None,
                    column_index=None,
                    page_number=line.page_number,
                )
            )
        if len(candidates) == 0:
            return _BindingResult(
                observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
            )
        if len(candidates) > 1:
            return _BindingResult(
                observed=None,
                failure_code="AMBIGUOUS_FIELD_BINDING",
                candidates=tuple(candidates),
            )
        return _BindingResult(observed=candidates[0], failure_code=None, candidates=())

    return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())


def _find_number_binding(
    *,
    docx_observation: _DocxObservation | None,
    pdf_observation: _PdfObservation | None,
    section_key: str,
    section_scopes: Mapping[str, tuple[int, int]],
) -> _BindingResult:
    """Find the number record in the target section by structural position.

    Per Corrective 5, the binding follows the renderer number
    contract: ``section heading → first non-empty, non-heading
    content record``. The helper walks the section's blocks/lines
    in document order, skips headings and empty records, and
    BINDS the FIRST record that parses as a number paragraph.

    Subsequent records (e.g. body prose that happens to contain a
    number) are NOT collected as candidates. This ensures that
    later-numbered body text does not cause false-fail (the
    earlier P1-3 round incorrectly returned
    ``AMBIGUOUS_FIELD_BINDING`` when the section had 2+ number-
    parseable lines).

    AMBIGUOUS_FIELD_BINDING is reserved for structural ambiguity
    (e.g. a section whose first parseable record is followed by a
    second structurally identical candidate) and is NOT raised
    here. The first record is always taken.

    Returns ``_BindingResult(observed, None, ())`` on success.
    Returns ``_BindingResult(None, "MISSING_FIELD_BINDING", ())``
    when the first non-empty, non-heading record cannot be
    parsed as a number. Returns ``_BindingResult(None,
    "MISSING_SECTION", ())`` when the section's heading is not
    in the artifact.
    """

    if docx_observation is not None:
        if section_key not in section_scopes:
            return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())
        start, end = section_scopes[section_key]
        for idx in range(start, end):
            block = docx_observation.body_blocks[idx]
            if block.kind != "paragraph":
                continue
            if block.heading_level is not None:
                continue
            text = _fold_whitespace(block.text)
            if not text:
                continue
            parsed = _split_number_paragraph(text)
            if parsed is None:
                # First non-empty, non-heading record is not a
                # number — fail closed.
                return _BindingResult(
                    observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
                )
            value, unit = parsed
            return _BindingResult(
                observed=_ObservedNumericField(
                    field_path="",
                    section_key=section_key,
                    binding_kind=_BINDING_KIND_NUMBER,
                    display_value=value,
                    display_unit=unit,
                    row_index=None,
                    column_index=None,
                ),
                failure_code=None,
                candidates=(),
            )
        # No non-empty, non-heading record at all.
        return _BindingResult(observed=None, failure_code="MISSING_FIELD_BINDING", candidates=())

    if pdf_observation is not None:
        if section_key not in section_scopes:
            return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())
        start, end = section_scopes[section_key]
        for idx in range(start, end):
            line = pdf_observation.all_lines[idx]
            # Skip visual headings: in PDF, the section heading
            # appears as a line with a larger font size.
            if line.max_font_size >= 13.0:
                continue
            text = _fold_whitespace(line.text)
            if not text:
                continue
            parsed = _split_number_paragraph(text)
            if parsed is None:
                return _BindingResult(
                    observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
                )
            value, unit = parsed
            return _BindingResult(
                observed=_ObservedNumericField(
                    field_path="",
                    section_key=section_key,
                    binding_kind=_BINDING_KIND_NUMBER,
                    display_value=value,
                    display_unit=unit,
                    row_index=None,
                    column_index=None,
                    page_number=line.page_number,
                ),
                failure_code=None,
                candidates=(),
            )
        return _BindingResult(observed=None, failure_code="MISSING_FIELD_BINDING", candidates=())

    return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())


def _strip_renderer_unit_wrapper(token: str) -> str:
    """Strip the renderer's single-layer outer parentheses used to render units.

    The real DOCX/PDF renderer wraps unit labels in a single pair of
    parentheses (e.g. ``"(kW(e))"``) so that parens are visible in the
    emitted text. The verifier strips exactly ONE outer pair of
    parentheses from the *rendered* unit token so it can be compared
    against the localized expected unit (which is the bare token
    ``"kW(e)"``). No inner parentheses, brackets, or aliases are
    rewritten; no fuzzy equivalence is applied.
    """
    folded = _fold_whitespace(token)
    if folded.startswith("(") and folded.endswith(")") and len(folded) >= 2:
        return folded[1:-1].strip()
    return folded


def _find_docx_table_candidate(
    *,
    section_blocks: tuple[_DocxBlock, ...],
    start: int,
    end: int,
    expected_headers: tuple[str, ...],
) -> tuple[_DocxBlock, ...]:
    """Return the DOCX table blocks within the section that match the
    localized ``expected_headers`` (folded-exact). 0 or 1+ candidates.
    """

    candidates: list[_DocxBlock] = []
    for idx in range(start, end):
        block = section_blocks[idx]
        if block.kind != "table" or block.cells is None:
            continue
        header_row = block.cells[0]
        if len(header_row) != len(expected_headers):
            continue
        folded_actual = tuple(_fold_whitespace(c) for c in header_row)
        if folded_actual == tuple(_fold_whitespace(h) for h in expected_headers):
            candidates.append(block)
    return tuple(candidates)


def _find_table_cell_binding_via_logical_table(
    *,
    pdf_logical_tables: tuple[_PdfLogicalTable, ...],
    section_key: str,
    table_section_key: str,
    row_index: int,
    column_index: int,
    expected_unit_codes: tuple[str, ...],
    expected_headers: tuple[str, ...],
    template_unit_row_enabled: bool,
) -> _BindingResult:
    """Resolve (row, column) using _PdfLogicalTable structural identity.

    Per P1-3 fourth corrective: when grid geometry is available, this
    path is the PRIMARY binding authority. Strategy:

      1. Filter logical tables by section_key (== table_section_key).
      2. Match each candidate by header folded-exact match.
      3. 0 matches -> USE_SECTION_TABLES_FALLBACK (caller should
         fall back to text-only).
      4. >=2 matches -> AMBIGUOUS_FIELD_BINDING.
      5. Use the unique logical_table.data_rows[row_index] as the data
         row. The unit row, if present, is the _PdfLogicalRow with
         ``row_kind == 'unit'`` in the FIRST segment (since each
         segment carries its own unit row). If no logical_table
         unit_row exists or row_kind != 'unit', observed_unit = ''.
    """
    if section_key != table_section_key:
        return _BindingResult(observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=())
    candidates = [
        t
        for t in pdf_logical_tables
        if t.section_key == section_key
        and tuple(_fold_whitespace(c.text) for c in t.header.cells)
        == tuple(_fold_whitespace(h) for h in expected_headers)
    ]
    if not candidates:
        return _BindingResult(
            observed=None,
            failure_code="USE_SECTION_TABLES_FALLBACK",
            candidates=(),
        )
    if len(candidates) > 1:
        return _BindingResult(
            observed=None,
            failure_code="AMBIGUOUS_FIELD_BINDING",
            candidates=(),
        )
    target = candidates[0]
    if row_index < 0 or row_index >= len(target.data_rows):
        return _BindingResult(observed=None, failure_code="TABLE_ROW_MISMATCH", candidates=())
    data_row = target.data_rows[row_index]
    cells = data_row.cells
    if column_index < 0 or column_index >= len(cells):
        return _BindingResult(observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=())
    cell = cells[column_index]
    cell_value = cell.text
    observed_unit = ""
    if target.segments:
        unit_row = target.segments[0].unit_row
        if (
            unit_row is not None
            and unit_row.row_kind == "unit"
            and column_index < len(unit_row.cells)
        ):
            observed_unit = _strip_renderer_unit_wrapper(unit_row.cells[column_index].text)
    return _BindingResult(
        observed=_ObservedNumericField(
            field_path="",
            section_key=section_key,
            binding_kind=_BINDING_KIND_TABLE_CELL,
            display_value=cell_value,
            display_unit=observed_unit,
            row_index=row_index,
            column_index=column_index,
        ),
        failure_code=None,
        candidates=(),
    )


def _find_table_cell_binding(
    *,
    docx_observation: _DocxObservation | None,
    docx_resolved_scopes: Mapping[str, tuple[int, int]],
    pdf_section_tables: tuple[_PdfSectionTable, ...],
    pdf_logical_tables: tuple[_PdfLogicalTable, ...] = (),
    section_key: str,
    table_section_key: str,
    row_index: int,
    column_index: int,
    expected_unit_codes: tuple[str, ...],
    expected_headers: tuple[str, ...],
    template_unit_row_enabled: bool,
    num_data_rows: int = 1,
) -> _BindingResult:
    """Find the (row, column) cell in the table located in the target section.

    Per Corrective 1+2+3+4, the binding:

      - Selects the table by structural identity (localized headers
        folded-exact match against the artifact's header row), NOT
        by the first table or by expected numeric value.
      - Determines unit-row presence with renderer parity:
        ``template_unit_row_enabled and any(expected_unit_codes)``.
      - Reads the OBSERVED unit from the artifact's unit row (after
        stripping the renderer's single-layer outer paren wrapper).
        The unit is NOT copied from the localized expected unit.
      - When multiple table candidates match the localized headers,
        the binding returns ``AMBIGUOUS_FIELD_BINDING`` (fail-closed).
    """

    # P1-3 fourth corrective: when ANY logical table exists for this
    # section, use the LOGICAL-TABLE binding path as the PRIMARY
    # authority. The text-only section-tables fallback is only used
    # when (a) there are no logical tables OR (b) logical-table
    # binding returns USE_SECTION_TABLES_FALLBACK (no header match).
    if pdf_logical_tables and pdf_section_tables is not None:
        result = _find_table_cell_binding_via_logical_table(
            pdf_logical_tables=pdf_logical_tables,
            section_key=section_key,
            table_section_key=table_section_key,
            row_index=row_index,
            column_index=column_index,
            expected_unit_codes=expected_unit_codes,
            expected_headers=expected_headers,
            template_unit_row_enabled=template_unit_row_enabled,
        )
        if result.failure_code != "USE_SECTION_TABLES_FALLBACK":
            return result
    if docx_observation is not None:
        if section_key != table_section_key:
            return _BindingResult(
                observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=()
            )
        if section_key not in docx_resolved_scopes:
            return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())
        start, end = docx_resolved_scopes[section_key]
        candidate_tables = _find_docx_table_candidate(
            section_blocks=docx_observation.body_blocks,
            start=start,
            end=end,
            expected_headers=expected_headers,
        )
        if not candidate_tables:
            return _BindingResult(
                observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
            )
        if len(candidate_tables) > 1:
            return _BindingResult(
                observed=None,
                failure_code="AMBIGUOUS_FIELD_BINDING",
                candidates=(),
            )
        cells = candidate_tables[0].cells
        assert cells is not None
        # Renderer parity: the unit row is present in the
        # artifact when the template enables it AND the
        # canonical has at least one non-empty expected unit
        # AND the artifact has an extra row (i.e. ``len(cells)
        # > 1 + num_data_rows``). The data row offset follows
        # the canonical's expected structure:
        #   - unit row expected AND present → data at
        #     cells[2 + row_index]
        #   - unit row not expected → data at
        #     cells[1 + row_index]
        #   - unit row expected BUT artifact missing the unit
        #     row → TABLE_ROW_MISMATCH (canonical structure
        #     wins)
        # For the symmetric comparison (Corrective 1), if the
        # canonical has no expected unit BUT the artifact has
        # an extra row, the extra row is still read as a unit
        # row and compared (fail-closed on unexpected unit).
        has_unit_row_expected = template_unit_row_enabled and any(expected_unit_codes)
        artifact_has_extra_row = len(cells) > 1 + num_data_rows
        if has_unit_row_expected:
            # Unit row is expected. The data row index is
            # based on the canonical's expectation.
            data_row_idx = 1 + 1 + row_index
            unit_row_present = artifact_has_extra_row
        elif artifact_has_extra_row:
            # No expected unit, but artifact has an extra row
            # (synthetic case or renderer bug). Read the extra
            # row as the unit row (Corrective 1).
            data_row_idx = 1 + 1 + row_index
            unit_row_present = True
        else:
            # No unit row in the artifact.
            data_row_idx = 1 + row_index
            unit_row_present = False
        if data_row_idx < 0 or data_row_idx >= len(cells):
            return _BindingResult(observed=None, failure_code="TABLE_ROW_MISMATCH", candidates=())
        if column_index < 0 or column_index >= len(cells[data_row_idx]):
            return _BindingResult(
                observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=()
            )
        cell_value = cells[data_row_idx][column_index]
        # Read the observed unit from the artifact's unit row.
        if unit_row_present and len(cells) > 1:
            raw_unit_token = cells[1][column_index] if column_index < len(cells[1]) else ""
            observed_unit = _strip_renderer_unit_wrapper(raw_unit_token)
        else:
            observed_unit = ""
        return _BindingResult(
            observed=_ObservedNumericField(
                field_path="",
                section_key=section_key,
                binding_kind=_BINDING_KIND_TABLE_CELL,
                display_value=cell_value,
                display_unit=observed_unit,
                row_index=row_index,
                column_index=column_index,
            ),
            failure_code=None,
            candidates=(),
        )

    if pdf_section_tables is not None:
        if section_key != table_section_key:
            return _BindingResult(
                observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=()
            )
        # The PDF section-local table's first row is the HEADER
        # row (anchored by ``_build_section_local_tables``).
        # Subsequent rows are body rows. Use column_centers to
        # map a body-line's x to its column index.
        section_tables = tuple(tbl for tbl in pdf_section_tables if tbl.section_key == section_key)
        if not section_tables:
            return _BindingResult(
                observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
            )
        # Each section-local table's first row is the header.
        # The header row text must match expected_headers
        # (folded-exact). 0 → MISSING; 1 → bind; >1 → AMBIGUOUS.
        matched: list[_PdfSectionTable] = []
        for tbl in section_tables:
            if not tbl.rows:
                continue
            header_row = tbl.rows[0]
            if len(header_row) != len(expected_headers):
                continue
            folded_row = tuple(_fold_whitespace(ln.text) for ln in header_row)
            if folded_row == tuple(_fold_whitespace(h) for h in expected_headers):
                matched.append(tbl)
        if not matched:
            return _BindingResult(
                observed=None, failure_code="MISSING_FIELD_BINDING", candidates=()
            )
        if len(matched) > 1:
            return _BindingResult(
                observed=None,
                failure_code="AMBIGUOUS_FIELD_BINDING",
                candidates=(),
            )
        target_table = matched[0]
        # Renderer parity: unit row visible iff template enables it
        # AND any localized expected unit is non-empty.
        has_unit_row_expected = template_unit_row_enabled and any(expected_unit_codes)
        # Body rows start at index 1 (after the header). The
        # ``expected`` unit row, if present, is body row 0;
        # data rows start at body row 1.
        body_rows = target_table.rows[1:]
        if not body_rows:
            return _BindingResult(observed=None, failure_code="TABLE_ROW_MISMATCH", candidates=())
        # Per Corrective 3, the unit row MAY be physically absent
        # from the artifact (e.g. when the unit row is empty in
        # the renderer and emits no spans). The unit row's
        # presence in the artifact is detected by row count: if
        # body has 2+ rows and we expect a unit row, the first
        # body row is the unit row. If body has 1 row, the unit
        # row is empty / absent.
        cell_value = ""
        observed_unit = ""
        # P1-3 fourth corrective: unit-row presence is determined
        # STRUCTURALLY. A body row is the unit row iff at least one
        # of its non-empty cells, located in a non-leftmost column
        # band, is a renderer unit token. The leftmost band is
        # treated as a label column and NOT required to be a unit
        # token (it may contain short scheme-name text in some
        # renderers, or a label like \"单位\" in others).
        # The ``has_unit_row_expected`` flag ONLY affects the post-
        # bind unit comparison; it MUST NOT influence row 0's
        # structural classification.
        body0 = body_rows[0]
        body0_non_empty = [ln for ln in body0 if _fold_whitespace(ln.text)]
        # column band of each non-empty line
        body0_band_to_text: list[tuple[int, str]] = []
        for ln in body0_non_empty:
            band = _body_line_band_for(ln, target_table.column_centers)
            if band is None:
                continue
            body0_band_to_text.append((band, ln.text))
        non_leftmost_texts = [t for b, t in body0_band_to_text if b > 0]
        body0_is_unit_structurally = any(_is_renderer_unit_token(t) for t in non_leftmost_texts)
        if body0_is_unit_structurally:
            unit_row_idx: int | None = 0
            data_row_idx = 1 + row_index
        else:
            unit_row_idx = None
            data_row_idx = 0 + row_index
        if data_row_idx < 0 or data_row_idx >= len(body_rows):
            return _BindingResult(observed=None, failure_code="TABLE_ROW_MISMATCH", candidates=())
        data_row = body_rows[data_row_idx]
        # Map data_row's lines to column indices via column_centers.
        if not target_table.column_centers:
            return _BindingResult(
                observed=None, failure_code="TABLE_COLUMN_MISMATCH", candidates=()
            )
        if len(target_table.column_centers) < 2:
            half_band = 100.0
        else:
            sorted_centers = sorted(target_table.column_centers)
            half_band = max(
                (sorted_centers[i + 1] - sorted_centers[i]) / 2
                for i in range(len(sorted_centers) - 1)
            )
        # Pick the line whose x-center is closest to the target
        # column's center.
        target_center = target_table.column_centers[
            min(column_index, len(target_table.column_centers) - 1)
        ]
        cell_line = None
        for ln in data_row:
            ln_center = (ln.bbox[0] + ln.bbox[2]) / 2
            if abs(ln_center - target_center) <= half_band and (
                cell_line is None
                or abs(ln_center - target_center)
                < abs((cell_line.bbox[0] + cell_line.bbox[2]) / 2 - target_center)
            ):
                cell_line = ln
        if cell_line is not None:
            cell_value = cell_line.text
        # Read the observed unit from the artifact's unit row.
        if unit_row_idx is not None and unit_row_idx < len(body_rows):
            unit_row = body_rows[unit_row_idx]
            unit_line = None
            for ln in unit_row:
                ln_center = (ln.bbox[0] + ln.bbox[2]) / 2
                if abs(ln_center - target_center) <= half_band and (
                    unit_line is None
                    or abs(ln_center - target_center)
                    < abs((unit_line.bbox[0] + unit_line.bbox[2]) / 2 - target_center)
                ):
                    unit_line = ln
            if unit_line is not None:
                observed_unit = _strip_renderer_unit_wrapper(unit_line.text)
        return _BindingResult(
            observed=_ObservedNumericField(
                field_path="",
                section_key=section_key,
                binding_kind=_BINDING_KIND_TABLE_CELL,
                display_value=cell_value,
                display_unit=observed_unit,
                row_index=row_index,
                column_index=column_index,
                page_number=target_table.page_number,
            ),
            failure_code=None,
            candidates=(),
        )

    return _BindingResult(observed=None, failure_code="MISSING_SECTION", candidates=())


def _compare_field(
    *,
    observed: _ObservedNumericField,
    expected_value: str,
    expected_unit: str,
) -> str | None:
    """Compare an observed field against the localized expected.

    Returns a failure code (one of ``"VALUE_MISMATCH"``,
    ``"UNIT_MISSING"``, ``"UNIT_MISMATCH"``) or None on success.

    Per Corrective 1's unit-integrity rule, unit comparison is
    strictly symmetric:

      - ``expected==""`` and ``observed==""`` → unit OK
      - ``expected==""`` and ``observed!=""`` → ``UNIT_MISMATCH``
        (the artifact emitted a unit the canonical did not expect)
      - ``expected!=""`` and ``observed==""`` → ``UNIT_MISSING``
        (the canonical expects a unit, the artifact has none)
      - ``expected!=""`` and ``observed!=expected`` →
        ``UNIT_MISMATCH`` (folded-whitespace only)
      - units match → continue to value comparison.

    Whitespace folding is the ONLY allowed transformation. No
    fuzzy numeric tolerance, no unit aliasing, no parenthesized
    alias equivalence.
    """

    expected_u_folded = _fold_whitespace(expected_unit)
    observed_u_folded = _fold_whitespace(observed.display_unit)
    if expected_u_folded or observed_u_folded:
        if expected_u_folded and not observed_u_folded:
            return "UNIT_MISSING"
        if observed_u_folded and not expected_u_folded:
            return "UNIT_MISMATCH"
        if observed_u_folded != expected_u_folded:
            return "UNIT_MISMATCH"
    # Value comparison (whitespace-folded).
    if not _strings_equal_folded(observed.display_value, expected_value):
        return "VALUE_MISMATCH"
    return None


def _extract_text(fmt: ExportFormat, data: bytes) -> str:  # noqa: ARG001
    """Retained for backward compatibility with P1-1/P1-2 contracts.

    The P1-3 round replaces the global-text verifier with a
    structured observation + binding layer. This function is kept
    for any test/utility that still needs a flattened text view
    (e.g. heading-presence checks, or legacy helpers). Numeric
    semantic verification MUST go through ``_semantic_checks`` and
    the structured binding functions, not through this helper.
    """

    if fmt is ExportFormat.DOCX:
        document = Document(io.BytesIO(data))
        parts: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
        return "\n".join(parts)
    if fmt is ExportFormat.PDF:
        with fitz.open(stream=data, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)
    _fail("UNSUPPORTED_FORMAT", f"Unsupported report format: {fmt.value}")
    raise AssertionError("unreachable")


def _canonical_metrics(
    model: CanonicalReportRenderModel,
) -> tuple[CanonicalRenderMetric | CanonicalRenderTableCell, ...]:
    values: list[CanonicalRenderMetric | CanonicalRenderTableCell] = []
    for section in model.sections:
        if section.number is not None:
            values.append(section.number)
        values.extend(section.metrics)
        if section.table is not None:
            for row in section.table.rows:
                values.extend(
                    cell
                    for cell in row
                    if isinstance(cell.raw_value, int) or hasattr(cell.raw_value, "as_tuple")
                )
    return tuple(values)


def _managed_paths() -> tuple[Path, ...]:
    paths: list[Path] = [Path("pilot-run.json"), Path("pilot-summary.json")]
    for locale, fmt in _RENDER_MATRIX:
        base = Path("artifacts") / locale.value / fmt.value
        paths.extend(
            (
                base / f"report.{fmt.value}",
                base / "artifact-metadata.json",
                base / "semantic-checks.json",
            )
        )
    return tuple(paths)


def _verify_artifact_binding(
    *,
    artifact: Any,
    report_id: str,
    revision: Any,
    locale: ReportLocale,
    fmt: ExportFormat,
    template: Any,
) -> None:
    if artifact.status is not ArtifactStatus.COMPLETED:
        _fail(
            "ARTIFACT_NOT_COMPLETED",
            "Rendered artifact is not completed.",
            artifact_id=artifact.id,
            status=artifact.status.value,
        )
    if (
        artifact.report_id != report_id
        or artifact.report_revision_id != revision.id
        or artifact.revision_number != revision.revision_number
    ):
        _fail(
            "REPORT_REVISION_MISMATCH",
            "Artifact does not bind to the one pilot report revision.",
            artifact_id=artifact.id,
        )
    if artifact.format is not fmt:
        _fail("ARTIFACT_METADATA_MISMATCH", "Artifact format mismatch.")
    if artifact.locale is not locale:
        _fail("LOCALE_BINDING_MISMATCH", "Artifact locale mismatch.")
    if artifact.template_locale is not locale:
        _fail("TEMPLATE_LOCALE_MISMATCH", "Template locale mismatch.")
    if artifact.source_content_hash != revision.content_hash:
        _fail("SOURCE_CONTENT_HASH_MISMATCH", "Artifact source hash mismatch.")
    manifest = artifact.render_manifest_json
    if manifest.get("render_mode") != "draft":
        _fail("UNSUPPORTED_RENDER_MODE", "Pilot artifact is not a draft render.")
    if manifest.get("template_content_hash") != template.template_content_hash:
        _fail("TEMPLATE_PROVENANCE_MISMATCH", "Template content hash mismatch.")
    catalog = get_catalog(locale)
    catalog_hash = compute_catalog_content_hash(locale)
    if not artifact.translation_catalog_version:
        _fail(
            "TRANSLATION_CATALOG_IDENTITY_MISSING",
            "Translation catalog version is empty.",
        )
    if (
        artifact.translation_catalog_version != catalog.version
        or artifact.translation_catalog_content_hash != catalog_hash
    ):
        _fail(
            "TRANSLATION_CATALOG_IDENTITY_MISMATCH",
            "Translation catalog identity mismatch.",
        )
    if not _is_sha256(artifact.localized_template_content_hash):
        _fail(
            "LOCALIZED_TEMPLATE_HASH_MISSING",
            "Localized template content hash is missing or malformed.",
        )


def _lookup_template_table_unit_row(
    *,
    template_manifest_json: dict[str, Any] | None,
    table_key: str | None,
) -> bool:
    """Return the template's canonical ``tables[table_key].unit_row`` bool.

    The production authority is the canonical ``TemplateManifest``:

        TemplateManifest.from_manifest_json(template_manifest_json)
        .tables[table_key].unit_row

    This helper MUST NOT maintain a second manifest schema parser.
    It delegates to the production ``TemplateManifest`` and reads
    ``unit_row`` from ``manifest.tables[table_key]``. If the bool
    cannot be resolved (unknown ``table_key``, missing field,
    schema mismatch), the renderer default (``True``) is used.

    ``table_key`` is the canonical table identity (NOT the
    section_key). The call chain is responsible for propagating the
    canonical table_key from the localized section.table.canonical
    down to this helper.
    """

    if not isinstance(template_manifest_json, dict):
        return True
    if not table_key:
        return True
    try:
        from cold_storage.modules.reports.domain.render_model import (
            TemplateManifest as _TemplateManifest,
        )

        manifest = _TemplateManifest.from_manifest_json(template_manifest_json)
    except Exception:  # pragma: no cover -- defensive: schema mismatch
        return True
    config = manifest.tables.get(table_key)
    if config is None:
        return True
    return bool(config.unit_row)


# ── Heading-scope-based required-section authority ────────────────────────


def _build_missing_sections_from_scopes(
    *,
    localized_sections: tuple[Any, ...],
    resolved_scopes: Mapping[str, tuple[int, int]],
) -> list[str]:
    """Build the ``missing_sections`` list from structural section scopes.

    Per Corrective 6, the required-section authority is the
    section-scope map (whether the localized heading was found in
    the artifact as a structural section divider), NOT a
    flattened-text substring check. A section is ``missing`` if
    its ``section_key`` is absent from ``resolved_scopes``. The
    returned list contains the localized titles for the missing
    sections (so the schema-level field still carries the
    human-readable heading text).
    """

    missing_titles: list[str] = []
    for section in localized_sections:
        if section.section_key not in resolved_scopes:
            missing_titles.append(section.title)
    return missing_titles


def _build_observed_localized_headings(
    *,
    localized_sections: tuple[Any, ...],
    resolved_scopes: Mapping[str, tuple[int, int]],
) -> list[str]:
    """Build the ``observed_localized_headings`` list from structural scopes.

    Per Corrective 6, this list now reflects which localized
    headings were structurally located as section dividers (i.e.
    whose ``section_key`` is in ``resolved_scopes``). The
    flattened-text substring heuristic is no longer used to
    determine section presence.
    """

    return [
        section.title for section in localized_sections if section.section_key in resolved_scopes
    ]


def _semantic_checks(
    *,
    canonical_model: CanonicalReportRenderModel,
    template: Any,
    locale: ReportLocale,
    fmt: ExportFormat,
    artifact_bytes: bytes,
) -> dict[str, Any]:
    """Run structured, field-bound semantic verification on a downloaded artifact.

    The P1-3 fix replaces the previous global substring search
    (metric.display_value in extracted_text) with a structured
    observation + binding layer. The function:

      1. Localizes the canonical model via ``localize_render_model``.
      2. Observes the downloaded artifact (DOCX/PDF bytes) into
         structured blocks/lines + section scopes.
      3. For each canonical numeric field (metric, number, table
         cell), finds the unique structural binding at the
         section/field/row/column position.
      4. Compares the observed value/unit against the localized
         expected value/unit (whitespace-folded only).
      5. Records any mismatch in ``numeric_mismatches`` (value/unit
         error), ``missing_units`` (unit absent/mismatched), and
         sets ``semantic_result`` to ``FAIL`` on any failure.

    The function NEVER falls back to global substring search. The
    observed value/unit come from the downloaded artifact bytes,
    not from the localized model. On binding failure, the
    ``display_value`` and ``display_unit`` fields are empty
    strings (NOT copied from the expected/canonical model) so the
    audit consumer can see the failure is real.
    """

    template_manifest_json = template.manifest_json if hasattr(template, "manifest_json") else None

    localized = localize_render_model(
        canonical_model,
        locale=locale,
        template_manifest_json=template_manifest_json,
        format=fmt.value,
    )

    # Per Corrective 6, the required-section authority is the
    # structural section-scope map (resolved_scopes), NOT a
    # flattened-text substring check. The flattened_text is
    # computed only for legacy diagnostic consumers; the
    # section-presence gate is now strictly scope-based.
    _ = _extract_text(fmt, artifact_bytes)

    # Section scopes from the localized model.
    section_scopes_spec = _build_section_scopes(localized_sections=localized.sections)

    # Structured observation.
    docx_observation: _DocxObservation | None = None
    pdf_observation: _PdfObservation | None = None
    pdf_section_tables: tuple[_PdfSectionTable, ...] = ()
    pdf_logical_tables: tuple[_PdfLogicalTable, ...] = ()
    if fmt is ExportFormat.DOCX:
        docx_observation = _observe_docx(artifact_bytes)
        resolved_scopes = _resolve_docx_section_scopes(
            observation=docx_observation, section_scopes=section_scopes_spec
        )
    elif fmt is ExportFormat.PDF:
        pdf_observation = _observe_pdf(artifact_bytes)
        resolved_scopes = _resolve_pdf_section_scopes(
            observation=pdf_observation, section_scopes=section_scopes_spec
        )
        # Per Corrective 3+4, tables are reconstructed per section
        # using the localized table headers (structural identity),
        # not page-global y-cluster heuristics. The header map is
        # section_key -> (expected_header_text_by_column).
        section_table_headers: dict[str, tuple[str, ...]] = {
            section.section_key: section.table.headers
            for section in localized.sections
            if section.table is not None
        }
        pdf_section_tables = _build_section_local_tables(
            pdf_observation=pdf_observation,
            section_scopes=resolved_scopes,
            section_table_headers=section_table_headers,
        )
        # Per Correctives 2/3/4/5, ALSO build grid-geometry-based
        # ``_PdfLogicalTable``s when the artifact has vector drawing
        # segments (real renderer output). When grid geometry is
        # available, this path produces cross-page-merged logical
        # tables (Corrective 4) and structural unit-row detection
        # (Corrective 3) and wrapped-cell folding (Corrective 5).
        # The grid-based logical tables are STORED for diagnostic
        # purposes; the cell-binding still uses
        # ``pdf_section_tables`` as the primary path for
        # backward-compat with synthetic PDFs that have no grid
        # lines. When grid geometry IS available, the verifier's
        # cell-binding path finds the unique header match by
        # structural identity — see ``_find_table_cell_binding``.
        if pdf_observation.grid_segments:
            for section_key, scope_range in resolved_scopes.items():
                expected = section_table_headers.get(section_key, ())
                if not expected:
                    continue
                tables = _build_logical_tables_for_section(
                    pdf_observation=pdf_observation,
                    section_key=section_key,
                    section_line_range=scope_range,
                    expected_headers=expected,
                )
                pdf_logical_tables = pdf_logical_tables + tables
        # Store on observation for diagnostic + downstream binding.
        pdf_observation = _PdfObservation(
            all_lines=pdf_observation.all_lines,
            section_scopes=pdf_observation.section_scopes,
            text_spans=pdf_observation.text_spans,
            grid_segments=pdf_observation.grid_segments,
        )
    else:
        _fail("UNSUPPORTED_FORMAT", f"Unsupported report format: {fmt.value}")
        raise AssertionError("unreachable")

    canonical_fields: list[dict[str, str]] = []
    observed_fields: list[dict[str, Any]] = []
    missing_units: list[str] = []
    numeric_mismatches: list[str] = []

    def _record_failure(
        *,
        field_path: str,
        section_key: str,
        binding_kind: str,
        failure_code: str,
        candidates: tuple[_ObservedNumericField, ...],
        row_index: int | None,
        column_index: int | None,
    ) -> None:
        """Append a binding-failure observed record with NO expected copy.

        Per Corrective 1, on binding failure the observed record's
        ``display_value`` and ``display_unit`` MUST be empty strings
        (not copied from the expected/canonical model). For
        AMBIGUOUS, the artifact-derived candidate observations are
        exposed for audit (so the failure is not silent).
        """
        record: dict[str, Any] = {
            "field_path": field_path,
            "section_key": section_key,
            "binding_kind": binding_kind,
            "display_value": "",
            "display_unit": "",
            "row_index": row_index,
            "column_index": column_index,
            "page_number": None,
            "binding_status": failure_code,
        }
        if candidates:
            record["candidate_count"] = len(candidates)
            record["candidate_values"] = [c.display_value for c in candidates]
            record["candidate_units"] = [c.display_unit for c in candidates]
            record["candidate_locations"] = [
                {
                    "row_index": c.row_index,
                    "column_index": c.column_index,
                    "page_number": c.page_number,
                }
                for c in candidates
            ]
        observed_fields.append(record)
        if failure_code in ("UNIT_MISSING", "UNIT_MISMATCH"):
            missing_units.append(field_path)
        else:
            numeric_mismatches.append(field_path)

    def _record_bound(
        *,
        field_path: str,
        observed: _ObservedNumericField,
    ) -> None:
        """Append a successfully-bound observed record with artifact value/unit."""

        observed_fields.append(
            {
                "field_path": field_path,
                "section_key": observed.section_key,
                "binding_kind": observed.binding_kind,
                "display_value": observed.display_value,
                "display_unit": observed.display_unit,
                "row_index": observed.row_index,
                "column_index": observed.column_index,
                "page_number": observed.page_number,
                "binding_status": "BOUND",
            }
        )

    def _inspect_metric(
        metric: Any,
        *,
        section_key: str,
        binding_label: str,
    ) -> None:
        canonical_fields.append(
            {
                "field_path": metric.canonical.field_path,
                "raw_value": str(metric.canonical.raw_value),
                "unit_code": metric.canonical.unit_code,
            }
        )
        result = _find_metric_binding(
            docx_observation=docx_observation,
            pdf_observation=pdf_observation,
            section_key=section_key,
            section_scopes=resolved_scopes,
            expected_label=binding_label,
            expected_value=metric.display_value,
            expected_unit=metric.display_unit,
        )
        if result.failure_code is not None:
            _record_failure(
                field_path=metric.canonical.field_path,
                section_key=section_key,
                binding_kind=_BINDING_KIND_METRIC,
                failure_code=result.failure_code,
                candidates=result.candidates,
                row_index=None,
                column_index=None,
            )
            return
        if result.observed is None:  # pragma: no cover
            numeric_mismatches.append(metric.canonical.field_path)
            return
        _record_bound(field_path=metric.canonical.field_path, observed=result.observed)
        cmp = _compare_field(
            observed=result.observed,
            expected_value=metric.display_value,
            expected_unit=metric.display_unit,
        )
        if cmp == "UNIT_MISSING" or cmp == "UNIT_MISMATCH":
            missing_units.append(metric.canonical.field_path)
        elif cmp is not None:
            numeric_mismatches.append(metric.canonical.field_path)

    def _inspect_number(
        number_metric: Any,
        *,
        section_key: str,
    ) -> None:
        canonical_fields.append(
            {
                "field_path": number_metric.canonical.field_path,
                "raw_value": str(number_metric.canonical.raw_value),
                "unit_code": number_metric.canonical.unit_code,
            }
        )
        # Corrective 2: number binding is by structural position; the
        # expected value/unit are passed in only for the comparison
        # step, NOT for candidate filtering. The helper no longer
        # accepts expected_value/expected_unit parameters.
        result = _find_number_binding(
            docx_observation=docx_observation,
            pdf_observation=pdf_observation,
            section_key=section_key,
            section_scopes=resolved_scopes,
        )
        if result.failure_code is not None:
            _record_failure(
                field_path=number_metric.canonical.field_path,
                section_key=section_key,
                binding_kind=_BINDING_KIND_NUMBER,
                failure_code=result.failure_code,
                candidates=result.candidates,
                row_index=None,
                column_index=None,
            )
            return
        if result.observed is None:  # pragma: no cover
            numeric_mismatches.append(number_metric.canonical.field_path)
            return
        _record_bound(field_path=number_metric.canonical.field_path, observed=result.observed)
        cmp = _compare_field(
            observed=result.observed,
            expected_value=number_metric.display_value,
            expected_unit=number_metric.display_unit,
        )
        if cmp == "UNIT_MISSING" or cmp == "UNIT_MISMATCH":
            missing_units.append(number_metric.canonical.field_path)
        elif cmp is not None:
            numeric_mismatches.append(number_metric.canonical.field_path)

    def _inspect_table_cell(
        cell: CanonicalRenderTableCell,
        localized_cell: Any,
        *,
        section_key: str,
        table_section_key: str,
        row_index: int,
        column_index: int,
        expected_unit_codes: tuple[str, ...],
        expected_headers: tuple[str, ...],
        template_unit_row_enabled: bool,
        num_data_rows: int = 1,
    ) -> None:
        canonical_fields.append(
            {
                "field_path": cell.field_path,
                "raw_value": str(cell.raw_value),
                "unit_code": cell.unit_code,
            }
        )
        result = _find_table_cell_binding(
            docx_observation=docx_observation,
            docx_resolved_scopes=resolved_scopes,
            pdf_section_tables=pdf_section_tables,
            pdf_logical_tables=pdf_logical_tables,
            section_key=section_key,
            table_section_key=table_section_key,
            row_index=row_index,
            column_index=column_index,
            expected_unit_codes=expected_unit_codes,
            expected_headers=expected_headers,
            template_unit_row_enabled=template_unit_row_enabled,
            num_data_rows=num_data_rows,
        )
        if result.failure_code is not None:
            _record_failure(
                field_path=cell.field_path,
                section_key=section_key,
                binding_kind=_BINDING_KIND_TABLE_CELL,
                failure_code=result.failure_code,
                candidates=result.candidates,
                row_index=row_index,
                column_index=column_index,
            )
            return
        if result.observed is None:  # pragma: no cover
            numeric_mismatches.append(cell.field_path)
            return
        _record_bound(field_path=cell.field_path, observed=result.observed)
        # Per Corrective 3's "expected authority" rule: the expected
        # value/unit come from the LOCALIZED render model directly
        # (``localized_cell.display_value`` /
        # ``localized_table.unit_row[column_index]``), NOT from a
        # second ``format_decimal(cell.raw_value, locale)`` reformat.
        # No table expected reformatting is performed in this module.
        expected_dv = localized_cell.display_value
        expected_du = (
            expected_unit_codes[column_index] if column_index < len(expected_unit_codes) else ""
        )
        cmp = _compare_field(
            observed=result.observed,
            expected_value=expected_dv,
            expected_unit=expected_du,
        )
        if cmp == "UNIT_MISSING" or cmp == "UNIT_MISMATCH":
            missing_units.append(cell.field_path)
        elif cmp is not None:
            numeric_mismatches.append(cell.field_path)

    for section in localized.sections:
        for metric in section.metrics:
            _inspect_metric(metric, section_key=section.section_key, binding_label=metric.label)
        if section.number is not None:
            _inspect_number(section.number, section_key=section.section_key)
        if section.table is not None:
            # The expected unit_codes come from the localized
            # table's unit_row (already formatted by
            # ``localize_render_model``). No reformatting here.
            localized_unit_codes: tuple[str, ...] = section.table.unit_row
            localized_headers: tuple[str, ...] = section.table.headers
            # Renderer parity: look up the template's
            # ``table.unit_row`` bool for this section's canonical
            # table. The production authority is
            # ``TemplateManifest.from_manifest_json(template_manifest_json)
            # .tables[table_key].unit_row``. The canonical
            # table_key comes from the localized
            # ``section.table.canonical.table_key`` (NOT from
            # ``section_key``). The default is ``True`` (matches
            # the renderer's default).
            _table_key: str | None = None
            if section.table.canonical is not None:
                _table_key = section.table.canonical.table_key or None
            template_unit_row_enabled = _lookup_template_table_unit_row(
                template_manifest_json=template_manifest_json,
                table_key=_table_key,
            )
            for row_idx, row in enumerate(section.table.rows):
                for col_idx, cell in enumerate(row):
                    raw = cell.canonical.raw_value
                    if isinstance(raw, (int, Decimal)) or hasattr(raw, "as_tuple"):
                        _inspect_table_cell(
                            cell.canonical,
                            cell,
                            section_key=section.section_key,
                            table_section_key=section.section_key,
                            row_index=row_idx,
                            column_index=col_idx,
                            expected_unit_codes=localized_unit_codes,
                            expected_headers=localized_headers,
                            template_unit_row_enabled=template_unit_row_enabled,
                            num_data_rows=len(section.table.rows),
                        )

    sections_ok = not _build_missing_sections_from_scopes(
        localized_sections=localized.sections, resolved_scopes=resolved_scopes
    )
    units_ok = not missing_units
    mismatches_ok = not numeric_mismatches
    result = "PASS" if sections_ok and units_ok and mismatches_ok else "FAIL"
    return {
        "schema_version": PILOT_RESULT_SCHEMA_VERSION,
        "locale": locale.value,
        "format": fmt.value,
        "canonical_section_keys": [section.section_key for section in canonical_model.sections],
        "required_heading_keys": [
            f"section.{section.section_key}" for section in canonical_model.sections
        ],
        "observed_localized_headings": _build_observed_localized_headings(
            localized_sections=localized.sections, resolved_scopes=resolved_scopes
        ),
        "canonical_numeric_fields": canonical_fields,
        "observed_numeric_fields": observed_fields,
        "missing_sections": _build_missing_sections_from_scopes(
            localized_sections=localized.sections, resolved_scopes=resolved_scopes
        ),
        "missing_units": sorted(set(missing_units)),
        "numeric_mismatches": sorted(set(numeric_mismatches)),
        "semantic_result": result,
    }


def verify_multilingual_report_pilot(
    *,
    report_service: Any,
    render_service: Any,
    template_repository: Any,
    project_id: str,
    project_version_id: str,
    source_commit_sha: str,
    source_manifest_sha: str,
    output_root: Path,
    repeat_index: int,
    run_identity: Mapping[str, Any],
    download_artifact: DownloadArtifact,
    actor: str = "task011-pilot",
) -> dict[str, Any]:
    """Run and verify the frozen four-render pilot from one report revision."""
    if not output_root.is_absolute():
        _fail("UNSAFE_OUTPUT_ROOT", "Pilot output root must be absolute.")
    if repeat_index not in (1, 2):
        _fail("INFRASTRUCTURE_ERROR", "repeat_index must be 1 or 2.")
    if not _is_sha256(source_manifest_sha):
        _fail("SOURCE_BINDING_MISMATCH", "Manifest SHA-256 is malformed.")
    assert_no_managed_artifacts(root=output_root, managed_paths=_managed_paths())

    report = report_service.create_report(
        project_id=project_id,
        project_version_id=project_version_id,
        report_type=ReportType.COLD_STORAGE_CONCEPT_DESIGN,
        actor=actor,
    )
    revision = report_service.generate_revision(report.id, actor)
    if revision.report_id != report.id:
        _fail("REPORT_REVISION_MISMATCH", "Generated revision report ID mismatch.")

    canonical_model = build_canonical_render_model(
        content=revision.content_json,
        report_id=report.id,
        revision_number=revision.revision_number,
        content_hash=revision.content_hash,
        generated_by=revision.generated_by,
        generated_at=revision.generated_at.isoformat(),
        template_code="cold_storage_concept_design",
        template_version="1.0.0",
        approval_snapshot=None,
    )
    section_keys = [section.section_key for section in canonical_model.sections]
    metrics = _canonical_metrics(canonical_model)

    started_at = datetime.now(UTC).isoformat()
    pilot_run = {
        "schema_version": PILOT_RESULT_SCHEMA_VERSION,
        "pilot_check_id": PILOT_CHECK_ID,
        "source_commit_sha": source_commit_sha,
        "source_manifest_sha": source_manifest_sha,
        **dict(run_identity),
        "repeat_index": repeat_index,
        "project_id": project_id,
        "project_version_id": project_version_id,
        "report_id": report.id,
        "report_revision_id": revision.id,
        "revision_number": revision.revision_number,
        "report_revision_content_hash": revision.content_hash,
        "report_type": report.report_type.value,
        "report_schema_version": revision.schema_version,
        "started_at": started_at,
    }
    atomic_write_json(path=output_root / "pilot-run.json", data=pilot_run)

    artifact_rows: list[dict[str, Any]] = []
    semantic_results: list[str] = []
    for locale, fmt in _RENDER_MATRIX:
        artifact = render_service.render(
            report_id=report.id,
            revision_number=revision.revision_number,
            format=fmt.value,
            template_version=None,
            mode="draft",
            actor=actor,
            locale=locale,
        )
        template = template_repository.get_template(artifact.template_id)
        if template is None:
            _fail(
                "TEMPLATE_PROVENANCE_MISMATCH",
                "Persisted template for artifact is missing.",
                template_id=artifact.template_id,
            )
        _verify_artifact_binding(
            artifact=artifact,
            report_id=report.id,
            revision=revision,
            locale=locale,
            fmt=fmt,
            template=template,
        )
        downloaded, download_headers = download_artifact(report.id, artifact.id, actor)
        downloaded_hash = _sha256(downloaded)
        if downloaded_hash != artifact.file_sha256 or len(downloaded) != artifact.file_size_bytes:
            _fail(
                "DOWNLOAD_INTEGRITY_MISMATCH",
                "Downloaded artifact bytes do not match persisted metadata.",
                artifact_id=artifact.id,
            )
        expected_headers = {
            "X-Content-SHA256": downloaded_hash,
            "X-Source-Content-Hash": revision.content_hash,
            "X-Report-Locale": locale.value,
            "X-Template-Locale": locale.value,
            "X-Translation-Catalog-Version": artifact.translation_catalog_version,
            "X-Translation-Catalog-Content-Hash": artifact.translation_catalog_content_hash,
            "X-Localized-Template-Content-Hash": artifact.localized_template_content_hash,
        }
        if any(download_headers.get(key) != value for key, value in expected_headers.items()):
            _fail(
                "DOWNLOAD_INTEGRITY_MISMATCH",
                "Download response header binding mismatch.",
                artifact_id=artifact.id,
            )

        checks = _semantic_checks(
            canonical_model=canonical_model,
            template=template,
            locale=locale,
            fmt=fmt,
            artifact_bytes=downloaded,
        )
        if checks["semantic_result"] != "PASS":
            _fail(
                "NUMERIC_SEMANTIC_MISMATCH",
                "Downloaded report failed section or numeric semantic verification.",
                locale=locale.value,
                format=fmt.value,
                checks=checks,
            )
        semantic_results.append(str(checks["semantic_result"]))

        artifact_dir = output_root / "artifacts" / locale.value / fmt.value
        metadata = {
            "schema_version": PILOT_RESULT_SCHEMA_VERSION,
            "artifact_id": artifact.id,
            "report_id": artifact.report_id,
            "report_revision_id": artifact.report_revision_id,
            "revision_number": artifact.revision_number,
            "format": artifact.format.value,
            "locale": artifact.locale.value,
            "template_locale": artifact.template_locale.value,
            "render_mode": artifact.render_manifest_json.get("render_mode"),
            "template_version": artifact.template_version,
            "template_content_hash": template.template_content_hash,
            "template_schema_version": template.schema_version,
            "source_content_hash": artifact.source_content_hash,
            "translation_catalog_version": artifact.translation_catalog_version,
            "translation_catalog_content_hash": artifact.translation_catalog_content_hash,
            "localized_template_content_hash": artifact.localized_template_content_hash,
            "artifact_status": artifact.status.value,
            "file_name": artifact.file_name,
            "file_size_bytes": artifact.file_size_bytes,
            "file_sha256": artifact.file_sha256,
            "download_headers": dict(download_headers),
            "integrity_result": "PASS",
        }
        atomic_write_bytes(path=artifact_dir / f"report.{fmt.value}", data=downloaded)
        atomic_write_json(path=artifact_dir / "artifact-metadata.json", data=metadata)
        atomic_write_json(path=artifact_dir / "semantic-checks.json", data=checks)
        artifact_rows.append(metadata)

    identities = {
        (
            item["report_id"],
            item["report_revision_id"],
            item["revision_number"],
            item["source_content_hash"],
        )
        for item in artifact_rows
    }
    if len(identities) != 1:
        _fail("REPORT_REVISION_MISMATCH", "Four renders do not share one revision.")
    if any(item["source_content_hash"] != revision.content_hash for item in artifact_rows):
        _fail("SOURCE_BINDING_MISMATCH", "Four renders do not share source content.")
    if not section_keys or not metrics:
        _fail(
            "REQUIRED_SECTION_MISSING",
            "Canonical report has no sections or numeric fields to verify.",
        )

    managed_hashes = {
        str(path.relative_to(output_root)): _sha256(path.read_bytes())
        for path in sorted(output_root.rglob("*"))
        if path.is_file() and path.name != "pilot-summary.json"
    }
    summary = {
        "schema_version": PILOT_RESULT_SCHEMA_VERSION,
        "pilot_check_id": PILOT_CHECK_ID,
        "source_commit_sha": source_commit_sha,
        "source_manifest_sha": source_manifest_sha,
        **dict(run_identity),
        "repeat_index": repeat_index,
        "started_at": started_at,
        "completed_at": datetime.now(UTC).isoformat(),
        "render_matrix": [
            {"locale": locale.value, "format": fmt.value, "mode": "draft"}
            for locale, fmt in _RENDER_MATRIX
        ],
        "source_binding_result": "PASS",
        "artifact_integrity_result": "PASS",
        "semantic_result": (
            "PASS" if all(result == "PASS" for result in semantic_results) else "FAIL"
        ),
        "overall_result": "PASS",
        "managed_file_sha256": managed_hashes,
    }
    atomic_write_json(path=output_root / "pilot-summary.json", data=summary)
    return summary


__all__ = [
    "PILOT_CHECK_ID",
    "PILOT_RESULT_SCHEMA_VERSION",
    "PilotVerificationError",
    "verify_multilingual_report_pilot",
]
