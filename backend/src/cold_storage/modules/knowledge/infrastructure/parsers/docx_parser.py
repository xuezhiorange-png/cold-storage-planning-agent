"""DOCX file parser — paragraphs and tables from python-docx."""

from __future__ import annotations

import io
import unicodedata

from cold_storage.modules.knowledge.domain.models import ParsedBlock
from cold_storage.modules.knowledge.infrastructure.parsers.base import (
    PARSER_VERSION,
    ParseResult,
    register_parser,
)

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[assignment]


class DocxParser:
    """Parse .docx files using python-docx.

    Extracts paragraphs and tables as ParsedBlock list.
    page_start/page_end are always None (python-docx does not provide page info).
    """

    name: str = "docx"

    def parse(self, content: bytes, filename: str) -> list[ParsedBlock]:
        """Parse a .docx file into structured blocks.

        Raises
        ------
        ImportError
            If python-docx is not installed.
        """
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing: pip install python-docx")

        buf = io.BytesIO(content)
        doc = DocxDocument(buf)

        blocks: list[ParsedBlock] = []
        order = 0
        para_index = 0
        table_index = 0
        section_path = ""

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                para = doc.paragraphs[para_index] if para_index < len(doc.paragraphs) else None
                para_index += 1
                if para is None:
                    continue
                text = unicodedata.normalize("NFKC", para.text)
                if not text.strip():
                    continue

                # Detect headings for section_path
                style = para.style
                style_name: str = style.name if style is not None else ""
                if style_name.startswith("Heading"):
                    section_path = text

                blocks.append(
                    ParsedBlock(
                        text=text,
                        block_type="heading" if style_name.startswith("Heading") else "paragraph",
                        section_path=section_path,
                        page_start=None,
                        page_end=None,
                        sheet_name=None,
                        paragraph_index=order,
                        source_order=order,
                        metadata={
                            "parser_version": PARSER_VERSION,
                            "style": style_name,
                        },
                    )
                )
                order += 1

            elif tag == "tbl":
                # Table
                table = doc.tables[table_index] if table_index < len(doc.tables) else None
                table_index += 1
                if table is None:
                    continue

                rows_text: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))

                table_text = "\n".join(rows_text)
                table_text = unicodedata.normalize("NFKC", table_text)
                if not table_text.strip():
                    continue

                blocks.append(
                    ParsedBlock(
                        text=table_text,
                        block_type="table",
                        section_path=section_path,
                        page_start=None,
                        page_end=None,
                        sheet_name=None,
                        row_start=1,
                        row_end=len(table.rows),
                        table_index=table_index - 1,
                        source_order=order,
                        metadata={
                            "parser_version": PARSER_VERSION,
                            "rows": len(table.rows),
                            "columns": len(table.columns) if table.rows else 0,
                        },
                    )
                )
                order += 1

        return blocks

    def parse_with_metadata(self, content: bytes, filename: str) -> ParseResult:
        """Parse and return ParseResult with blocks and metadata."""
        blocks = self.parse(content, filename)
        return ParseResult(blocks=blocks)


register_parser(".docx", DocxParser())
