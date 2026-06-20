from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook


@dataclass(frozen=True)
class ReportArtifact:
    report_id: str
    word_path: Path
    excel_path: Path


class ReportService:
    def generate(
        self,
        report_id: str,
        output_dir: Path,
        calculation_results: list[dict[str, object]],
    ) -> ReportArtifact:
        output_dir.mkdir(parents=True, exist_ok=True)
        word_path = output_dir / f"{report_id}.docx"
        excel_path = output_dir / f"{report_id}.xlsx"

        document = Document()
        document.add_heading("冷库规划设计方案书", level=1)
        document.add_paragraph("本报告为规划和概念设计辅助输出，需专业人员复核。")
        for item in calculation_results:
            document.add_heading(str(item.get("calculator_name", "计算结果")), level=2)
            document.add_paragraph(str(item.get("result", {})))
        document.save(str(word_path))

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "计算结果"
        sheet.append(["calculator_name", "result"])
        for item in calculation_results:
            sheet.append([str(item.get("calculator_name", "")), str(item.get("result", {}))])
        workbook.save(excel_path)

        return ReportArtifact(report_id=report_id, word_path=word_path, excel_path=excel_path)
