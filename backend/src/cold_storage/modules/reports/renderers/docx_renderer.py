"""DOCX renderer — produces a .docx file from a ReportRenderModel.

Uses python-docx.  No macros, no external resource loading, no template
expressions.  Font fallback: SimSun → Times New Roman.  A4 page size.

P0-5: Fully manifest-driven headers, footers, watermarks, fonts, and layout.
"""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# Register VML namespaces for watermark XML
from docx.oxml.ns import nsmap as _nsmap
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from cold_storage.modules.reports.domain.errors import TemplateManifestError

_nsmap["v"] = "urn:schemas-microsoft-com:vml"
_nsmap["o"] = "urn:schemas-microsoft-com:office:office"
_nsmap["w14"] = "http://schemas.microsoft.com/office/word/2010/wordml"

if TYPE_CHECKING:
    from cold_storage.modules.reports.domain.render_model import (
        RenderSection,
        RenderTable,
        ReportRenderModel,
    )

# ---------------------------------------------------------------------------
# Constants (defaults — overridden by template manifest)
# ---------------------------------------------------------------------------
_BODY_FONT = "SimSun"
_BODY_FONT_FALLBACK = "Times New Roman"
_HEADING_FONT = "Times New Roman"
_A4_WIDTH_PT = 21.0 * 28.3465  # A4 width in points
_A4_HEIGHT_PT = 29.7 * 28.3465  # A4 height in points
_MARGIN_PT = 2.0 * 28.3465  # 2 cm margins
_PT_TO_CM = 0.0352778

_BODY_FONT_SIZE = 10.5
_HEADING1_SIZE = 16
_HEADING2_SIZE = 14
_HEADING3_SIZE = 12


# ---------------------------------------------------------------------------
# Manifest Settings Loader (P0-5)
# ---------------------------------------------------------------------------


def _load_manifest_settings(
    model: ReportRenderModel,
) -> dict[str, Any]:
    """Extract rendering settings from template manifest with defaults.

    Supports both legacy and canonical TemplateManifest structures.
    Mirrors PDF renderer's _load_manifest_settings pattern.
    """
    render_settings = model.manifest.render_settings
    settings: dict[str, Any] = {}

    # Page settings
    page = render_settings.get("page", {})
    settings["page_width_pt"] = page.get("width_pt", _A4_WIDTH_PT)
    settings["page_height_pt"] = page.get("height_pt", _A4_HEIGHT_PT)
    settings["margin_top_pt"] = page.get("margin_top_pt", page.get("margin_pt", _MARGIN_PT))
    settings["margin_bottom_pt"] = page.get("margin_bottom_pt", page.get("margin_pt", _MARGIN_PT))
    settings["margin_left_pt"] = page.get("margin_left_pt", page.get("margin_pt", _MARGIN_PT))
    settings["margin_right_pt"] = page.get("margin_right_pt", page.get("margin_pt", _MARGIN_PT))

    # Default page orientation
    settings["orientation"] = page.get("orientation", "portrait")

    # Font settings
    fonts = render_settings.get("fonts", {})
    settings["body_font_name"] = fonts.get("body_name", _BODY_FONT)
    settings["body_font_size"] = fonts.get("body_size_pt", _BODY_FONT_SIZE)
    settings["heading_font_name"] = _HEADING_FONT
    settings["heading1_size"] = fonts.get("heading1_size_pt", _HEADING1_SIZE)
    settings["heading2_size"] = fonts.get("heading2_size_pt", _HEADING2_SIZE)
    settings["heading3_size"] = fonts.get("heading3_size_pt", _HEADING3_SIZE)
    settings["table_header_size"] = fonts.get("table_header_size_pt", 9.5)
    settings["table_body_size"] = fonts.get("table_body_size_pt", 9)
    settings["footer_size"] = fonts.get("footer_size_pt", 8)
    settings["header_size"] = fonts.get("header_size_pt", 8)

    # Header/footer text (canonical: header.left/right/center)
    header = render_settings.get("header", {})
    footer = render_settings.get("footer", {})
    settings["header_left"] = header.get("left", "")
    settings["header_center"] = header.get("center", "")
    settings["header_right"] = header.get("right", "")
    settings["footer_left"] = footer.get("left", "")
    settings["footer_center"] = footer.get("center", "")
    settings["footer_right"] = footer.get("right", "")

    # Draft watermark
    wm = render_settings.get("watermark", {})
    draft_wm = render_settings.get("draft_watermark", wm)
    settings["draft_watermark_text"] = draft_wm.get("text", "DRAFT")
    settings["draft_watermark_size"] = draft_wm.get("font_size_pt", draft_wm.get("size", 72))
    settings["draft_watermark_color"] = draft_wm.get("color", "#CCCCCC")
    settings["draft_watermark_opacity"] = draft_wm.get("opacity", 0.3)
    settings["draft_watermark_angle"] = draft_wm.get("angle", 45)

    # Empty section placeholder
    esb = render_settings.get("empty_section_behavior", {})
    if not isinstance(esb, dict):
        esb = {}
    settings["placeholder_text"] = esb.get("placeholder_text", {})
    settings["empty_section_behavior"] = esb

    # Table configs
    settings["table_configs"] = render_settings.get("tables", {})
    settings["table_repeat_header"] = render_settings.get("tables", {}).get("repeat_header", True)

    # Landscape sections
    settings["landscape_sections"] = render_settings.get("landscape_sections", [])
    if not settings["landscape_sections"]:
        settings["landscape_sections"] = render_settings.get("page", {}).get(
            "landscape_sections", []
        )

    return settings


# ---------------------------------------------------------------------------
# Variable Substitution
# ---------------------------------------------------------------------------


def _substitute_vars(text: str, meta: Any, page_num: int = 0) -> str:
    """Substitute template variables like {project_name} in text."""
    return (
        text.replace("{project_name}", getattr(meta, "project_name", "") or "")
        .replace("{report_type}", getattr(meta, "report_type", "") or "")
        .replace("{revision_number}", str(getattr(meta, "revision_number", "")))
        .replace("{generated_at}", (getattr(meta, "generated_at", "") or "")[:10])
        .replace("{content_hash_short}", getattr(meta, "content_hash_short", "") or "")
        .replace("{confidentiality}", getattr(meta, "confidentiality", "") or "")
        .replace("{page_number}", str(page_num))
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _set_cell_border(cell: Any, **kwargs: Any) -> None:
    """Set cell border on a table cell.

    Usage: ``_set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"})``
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            attrs = kwargs[edge]
            el = OxmlElement(f"w:{edge}")
            for k, v in attrs.items():
                el.set(qn(f"w:{k}"), str(v))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_run_font(
    run: Any,
    *,
    font_name: str = _BODY_FONT,
    size: Pt | None = None,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    """Configure font on a run with CJK fallback."""
    run.font.name = font_name
    # Set East Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), _BODY_FONT)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

    if size is not None:
        run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_tab_field_run(p: Any) -> None:
    """Add a tab character as a field run (for left/center/right layout)."""
    tab_run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    tab_run._r.append(fldChar)
    tab_run2 = p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.text = "\t"
    tab_run2._r.append(instrText)
    tab_run3 = p.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    tab_run3._r.append(fldChar2)


def _add_page_field(p: Any, text_template: str, font_size: int) -> None:
    """Add text with {page_number} replaced by a PAGE field.

    P0-5: Handles left/right footer positions that contain {page_number}.
    """
    parts = text_template.split("{page_number}")
    prefix = parts[0]
    suffix = parts[1] if len(parts) > 1 else ""
    if prefix:
        run_pre = p.add_run(prefix)
        _set_run_font(run_pre, size=Pt(font_size), color=RGBColor(0x80, 0x80, 0x80))
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run_field = p.add_run()
    run_field._r.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run_instr = p.add_run()
    run_instr._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_end = p.add_run()
    run_end._r.append(fldChar2)
    if suffix:
        run_suf = p.add_run(suffix)
        _set_run_font(run_suf, size=Pt(font_size), color=RGBColor(0x80, 0x80, 0x80))


def _add_footer(section: Any, settings: dict[str, Any], meta: Any, page_num: int) -> None:
    """Add manifest-driven footer with left/center/right text and variable substitution.

    P0-5: Replaces the hardcoded _add_page_number_footer().
    P0-7: Reads actual section dimensions for landscape-correct tab stops.
    {page_number} in left/right positions generates PAGE field.
    """
    footer_left = settings.get("footer_left", "")
    footer_center = settings.get("footer_center", "")
    footer_right = settings.get("footer_right", "")
    footer_size = settings.get("footer_size", 8)

    # If all empty, use legacy default
    if not footer_left and not footer_center and not footer_right:
        footer_center = f"\u2014 {page_num} \u2014"

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # For combined left+center+right, use tab stops
    has_multiple = sum(bool(x) for x in [footer_left, footer_center, footer_right]) > 1

    # P0-7: Calculate tab stop positions from ACTUAL section dimensions
    # (python-docx page_width returns EMU; 1 pt = 12700 EMU)
    actual_width_pt = section.page_width / 12700
    actual_margin_left = section.left_margin / 12700
    actual_margin_right = section.right_margin / 12700
    content_width_pt = actual_width_pt - actual_margin_left - actual_margin_right
    # Word tab stops in twips (1/20 pt)
    center_pos = int((content_width_pt / 2) * 20)
    right_pos = int(content_width_pt * 20)

    if has_multiple:
        # Add tab stops for left/center/right layout
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab_center = OxmlElement("w:tab")
        tab_center.set(qn("w:val"), "center")
        tab_center.set(qn("w:pos"), str(center_pos))
        tabs.append(tab_center)
        tab_right = OxmlElement("w:tab")
        tab_right.set(qn("w:val"), "right")
        tab_right.set(qn("w:pos"), str(right_pos))
        tabs.append(tab_right)
        pPr.append(tabs)

    first = True
    if footer_left:
        # P0-5: Handle {page_number} in left position as PAGE field
        if "{page_number}" in footer_left:
            _add_page_field(p, footer_left, footer_size)
        else:
            text = _substitute_vars(footer_left, meta, page_num)
            run = p.add_run(text)
            _set_run_font(run, size=Pt(footer_size), color=RGBColor(0x80, 0x80, 0x80))
        if has_multiple:
            _add_tab_field_run(p)
        first = False

    if footer_center:
        text = _substitute_vars(footer_center, meta, page_num)
        # Check if it contains {page_number} — if so, insert PAGE field
        if "{page_number}" in footer_center:
            _add_page_field(p, footer_center, footer_size)
        else:
            run = p.add_run(text)
            _set_run_font(run, size=Pt(footer_size), color=RGBColor(0x80, 0x80, 0x80))
        if has_multiple:
            _add_tab_field_run(p)
        first = False  # noqa: F841

    if footer_right:
        # P0-5: Handle {page_number} in right position as PAGE field
        if "{page_number}" in footer_right:
            _add_page_field(p, footer_right, footer_size)
        else:
            text = _substitute_vars(footer_right, meta, page_num)
            run = p.add_run(text)
            _set_run_font(run, size=Pt(footer_size), color=RGBColor(0x80, 0x80, 0x80))


def _add_header(section: Any, settings: dict[str, Any], meta: Any, page_num: int) -> None:
    """Add manifest-driven header with left/center/right text and variable substitution.

    P0-5: Replaces the hardcoded _add_header().
    P0-7: Reads actual section dimensions for landscape-correct tab stops.
    """
    header_left = settings.get("header_left", "")
    header_center = settings.get("header_center", "")
    header_right = settings.get("header_right", "")
    header_size = settings.get("header_size", 8)

    # If all empty, use legacy default
    if not header_left and not header_center and not header_right:
        pname = getattr(meta, "project_name", "")
        rtype = getattr(meta, "report_type", "")
        header_right = f"{pname} \u2014 {rtype}"

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    has_multiple = sum(bool(x) for x in [header_left, header_center, header_right]) > 1

    if has_multiple:
        # P0-7: Calculate tab stop positions from ACTUAL section dimensions
        # (python-docx page_width returns EMU; 1 pt = 12700 EMU)
        actual_width_pt = section.page_width / 12700
        actual_margin_left = section.left_margin / 12700
        actual_margin_right = section.right_margin / 12700
        h_content_width_pt = actual_width_pt - actual_margin_left - actual_margin_right
        h_center_pos = int((h_content_width_pt / 2) * 20)  # twips
        h_right_pos = int(h_content_width_pt * 20)

        # Add tab stops for left/center/right layout
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab_center = OxmlElement("w:tab")
        tab_center.set(qn("w:val"), "center")
        tab_center.set(qn("w:pos"), str(h_center_pos))
        tabs.append(tab_center)
        tab_right = OxmlElement("w:tab")
        tab_right.set(qn("w:val"), "right")
        tab_right.set(qn("w:pos"), str(h_right_pos))
        tabs.append(tab_right)
        pPr.append(tabs)

    if header_left:
        text = _substitute_vars(header_left, meta, page_num)
        run = p.add_run(text)
        _set_run_font(run, size=Pt(header_size), color=RGBColor(0x80, 0x80, 0x80))
        if has_multiple:
            _add_tab_field_run(p)

    if header_center:
        text = _substitute_vars(header_center, meta, page_num)
        run = p.add_run(text)
        _set_run_font(run, size=Pt(header_size), color=RGBColor(0x80, 0x80, 0x80))
        if has_multiple:
            _add_tab_field_run(p)

    if header_right:
        text = _substitute_vars(header_right, meta, page_num)
        run = p.add_run(text)
        _set_run_font(run, size=Pt(header_size), color=RGBColor(0x80, 0x80, 0x80))


def _add_draft_watermark(doc: Any, settings: dict[str, Any]) -> None:
    """Add draft watermark to every section header using VML/XML.

    P0-5: Adds watermark as a SEPARATE paragraph in each section's header,
    so it coexists with header text without overwriting it.
    P0-7: Uses actual section dimensions for watermark size.
    """
    wm_text = settings.get("draft_watermark_text", "DRAFT")
    if not wm_text:
        return

    wm_size = settings.get("draft_watermark_size", 72)
    wm_color_str = settings.get("draft_watermark_color", "#CCCCCC")
    wm_fillcolor = wm_color_str.lstrip("#")
    wm_opacity = settings.get("draft_watermark_opacity", 0.3)
    wm_angle = settings.get("draft_watermark_angle", 45)

    for section in doc.sections:
        header = section.header
        if not header.is_linked_to_previous:
            # Only add watermark to non-linked headers
            _add_vml_watermark(
                header,
                wm_text,
                wm_fillcolor,
                wm_opacity,
                wm_angle,
                wm_size,
                settings,
                section=section,
            )


def _add_vml_watermark(
    header: Any,
    text: str,
    fillcolor: str,
    opacity: float,
    angle: float,
    font_size: int,
    settings: dict[str, Any],
    section: Any = None,
) -> None:
    """Add VML-based watermark to the document header.

    P0-5: Adds watermark as a SEPARATE paragraph so it coexists with header text.
    P0-7: Uses actual section dimensions (EMU) for watermark width/height.
    Uses proper VML rotation via style attribute and o:opacity for transparency.
    """
    # P0-7: Compute watermark dimensions from actual section size
    if section is not None:
        actual_width_pt = section.page_width / 12700
        actual_height_pt = section.page_height / 12700
        wm_width_pt = actual_width_pt * 0.85  # ~85% of page width
        wm_height_pt = actual_height_pt * 0.30  # ~30% of page height
    else:
        wm_width_pt = 527.85
        wm_height_pt = 131.95

    # Add a NEW paragraph for the watermark (don't modify existing header text)
    wm_para = header.add_paragraph()
    wm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build VML shape for the watermark
    r = OxmlElement("w:r")
    r.set(qn("w14:paraId"), "10661244")
    r.set(qn("w14:textId"), "77777777")

    wm_para._p.append(r)

    # Add shape with VML
    pict = OxmlElement("w:pict")
    pict.set(qn("w14:shapeId"), "PowerPlusWaterMarkObject")

    shapetype = OxmlElement("v:shapetype")
    shapetype.set("id", "_x0000_t136")
    shapetype.set("coordsize", "21600,21600")
    shapetype.set(qn("o:spt"), "136")
    shapetype.set("adj", "10800")

    shape = OxmlElement("v:shape")
    shape.set("id", "PowerPlusWaterMarkObject")
    shape.set(qn("o:spid"), "_x0000_s2049")
    shape.set("type", "#_x0000_t136")
    # P0-7: VML style with rotation, using section-aware dimensions
    vml_style = (
        "position:absolute;"
        "margin-left:0;"
        "margin-top:0;"
        f"width:{wm_width_pt:.2f}pt;"
        f"height:{wm_height_pt:.2f}pt;"
        f"z-index:-251658752;"
        f"rotation:{int(angle)};"
        "mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        "mso-position-vertical-relative:margin;"
    )
    shape.set("style", vml_style)
    shape.set("fillcolor", fillcolor)
    shape.set("strokecolor", "none")
    shape.set("strokeweight", "0")

    # Fill element with opacity — use o:opacity for VML
    fill = OxmlElement("v:fill")
    fill.set("type", "solid")
    fill.set("color", fillcolor)
    fill.set("opacity", str(opacity))
    shape.append(fill)

    # Office namespace opacity attribute (o:opacity)
    o_opacity = OxmlElement("o:opacity")
    o_opacity.set("value", str(int(opacity * 100)) + "%")
    shape.append(o_opacity)

    # TextPath element
    textpath = OxmlElement("v:textpath")
    textpath.set("on", "t")
    shape.append(textpath)

    # Lock aspect ratio
    lock = OxmlElement("v:lock")
    lock.set("type", "none")
    shape.append(lock)

    pict.append(shape)

    # Create the actual text run inside the shape
    # Word uses <w:t> with font properties inside the shape
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")

    # For proper rendering, create a textbox inside the shape
    textbox = OxmlElement("v:textbox")
    textbox.set("style", "mso-fit-shape-to-text:false")
    txContent = OxmlElement("w:txbxContent")

    # Create paragraph with text
    txP = OxmlElement("w:p")
    txPPR = OxmlElement("w:pPr")
    txPStyle = OxmlElement("w:pStyle")
    txPStyle.set(qn("w:val"), "a4")
    txPPR.append(txPStyle)
    txAlign = OxmlElement("w:jc")
    txAlign.set(qn("w:val"), "center")
    txPPR.append(txAlign)
    txP.append(txPPR)

    txR = OxmlElement("w:r")
    txRPr = OxmlElement("w:rPr")
    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "SimHei")
    rFonts.set(qn("w:hAnsi"), "SimHei")
    txRPr.append(rFonts)
    # Size
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))  # half-points
    txRPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    txRPr.append(szCs)
    # Color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), fillcolor)
    txRPr.append(color)
    # Bold
    b = OxmlElement("w:b")
    txRPr.append(b)
    txR.append(txRPr)
    txT = OxmlElement("w:t")
    txT.text = text
    txT.set(qn("xml:space"), "preserve")
    txR.append(txT)
    txP.append(txR)
    txContent.append(txP)
    textbox.append(txContent)
    shape.append(textbox)

    pict.append(shapetype)
    r.append(pict)


# ---------------------------------------------------------------------------
# DOCX Renderer
# ---------------------------------------------------------------------------


class DocxRenderer:
    """Render a ReportRenderModel to DOCX bytes."""

    def render(self, model: ReportRenderModel, *, is_draft: bool = False) -> bytes:
        """Render the model to DOCX bytes.

        Parameters
        ----------
        model:
            The complete render model.
        is_draft:
            If True, add a DRAFT watermark.

        Returns
        -------
        bytes
            The raw .docx file content.
        """
        # P0-5: Load manifest settings (canonical structure)
        settings = _load_manifest_settings(model)

        meta = model.metadata

        # Page size from manifest
        page_width_cm = settings["page_width_pt"] * _PT_TO_CM
        page_height_cm = settings["page_height_pt"] * _PT_TO_CM
        margin_top_cm = settings["margin_top_pt"] * _PT_TO_CM
        margin_bottom_cm = settings["margin_bottom_pt"] * _PT_TO_CM
        margin_left_cm = settings["margin_left_pt"] * _PT_TO_CM
        margin_right_cm = settings["margin_right_pt"] * _PT_TO_CM

        body_font_name = settings["body_font_name"]
        body_font_size = settings["body_font_size"]
        heading_font_name = settings["heading_font_name"]

        doc: Any = Document()

        # ---- Page setup (from manifest) ----
        default_orientation = settings.get("orientation", "portrait")
        for section in doc.sections:
            section.page_width = Cm(page_width_cm)
            section.page_height = Cm(page_height_cm)
            section.top_margin = Cm(margin_top_cm)
            section.bottom_margin = Cm(margin_bottom_cm)
            section.left_margin = Cm(margin_left_cm)
            section.right_margin = Cm(margin_right_cm)
            section.orientation = (
                WD_ORIENT.LANDSCAPE if default_orientation == "landscape" else WD_ORIENT.PORTRAIT
            )

        # ---- Default style ----
        style = doc.styles["Normal"]
        style.font.name = body_font_name
        style.font.size = Pt(body_font_size)
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.append(rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), body_font_name)

        # ---- Heading styles (from manifest) ----
        heading_sizes = {
            1: settings["heading1_size"],
            2: settings["heading2_size"],
            3: settings["heading3_size"],
        }
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                hs = doc.styles[style_name]
                hs.font.name = heading_font_name
                hs.font.size = Pt(heading_sizes.get(level, [16, 14, 12][level - 1]))
                hs.font.bold = True
                hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # ---- Document properties (metadata) ----
        cp = doc.core_properties
        cp.title = meta.project_name or "Report"
        cp.subject = meta.report_type
        cp.comments = f"Revision {meta.revision_number}"
        cp.author = meta.generated_by
        # Custom properties via XML
        docProps = doc.element.find(qn("w:docProps"))
        if docProps is None:
            docProps = OxmlElement("w:docProps")
            doc.element.append(docProps)

        def _add_custom_prop(name: str, value: str) -> None:
            prop = OxmlElement("w:customProp")
            prop.set(qn("w:name"), name)
            val_el = OxmlElement("w:str")
            val_el.set(qn("w:val"), value)
            prop.append(val_el)
            docProps.append(prop)

        _add_custom_prop("SourceContentHash", meta.content_hash)
        _add_custom_prop("TemplateVersion", meta.template_version)

        # ---- Cover page ----
        # Blank lines for centering
        for _ in range(6):
            doc.add_paragraph("")

        # Project name
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_name.add_run(meta.project_name or "\u9879\u76ee\u62a5\u544a")
        _set_run_font(
            run,
            font_name=heading_font_name,
            size=Pt(26),
            bold=True,
            color=RGBColor(0x1F, 0x38, 0x64),
        )

        # Report type
        p_type = doc.add_paragraph()
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_type.add_run(meta.report_type)
        _set_run_font(
            run,
            font_name=heading_font_name,
            size=Pt(18),
            color=RGBColor(0x40, 0x40, 0x40),
        )

        # Version line
        p_ver = doc.add_paragraph()
        p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_at = meta.generated_at if meta.generated_at else ""
        date_display = generated_at[:10] if len(generated_at) >= 10 else generated_at
        ver_str = f"\u7248\u672c {meta.revision_number}  |  {date_display}"
        run = p_ver.add_run(ver_str)
        _set_run_font(run, size=Pt(12), color=RGBColor(0x60, 0x60, 0x60))

        # ---- Document control info ----
        doc.add_page_break()

        p_title = doc.add_paragraph()
        run = p_title.add_run("\u6587\u4ef6\u63a7\u5236\u4fe1\u606f")
        _set_run_font(run, font_name=heading_font_name, size=Pt(14), bold=True)

        hash_val = meta.content_hash
        hash_display = hash_val[:16] + "\u2026" if len(hash_val) > 16 else hash_val
        control_items = [
            ("\u5185\u5bb9\u54c8\u5e0c", hash_display),
            ("\u6a21\u677f\u7248\u672c", meta.template_version),
            ("\u751f\u6210\u8005", meta.generated_by),
            ("\u751f\u6210\u65f6\u95f4", meta.generated_at),
            ("\u4fee\u8ba2\u53f7", str(meta.revision_number)),
        ]
        for label, value in control_items:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}\uff1a")
            _set_run_font(run_label, bold=True, size=Pt(10))
            run_val = p.add_run(value)
            _set_run_font(run_val, size=Pt(10))

        doc.add_page_break()

        # ---- Sections ----
        table_configs = settings.get("table_configs", {})
        landscape_sections = settings.get("landscape_sections", [])
        current_orientation = default_orientation
        for render_section in model.sections:
            # P0-4: Orientation resolution order:
            #   1. tables[section_key].orientation (from manifest)
            #   2. landscape_sections list
            #   3. page.orientation (default)
            section_config = table_configs.get(render_section.section_key, {})
            per_section_orientation = section_config.get("orientation", "")
            if per_section_orientation in ("landscape", "portrait"):
                target_orientation = per_section_orientation
            elif render_section.section_key in landscape_sections:
                target_orientation = "landscape"
            else:
                target_orientation = default_orientation

            if target_orientation != current_orientation:
                # Switch orientation — create new Word section
                new_section = doc.add_section()
                if target_orientation == "landscape":
                    new_section.orientation = WD_ORIENT.LANDSCAPE
                    new_section.page_width = Cm(page_height_cm)  # swapped
                    new_section.page_height = Cm(page_width_cm)  # swapped
                else:
                    new_section.orientation = WD_ORIENT.PORTRAIT
                    new_section.page_width = Cm(page_width_cm)
                    new_section.page_height = Cm(page_height_cm)
                new_section.top_margin = Cm(margin_top_cm)
                new_section.bottom_margin = Cm(margin_bottom_cm)
                new_section.left_margin = Cm(margin_left_cm)
                new_section.right_margin = Cm(margin_right_cm)
                current_orientation = target_orientation

            self._render_section(doc, render_section, settings, meta)

        # ---- Footer with page numbers (P0-5: manifest-driven) ----
        for idx, doc_section in enumerate(doc.sections):
            _add_footer(doc_section, settings, meta, idx + 1)

        # ---- Header (P0-5: manifest-driven) ----
        for idx, doc_section in enumerate(doc.sections):
            _add_header(doc_section, settings, meta, idx + 1)

        # ---- Watermark (P0-5: manifest-driven) ----
        if is_draft:
            _add_draft_watermark(doc, settings)

        # ---- Serialize ----
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Section rendering
    # ------------------------------------------------------------------

    def _render_section(
        self,
        doc: Any,
        section: RenderSection,
        settings: dict[str, Any] | None = None,
        meta: Any = None,
    ) -> None:
        """Render a single section into the document.

        P0-5: Uses manifest-driven placeholder text for empty sections.
        """
        if settings is None:
            settings = {}
        body_font_size = settings.get("body_font_size", _BODY_FONT_SIZE)

        if section.is_empty:
            doc.add_heading(section.title, level=section.level)
            # P0-5: Use manifest-driven placeholder text
            esb = settings.get("empty_section_behavior", {})
            placeholder_map = esb.get("placeholder_text", {})
            if not isinstance(placeholder_map, dict):
                placeholder_map = {}
            default_text = self._empty_reason_text(section.empty_reason)
            placeholder = placeholder_map.get(section.empty_reason, default_text)
            p = doc.add_paragraph(f"\uff08{placeholder}\uff09")
            target_run = p.runs[0] if p.runs else p.add_run("")
            _set_run_font(
                target_run,
                size=Pt(body_font_size),
                color=RGBColor(0x99, 0x99, 0x99),
            )
            return

        # Section heading
        doc.add_heading(section.title, level=section.level)

        if section.content_type == "text" and section.text:
            self._render_text_block(doc, section.text, body_font_size)
        elif section.content_type == "metrics" and section.metrics:
            for metric in section.metrics:
                p = doc.add_paragraph()
                run = p.add_run(f"{metric.label}: {metric.display_value} {metric.unit}".strip())
                _set_run_font(run, size=Pt(body_font_size))
            if section.number:
                self._render_number(doc, section, body_font_size)
        elif section.content_type == "number" and section.number:
            self._render_number(doc, section, body_font_size)
        elif section.content_type == "table" and section.table:
            if section.text:
                self._render_text_block(doc, section.text, body_font_size)
            self._render_table(doc, section.table, section.section_key, settings)
        elif section.content_type == "finding":
            if section.text:
                self._render_text_block(doc, section.text, body_font_size)
            if section.table:
                self._render_table(doc, section.table, section.section_key, settings)

        # Render paragraphs
        if section.paragraphs:
            for para in section.paragraphs:
                p = doc.add_paragraph()
                run = p.add_run(para)
                _set_run_font(run, size=Pt(body_font_size))

        # Render citations as numbered footnotes
        if section.citations:
            for idx, cite in enumerate(section.citations, 1):
                tool = cite.get("tool_name", "")
                src = cite.get("source_id", "")
                cite_text = f"[{idx}] {tool} \u2014 {src}"
                p = doc.add_paragraph()
                run = p.add_run(cite_text)
                _set_run_font(run, size=Pt(9), color=RGBColor(0x60, 0x60, 0x60))

    def _render_text_block(
        self,
        doc: Any,
        text: str,
        body_font_size: float = _BODY_FONT_SIZE,
    ) -> None:
        """Render a text block, preserving line breaks as paragraphs."""
        for line in text.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            _set_run_font(run, size=Pt(body_font_size))

    def _render_number(
        self, doc: Any, section: RenderSection, body_font_size: float = _BODY_FONT_SIZE
    ) -> None:
        """Render a number field with its value and unit."""
        num = section.number
        if num is not None:
            p = doc.add_paragraph()
            run = p.add_run(f"{num.display} {num.unit}")
            _set_run_font(run, size=Pt(body_font_size + 0.5), bold=True)
        if section.text:
            self._render_text_block(doc, section.text, body_font_size)

    def _render_table(
        self,
        doc: Any,
        table: RenderTable,
        section_key: str = "",
        settings: dict[str, Any] | None = None,
    ) -> None:
        """Render a RenderTable as a Word table.

        P0-5: Reads per-section table config from manifest for width_ratio,
        repeat_header, unit_row, and column alignment.
        P0-4: Table width calculated within content area (margins subtracted).
        """
        if settings is None:
            settings = {}
        table_configs = settings.get("table_configs", {})
        table_body_size = settings.get("table_body_size", 9)
        table_header_size = settings.get("table_header_size", 9.5)

        # Get section-specific table config
        section_config = table_configs.get(section_key, {})
        column_configs = section_config.get("columns", [])
        repeat_header = section_config.get("repeat_header", True)
        # P0-4: unit_row from manifest config; only create if config says true AND data has units
        show_unit_row = section_config.get("unit_row", True)

        num_cols = len(table.headers)
        num_rows = len(table.rows) + 1  # +1 for header
        has_unit_row = show_unit_row and table.unit_row and any(u for u in table.unit_row)
        if has_unit_row:
            num_rows += 1  # +1 for unit row

        # P0-4: Validate column count
        if column_configs and len(column_configs) != num_cols:
            raise TemplateManifestError(
                f"columns_config count ({len(column_configs)}) != headers count "
                f"({num_cols}) for section_key={section_key!r}"
            )

        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = "Table Grid"

        # P0-4: Compute column widths within content area (subtract margins)
        margin_left_pt = settings.get("margin_left_pt", _MARGIN_PT)
        margin_right_pt = settings.get("margin_right_pt", _MARGIN_PT)
        content_width_pt = (
            settings.get("page_width_pt", _A4_WIDTH_PT) - margin_left_pt - margin_right_pt
        )
        content_width_cm = content_width_pt * _PT_TO_CM

        # P0-5: Apply width_ratio from manifest to tblGrid (normalized)
        if column_configs:
            tblGrid = word_table._tbl.find(qn("w:tblGrid"))
            if tblGrid is None:
                tblGrid = OxmlElement("w:tblGrid")
                word_table._tbl.insert(0, tblGrid)

            # P0-4: Normalize width_ratio values
            ratios = [col.get("width_ratio", 1.0) for col in column_configs]
            total = sum(ratios)
            if total <= 0 or any(r < 0 for r in ratios):
                raise TemplateManifestError(
                    f"Invalid width_ratio values for section_key={section_key!r}: "
                    f"total={total}, ratios={ratios}"
                )
            norm = [r / total for r in ratios]

            # Convert to EMU for Word (1 cm = 360000 EMU)
            for i, gridCol in enumerate(tblGrid.findall(qn("w:gridCol"))):
                col_width_cm = content_width_cm * norm[i]
                gridCol.set(qn("w:w"), str(int(col_width_cm * 360000)))

        # Header row
        for col_idx, header in enumerate(table.headers):
            cell = word_table.rows[0].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # P0-5: Apply column alignment from manifest
            if col_idx < len(column_configs):
                align = column_configs[col_idx].get("align", "left")
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(header)
            _set_run_font(run, size=Pt(table_header_size), bold=True)
            # Gray background for header
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "D9E2F3")
            tcPr.append(shading)

        # P0-5: Mark header row with <w:tblHeader> for repeating on page breaks
        if repeat_header:
            header_tr = word_table.rows[0]._tr
            header_trPr = header_tr.get_or_add_trPr()
            tblHeader = OxmlElement("w:tblHeader")
            tblHeader.set(qn("w:val"), "true")
            header_trPr.append(tblHeader)

        # Unit row (if present)
        row_offset = 1
        if has_unit_row:
            for col_idx, unit in enumerate(table.unit_row):
                cell = word_table.rows[1].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(f"({unit})" if unit else "")
                _set_run_font(run, size=Pt(9), color=RGBColor(0x60, 0x60, 0x60))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_offset = 2

        # Data rows
        for row_idx, row_data in enumerate(table.rows):
            # P0-5: Mark data row with <w:cantSplit> to prevent mid-row breaks
            tr = word_table.rows[row_idx + row_offset]._tr
            trPr = tr.get_or_add_trPr()
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)

            for col_idx, cell_data in enumerate(row_data):
                word_cell = word_table.rows[row_idx + row_offset].cells[col_idx]
                word_cell.text = ""
                p = word_cell.paragraphs[0]
                # P0-4: Cell alignment priority: cell.align > column_align > "left"
                if cell_data.align is not None:
                    align = cell_data.align
                elif col_idx < len(column_configs):
                    align = column_configs[col_idx].get("align", "left")
                else:
                    align = "left"
                if align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_data.value)
                _set_run_font(run, size=Pt(table_body_size))

    @staticmethod
    def _empty_reason_text(reason: str) -> str:
        """Human-readable empty reason."""
        reasons = {
            "not_provided": "\u8be5\u90e8\u5206\u6570\u636e\u672a\u63d0\u4f9b",
            "not_calculated": "\u8be5\u90e8\u5206\u5c1a\u672a\u8ba1\u7b97",
        }
        return reasons.get(reason, "\u8be5\u90e8\u5206\u5185\u5bb9\u4e0d\u53ef\u7528")
